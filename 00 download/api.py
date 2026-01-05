"""
CIP Backend API
Flask REST API for Contract Intelligence Platform
Integrates orchestrator, config, and database operations
"""

import json
import sqlite3
from pathlib import Path
from datetime import datetime
from typing import Dict, Optional
from dataclasses import asdict

from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from werkzeug.utils import secure_filename
from werkzeug.exceptions import HTTPException

# Import Anthropic for Claude skill integration
from anthropic import Anthropic

# Import CIP components
try:
    from config import (
        ANTHROPIC_API_KEY,
        DEFAULT_MODEL,
        CONTRACTS_DB,
        REPORTS_DB,
        UPLOAD_DIRECTORY,
        SUPPORTED_FORMATS,
        MAX_FILE_SIZE_BYTES,
        API_HOST,
        API_PORT,
        DEBUG_MODE,
        LOG_LEVEL,
        get_prompt_composer_config
    )
    from prompt_composer import PromptComposer, create_composer
    from orchestrator import (
        ContractContext,
        RiskAssessment,
        create_orchestrator,
        # RA-ERR-10054: Error handling classes
        AnalyzeError,
        AnalyzeNetworkError,
        AnalyzeAuthError,
        AnalyzePayloadError,
        AnalyzeInternalError,
        classify_analyze_error
    )
except ImportError as e:
    print(f"Import error: {e}")
    print("Make sure config.py and orchestrator.py are in the same directory")
    raise

# Setup comprehensive logging
from logger_config import setup_logging, setup_api_logging, log_user_action
from extraction_service import extract_contract_metadata, find_related_contracts, check_dependencies, extract_text_from_docx

logger = setup_logging(__name__, LOG_LEVEL)
api_logger = setup_api_logging()

# ============================================================================
# FLASK APP INITIALIZATION
# ============================================================================

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = MAX_FILE_SIZE_BYTES
app.config['UPLOAD_FOLDER'] = str(UPLOAD_DIRECTORY)

# Enable CORS for frontend integration
CORS(app, resources={
    r"/api/*": {
        "origins": ["http://localhost:8501", "http://127.0.0.1:8501"],
        "methods": ["GET", "POST", "PUT", "DELETE", "OPTIONS"],
        "allow_headers": ["Content-Type", "Authorization"]
    }
})

# Initialize orchestrator
orchestrator = None

try:
    orchestrator = create_orchestrator()
    logger.info("Contract Orchestrator initialized successfully")
except Exception as e:
    logger.error(f"Failed to initialize orchestrator: {e}")
    orchestrator = None

# Initialize Claude client for skill integration
claude_client = None

try:
    if ANTHROPIC_API_KEY:
        claude_client = Anthropic(api_key=ANTHROPIC_API_KEY)
        logger.info("Claude client initialized for skill integration")
    else:
        logger.warning("No Anthropic API key found - Claude skill enhancements disabled")
except Exception as e:
    logger.error(f"Failed to initialize Claude client: {e}")
    claude_client = None

# Initialize Prompt Composer
prompt_composer = None

try:
    composer_config = get_prompt_composer_config()
    prompt_composer = create_composer(composer_config)
    logger.info(f"Prompt Composer initialized - Deck v{prompt_composer.deck_version} with {len(prompt_composer.deck.get('cards', {}))} patterns")
except Exception as e:
    logger.error(f"Failed to initialize Prompt Composer: {e}")
    prompt_composer = None


# ============================================================================
# HELPER FUNCTIONS
# ============================================================================

def get_db_connection(db_path: Optional[str] = None):
    """Get database connection"""
    db_path = db_path or str(CONTRACTS_DB)
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA foreign_keys = ON")
    return conn


def query_distinct(conn, table: str, column: str) -> list:
    """Query distinct values from a table column"""
    try:
        cursor = conn.execute(f"SELECT DISTINCT {column} FROM {table} WHERE {column} IS NOT NULL ORDER BY {column}")
        return [row[0] for row in cursor.fetchall()]
    except:
        return []


def allowed_file(filename: str) -> bool:
    """Check if file extension is allowed"""
    return Path(filename).suffix.lower() in SUPPORTED_FORMATS


def save_uploaded_file(file) -> Path:
    """Save uploaded file to disk"""
    filename = secure_filename(file.filename)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    unique_filename = f"{timestamp}_{filename}"
    file_path = UPLOAD_DIRECTORY / unique_filename
    file.save(str(file_path))
    logger.info(f"File saved: {file_path}")
    return file_path


def format_risk_assessment(assessment: RiskAssessment) -> Dict:
    """Format RiskAssessment for JSON response"""
    return {
        'contract_id': assessment.contract_id,
        'overall_risk': assessment.overall_risk,
        'confidence_score': assessment.confidence_score,
        'analysis_date': assessment.analysis_date,
        'dealbreakers': [asdict(item) for item in assessment.dealbreakers],
        'critical_items': [asdict(item) for item in assessment.critical_items],
        'important_items': [asdict(item) for item in assessment.important_items],
        'standard_items': [asdict(item) for item in assessment.standard_items],
        'context': asdict(assessment.context),
        # v1.6 additions
        'risk_by_category': assessment.risk_by_category or {},
        'clauses_reviewed': assessment.clauses_reviewed,
        'clauses_flagged': assessment.clauses_flagged,
        # v1.7 additions
        'severity_counts': assessment.severity_counts or {},
        'low_risk_clauses': assessment.low_risk_clauses or []
    }


def validate_claude_available():
    """Check if Claude API is available and return error tuple if not"""
    if not ANTHROPIC_API_KEY:
        error_msg = {
            'error': 'Claude API not configured',
            'details': 'ANTHROPIC_API_KEY environment variable not set',
            'action': 'Set ANTHROPIC_API_KEY in your environment or .env file'
        }
        return error_msg, 503

    if not claude_client:
        error_msg = {
            'error': 'Claude client not initialized',
            'details': 'Failed to initialize Anthropic client',
            'action': 'Check API key validity and network connection'
        }
        return error_msg, 503

    return None


def validate_orchestrator_available():
    """Check if orchestrator is available and return error tuple if not"""
    if not orchestrator:
        error_msg = {
            'error': 'Analysis service not available',
            'details': 'Contract orchestrator failed to initialize',
            'action': 'Check server logs for orchestrator initialization errors'
        }
        return error_msg, 503

    return None


def safe_db_connection(db_path: Optional[str] = None):
    """Get database connection with error handling"""
    try:
        return get_db_connection(db_path), None
    except sqlite3.Error as e:
        logger.error(f"Database connection failed: {e}")
        error_msg = {
            'error': 'Database connection failed',
            'details': str(e),
            'action': 'Check database file exists and is not corrupted'
        }
        return None, (error_msg, 500)


def get_contract_metadata(contract_id: int) -> Optional[Dict]:
    """Fetch contract metadata from database for skill enhancement"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        cursor.execute("""
            SELECT id, filename, position, leverage, narrative,
                   contract_type, parties, effective_date
            FROM contracts
            WHERE id = ?
        """, (contract_id,))

        row = cursor.fetchone()
        conn.close()

        if row:
            return dict(row)
        return None

    except Exception as e:
        logger.error(f"Failed to fetch contract metadata: {e}")
        return None


def enhance_comparison_with_claude(comparison_results: Dict, v1_contract: Dict, v2_contract: Dict) -> Dict:
    """
    Enhance Python comparison results with Claude's context-aware analysis.

    Args:
        comparison_results: Dict from compare_documents.py
        v1_contract: Dict with contract metadata from database
        v2_contract: Dict with contract metadata from database

    Returns:
        Enhanced comparison results with reasoning and context-aware classification
    """
    if not claude_client:
        logger.warning("Claude client not available - skipping skill enhancement")
        comparison_results['claude_enhanced'] = False
        comparison_results['enhancement_error'] = "Claude client not initialized"
        return comparison_results

    try:
        # Prepare input for Claude skill
        skill_input = {
            "changes": comparison_results.get('changes', []),
            "v1_context": {
                "contract_id": v1_contract.get('id'),
                "filename": v1_contract.get('filename', 'Unknown'),
                "position": v1_contract.get('position', 'Unknown'),
                "leverage": v1_contract.get('leverage', 'Unknown'),
                "narrative": v1_contract.get('narrative', '')
            },
            "v2_context": {
                "contract_id": v2_contract.get('id'),
                "filename": v2_contract.get('filename', 'Unknown'),
                "position": v2_contract.get('position', 'Unknown'),
                "leverage": v2_contract.get('leverage', 'Unknown'),
                "narrative": v2_contract.get('narrative', '')
            }
        }

        # Load skill prompt from file
        skill_path = Path(__file__).parent.parent / ".claude" / "skills" / "contract-comparison" / "skill.md"

        if not skill_path.exists():
            logger.error(f"Skill file not found: {skill_path}")
            comparison_results['claude_enhanced'] = False
            comparison_results['enhancement_error'] = "Skill file not found"
            return comparison_results

        with open(skill_path, 'r', encoding='utf-8') as f:
            skill_prompt = f.read()

        # Construct Claude prompt with emphasis on cumulative analysis
        full_prompt = f"""{skill_prompt}

# Analysis Request

Please analyze these contract changes using the context-aware classification logic defined above.

## Input Data

```json
{json.dumps(skill_input, indent=2)}
```

## CRITICAL Instructions - Optimized for Quality

**STEP 1: CUMULATIVE PATTERN ANALYSIS (KEY QUALITY DIFFERENTIATOR)**
Before analyzing individual changes, review ALL {len(skill_input['changes'])} changes together:
- Identify patterns: Do multiple changes weaken the same area?
- Note compound effects: Changes that together create significant risk
- Recognize systematic shifts: Is power/risk transferring from one party to another?

**STEP 2: DETAILED PER-CHANGE ANALYSIS**
For EACH change, provide:
1. Position-aware impact classification
2. Detailed reasoning (3-4 sentences) explaining business consequences
3. Specific risk factors
4. Actionable recommendation
5. **Cumulative notes**: How this change compounds with others (if applicable)

**STEP 3: STRATEGIC EXECUTIVE SUMMARY**
Generate executive summary that includes:
- Impact level counts
- **Key patterns identified** (most important output!)
- Top 3-5 risks (prioritized by cumulative severity)
- Overall strategic recommendation considering patterns + leverage

**STEP 4: RETURN COMPLETE JSON**
Return results matching output schema with:
- All original fields preserved
- Enhanced fields added (impact, reasoning, risk_factors, cumulative_notes)
- Executive summary with pattern analysis
- Valid, properly-escaped JSON structure

CRITICAL JSON REQUIREMENTS:
1. Return ONLY valid JSON - no markdown code blocks, no extra text
2. Ensure ALL string values are properly JSON-escaped:
   - Backslashes: \\
   - Quotes: \"
   - Newlines: \\n
   - Tabs: \\t
   - Special characters (curly quotes, etc.) must be escaped as Unicode: \\uXXXX
3. The v1_content and v2_content fields contain contract text with quotes and special characters - these MUST be properly escaped
4. Test that your output is valid JSON before returning"""

        logger.info("Calling Claude skill for comparison enhancement...")

        # Call Claude API with optimized parameters for quality
        response = claude_client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=16000,      # 2x larger for detailed per-change analysis
            temperature=0,          # Consistent classifications across changes
            messages=[{"role": "user", "content": full_prompt}]
        )

        # Extract and parse response
        response_text = response.content[0].text.strip()

        logger.info(f"Claude response received ({len(response_text)} chars)")

        # Multiple extraction strategies for robustness
        enhanced_data = None
        parse_errors = []

        # Strategy 1: Remove markdown code blocks if present
        cleaned_text = response_text
        if cleaned_text.startswith("```json"):
            cleaned_text = cleaned_text[7:]
        if cleaned_text.startswith("```"):
            cleaned_text = cleaned_text[3:]
        if cleaned_text.endswith("```"):
            cleaned_text = cleaned_text[:-3]
        cleaned_text = cleaned_text.strip()

        # Strategy 1.5: Basic JSON repair for common issues
        # Note: This is a best-effort fix and may not catch all issues
        def attempt_json_repair(text):
            """Attempt to fix common JSON formatting issues"""
            # Don't modify - this is too risky and could break valid JSON
            # Just return as-is for now
            return text

        cleaned_text = attempt_json_repair(cleaned_text)

        try:
            enhanced_data = json.loads(cleaned_text)
            logger.info("JSON parsed successfully using markdown removal")
        except json.JSONDecodeError as e:
            parse_errors.append(f"Markdown removal failed: {e}")

            # Strategy 2: Find outermost JSON braces (handles text before/after)
            first_brace = cleaned_text.find('{')
            last_brace = cleaned_text.rfind('}')

            if first_brace != -1 and last_brace != -1 and last_brace > first_brace:
                potential_json = cleaned_text[first_brace:last_brace + 1]
                try:
                    enhanced_data = json.loads(potential_json)
                    logger.info("JSON parsed successfully using brace extraction")
                except json.JSONDecodeError as e2:
                    parse_errors.append(f"Brace extraction failed: {e2}")

            if not enhanced_data:
                # Strategy 3: Save debug info and try one more time with original response
                debug_file = Path(__file__).parent.parent / "data" / "reports" / "claude_response_debug.txt"
                try:
                    with open(debug_file, 'w', encoding='utf-8') as f:
                        f.write(f"=== ORIGINAL RESPONSE ({len(response_text)} chars) ===\n")
                        f.write(response_text)
                        f.write("\n\n=== CLEANED TEXT ===\n")
                        f.write(cleaned_text)
                        f.write("\n\n=== PARSE ERRORS ===\n")
                        f.write('\n'.join(parse_errors))

                    logger.error(f"All JSON parsing strategies failed. Debug info saved to: {debug_file}")
                except Exception as debug_error:
                    logger.error(f"Could not save debug file: {debug_error}")

                # Try original response_text as last resort
                first_brace_orig = response_text.find('{')
                last_brace_orig = response_text.rfind('}')

                if first_brace_orig != -1 and last_brace_orig != -1:
                    try:
                        enhanced_data = json.loads(response_text[first_brace_orig:last_brace_orig + 1])
                        logger.info("JSON parsed using original text extraction")
                    except json.JSONDecodeError as e3:
                        parse_errors.append(f"Original text extraction failed: {e3}")
                        logger.error(f"All parsing strategies exhausted: {'; '.join(parse_errors)}")
                        raise json.JSONDecodeError("All parsing strategies failed", response_text, 0)

        if not enhanced_data:
            raise ValueError("Failed to extract valid JSON from Claude response")

        # Merge enhanced data with original results
        # IMPORTANT: Merge enhancements INTO original changes (preserve v1_content, v2_content)
        enhanced_changes = enhanced_data.get('changes', [])
        original_changes = comparison_results.get('changes', [])

        # Match and merge by section_number and section_title
        for orig_change in original_changes:
            section_num = orig_change.get('section_number')
            section_title = orig_change.get('section_title')

            # Find matching enhanced change
            for enh_change in enhanced_changes:
                if (enh_change.get('section_number') == section_num and
                    enh_change.get('section_title') == section_title):
                    # Merge enhanced fields into original change
                    orig_change['impact'] = enh_change.get('impact', orig_change.get('impact'))
                    orig_change['reasoning'] = enh_change.get('reasoning', '')
                    orig_change['risk_factors'] = enh_change.get('risk_factors', [])
                    orig_change['business_impact'] = enh_change.get('business_impact', '')
                    orig_change['recommendation'] = enh_change.get('recommendation', '')
                    orig_change['cumulative_notes'] = enh_change.get('cumulative_notes', '')
                    break

        comparison_results['executive_summary_enhanced'] = enhanced_data.get('executive_summary', {})
        comparison_results['claude_enhanced'] = True

        # Update impact breakdown from enhanced results
        impact_counts = {}
        for change in comparison_results['changes']:
            impact = change.get('impact', 'ADMINISTRATIVE')
            impact_counts[impact] = impact_counts.get(impact, 0) + 1

        comparison_results['impact_breakdown'] = impact_counts

        logger.info("Claude skill enhancement completed successfully")
        return comparison_results

    except json.JSONDecodeError as e:
        logger.error(f"Failed to parse Claude response as JSON: {e}")
        logger.error(f"Response text: {response_text[:500]}...")
        comparison_results['claude_enhanced'] = False
        comparison_results['enhancement_error'] = f"JSON parse error: {str(e)}"
        return comparison_results

    except Exception as e:
        logger.error(f"Claude skill enhancement failed: {e}", exc_info=True)
        comparison_results['claude_enhanced'] = False
        comparison_results['enhancement_error'] = str(e)
        return comparison_results


# ============================================================================
# ERROR HANDLERS
# ============================================================================

@app.errorhandler(HTTPException)
def handle_http_exception(e):
    """Handle HTTP exceptions"""
    logger.error(f"HTTP Error: {e.code} - {e.description}")
    return jsonify({
        'error': e.name,
        'message': e.description,
        'status_code': e.code
    }), e.code


@app.errorhandler(Exception)
def handle_exception(e):
    """Handle general exceptions"""
    logger.error(f"Unhandled exception: {str(e)}", exc_info=True)
    return jsonify({
        'error': 'Internal Server Error',
        'message': str(e),
        'status_code': 500
    }), 500


# ============================================================================
# HEALTH CHECK ENDPOINT
# ============================================================================

@app.route('/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    status = {
        'status': 'healthy',
        'service': 'CIP API',
        'timestamp': datetime.now().isoformat(),
        'orchestrator': orchestrator is not None,
        'api_key_configured': bool(ANTHROPIC_API_KEY),
        'database': {
            'contracts': CONTRACTS_DB.exists(),
            'reports': REPORTS_DB.exists()
        }
    }

    logger.debug("Health check requested")
    return jsonify(status), 200


# ============================================================================
# INTAKE WIZARD ENDPOINTS (Phase 2)
# ============================================================================

@app.route('/api/upload-enhanced', methods=['POST'])
def upload_enhanced():
    """
    Simplified upload for intake wizard

    Form Data:
        - file: Contract file (PDF, DOCX)

    Returns:
        JSON with contract_id, filename, file_path
    """
    logger.info("Intake wizard upload request received")

    # Check if file is present
    if 'file' not in request.files:
        logger.warning("No file in request")
        return jsonify({'error': 'No file provided'}), 400

    file = request.files['file']

    if file.filename == '':
        logger.warning("Empty filename")
        return jsonify({'error': 'No file selected'}), 400

    if not allowed_file(file.filename):
        logger.warning(f"Invalid file type: {file.filename}")
        return jsonify({
            'error': 'Invalid file type',
            'supported_formats': SUPPORTED_FORMATS
        }), 400

    try:
        # Save file
        file_path = save_uploaded_file(file)
        logger.info(f"File saved: {file_path}")

        # Create contract record with status = 'intake'
        conn = get_db_connection()
        cursor = conn.cursor()

        cursor.execute("""
            INSERT INTO contracts (
                filename, filepath, upload_date, status
            ) VALUES (?, ?, ?, ?)
        """, (
            file.filename,
            str(file_path),
            datetime.now().isoformat(),
            'intake'
        ))

        contract_id = cursor.lastrowid
        conn.commit()
        conn.close()

        logger.info(f"Contract created for intake wizard: ID={contract_id}")

        return jsonify({
            'contract_id': contract_id,
            'filename': file.filename,
            'file_path': str(file_path)
        }), 200

    except Exception as e:
        logger.error(f"Upload failed: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/parse-metadata', methods=['POST'])
def parse_metadata():
    """
    Extract contract metadata using Claude AI

    JSON Body:
        - contract_id: Contract ID

    Returns:
        JSON with extracted metadata:
        {
            'metadata': {
                'title': str,
                'counterparty': str,
                'contract_type': str,
                'effective_date': str (YYYY-MM-DD),
                'expiration_date': str (YYYY-MM-DD),
                'contract_value': float
            }
        }
    """
    logger.info("Parse metadata request received")

    # Validate Claude API availability
    claude_error = validate_claude_available()
    if claude_error:
        return jsonify(claude_error[0]), claude_error[1]

    # Validate orchestrator availability
    orch_error = validate_orchestrator_available()
    if orch_error:
        return jsonify(orch_error[0]), orch_error[1]

    data = request.get_json()

    if not data:
        return jsonify({'error': 'No data provided'}), 400

    contract_id = data.get('contract_id')

    if not contract_id:
        return jsonify({'error': 'contract_id required'}), 400

    try:
        # Get contract file path from database
        conn, db_error = safe_db_connection()
        if db_error:
            return jsonify(db_error[0]), db_error[1]

        cursor = conn.cursor()

        cursor.execute("SELECT filename, filepath FROM contracts WHERE id = ?", (contract_id,))
        row = cursor.fetchone()
        conn.close()

        if not row:
            logger.error(f"Contract {contract_id} not found")
            return jsonify({'error': 'Contract not found'}), 404

        filename = row[0]
        file_path = row[1]

        logger.info(f"Extracting metadata for contract {contract_id}: {filename}")

        doc_text = orchestrator.extractor.extract_text(file_path)

        if not doc_text or len(doc_text) < 100:
            logger.warning("Insufficient document content for metadata extraction")
            return jsonify({
                'metadata': {
                    'title': filename,
                    'counterparty': '',
                    'contract_type': 'Other',
                    'effective_date': None,
                    'expiration_date': None,
                    'contract_value': None
                }
            }), 200

        # Prepare Claude prompt for metadata extraction
        extraction_prompt = f"""You are a contract analysis expert. Extract key metadata from this contract document.

DOCUMENT FILENAME: {filename}

DOCUMENT TEXT (first 4000 chars):
{doc_text[:4000]}

Extract and return ONLY a valid JSON object with these exact fields:
{{
    "title": "Contract title or name (e.g., 'Master Service Agreement', 'Software License Agreement')",
    "counterparty": "The name of the other party in the contract (company or person)",
    "contract_type": "One of: NDA, MSA, SOW, License, Employment, Amendment, Other",
    "effective_date": "YYYY-MM-DD or null if not found",
    "expiration_date": "YYYY-MM-DD or null if not found",
    "contract_value": numeric dollar amount or null if not found (no currency symbols)
}}

IMPORTANT:
- Return ONLY the JSON object, no other text
- For contract_type, choose the BEST match from the list
- For contract_value, extract only the number (e.g., 50000 not "$50,000")
- Use null (not string "null") for missing values"""

        # Call Claude API
        if not claude_client:
            logger.error("Claude client not initialized")
            return jsonify({'error': 'AI service not available'}), 503

        response = claude_client.messages.create(
            model=DEFAULT_MODEL,
            max_tokens=2000,
            temperature=0.3,
            messages=[{
                "role": "user",
                "content": extraction_prompt
            }]
        )

        # Parse response
        response_text = response.content[0].text.strip()
        logger.debug(f"Claude response: {response_text}")

        # Extract JSON from response
        import re
        json_match = re.search(r'\{.*\}', response_text, re.DOTALL)

        if json_match:
            metadata = json.loads(json_match.group())
            logger.info(f"Metadata extracted successfully for contract {contract_id}")

            # Store parsed_metadata in database
            conn = get_db_connection()
            cursor = conn.cursor()
            cursor.execute("""
                UPDATE contracts
                SET parsed_metadata = ?
                WHERE id = ?
            """, (json.dumps(metadata), contract_id))
            conn.commit()
            conn.close()

            return jsonify({'metadata': metadata}), 200
        else:
            logger.error("Could not parse JSON from Claude response")
            return jsonify({'error': 'Failed to parse metadata from AI response'}), 500

    except json.JSONDecodeError as e:
        logger.error(f"JSON decode error: {e}")
        return jsonify({'error': f'Invalid JSON response from AI: {str(e)}'}), 500
    except Exception as e:
        logger.error(f"Metadata extraction failed: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/verify-metadata', methods=['POST'])
def verify_metadata():
    """
    Save user-verified metadata to database

    JSON Body:
        - contract_id: Contract ID
        - metadata: {
            'title': str,
            'counterparty': str,
            'contract_type': str,
            'contract_stage': str (MNDA|NDA|COMMERCIAL|EXECUTED),
            'position': str (customer|vendor|partner|landlord|tenant),
            'leverage': str (strong|balanced|weak),
            'effective_date': str,
            'expiration_date': str,
            'contract_value': float
        }

    Returns:
        Success status
    """
    logger.info("Verify metadata request received")

    data = request.get_json()

    if not data:
        return jsonify({'error': 'No data provided'}), 400

    contract_id = data.get('contract_id')
    metadata = data.get('metadata')

    if not contract_id or not metadata:
        return jsonify({'error': 'contract_id and metadata required'}), 400

    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        # Update contract with verified metadata (v1.4: includes all date fields + category)
        cursor.execute("""
            UPDATE contracts
            SET title = ?,
                counterparty = ?,
                contract_type = ?,
                contract_stage = ?,
                contract_category = ?,
                position = ?,
                leverage = ?,
                date_in = ?,
                date_executed = ?,
                effective_date = ?,
                expiration_date = ?,
                contract_value = ?,
                version_number = ?,
                parent_id = ?,
                narrative = ?,
                metadata_verified = 1,
                status = 'active',
                updated_at = CURRENT_TIMESTAMP
            WHERE id = ?
        """, (
            metadata.get('title'),
            metadata.get('counterparty'),
            metadata.get('contract_type'),
            metadata.get('contract_stage', 'COMMERCIAL'),
            metadata.get('contract_category', 'A'),  # v1.2: A/B/C/D
            metadata.get('position'),
            metadata.get('leverage'),
            metadata.get('date_in'),  # v1.4: Date received
            metadata.get('date_executed'),  # v1.4: Signature date
            metadata.get('effective_date'),
            metadata.get('expiration_date'),
            metadata.get('contract_value'),
            metadata.get('version_number', 1),
            metadata.get('parent_id'),
            metadata.get('business_owner'),  # Stored in narrative field
            contract_id
        ))

        conn.commit()
        conn.close()

        logger.info(f"Metadata verified and saved for contract {contract_id} (stage: {metadata.get('contract_stage', 'COMMERCIAL')})")

        return jsonify({
            'success': True,
            'contract_id': contract_id,
            'message': 'Metadata verified and saved'
        }), 200

    except Exception as e:
        logger.error(f"Metadata verification failed: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/find-similar', methods=['POST'])
def find_similar():
    """
    Find similar contracts based on metadata

    JSON Body:
        - contract_id: Current contract ID (to exclude from results)
        - counterparty: Counterparty name for matching
        - contract_type: Contract type for matching
        - title: Contract title for keyword matching

    Returns:
        JSON with list of similar contracts:
        {
            'similar': [
                {
                    'id': int,
                    'title': str,
                    'counterparty': str,
                    'contract_type': str,
                    'uploaded': str,
                    'status': str,
                    'match_reason': str  # 'counterparty', 'type', or 'title'
                },
                ...
            ]
        }
    """
    logger.info("Find similar contracts request received")

    data = request.get_json()

    if not data:
        return jsonify({'error': 'No data provided'}), 400

    contract_id = data.get('contract_id')
    counterparty = data.get('counterparty', '')
    contract_type = data.get('contract_type', '')
    title = data.get('title', '')

    if not contract_id:
        return jsonify({'error': 'contract_id required'}), 400

    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        similar_contracts = []
        seen_ids = set()

        # Strategy 1: Exact counterparty match (highest priority)
        if counterparty:
            cursor.execute("""
                SELECT id, title, counterparty, contract_type, upload_date, status
                FROM contracts
                WHERE id != ? AND counterparty = ?
                ORDER BY upload_date DESC
                LIMIT 10
            """, (contract_id, counterparty))

            for row in cursor.fetchall():
                if row[0] not in seen_ids:
                    similar_contracts.append({
                        'id': row[0],
                        'title': row[1] or 'Untitled',
                        'counterparty': row[2] or '',
                        'contract_type': row[3] or 'Unknown',
                        'uploaded': row[4] or '',
                        'status': row[5] or 'active',
                        'match_reason': 'counterparty'
                    })
                    seen_ids.add(row[0])

        # Strategy 2: Same contract type (if we need more results)
        if len(similar_contracts) < 10 and contract_type:
            cursor.execute("""
                SELECT id, title, counterparty, contract_type, upload_date, status
                FROM contracts
                WHERE id != ? AND contract_type = ?
                ORDER BY upload_date DESC
                LIMIT ?
            """, (contract_id, contract_type, 10 - len(similar_contracts)))

            for row in cursor.fetchall():
                if row[0] not in seen_ids:
                    similar_contracts.append({
                        'id': row[0],
                        'title': row[1] or 'Untitled',
                        'counterparty': row[2] or '',
                        'contract_type': row[3] or 'Unknown',
                        'uploaded': row[4] or '',
                        'status': row[5] or 'active',
                        'match_reason': 'type'
                    })
                    seen_ids.add(row[0])

        # Strategy 3: Title keyword match (if we still need more)
        if len(similar_contracts) < 10 and title:
            # Extract first significant word from title (ignore common words)
            keywords = [w for w in title.split() if len(w) > 3 and w.lower() not in ['the', 'and', 'for', 'with']]
            if keywords:
                keyword = keywords[0]
                cursor.execute("""
                    SELECT id, title, counterparty, contract_type, upload_date, status
                    FROM contracts
                    WHERE id != ? AND title LIKE ?
                    ORDER BY upload_date DESC
                    LIMIT ?
                """, (contract_id, f'%{keyword}%', 10 - len(similar_contracts)))

                for row in cursor.fetchall():
                    if row[0] not in seen_ids:
                        similar_contracts.append({
                            'id': row[0],
                            'title': row[1] or 'Untitled',
                            'counterparty': row[2] or '',
                            'contract_type': row[3] or 'Unknown',
                            'uploaded': row[4] or '',
                            'status': row[5] or 'active',
                            'match_reason': 'title'
                        })
                        seen_ids.add(row[0])

        conn.close()

        logger.info(f"Found {len(similar_contracts)} similar contracts for contract {contract_id}")

        return jsonify({'similar': similar_contracts}), 200

    except Exception as e:
        logger.error(f"Find similar failed: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/link-contracts', methods=['POST'])
def link_contracts():
    """
    Create relationships between contracts

    JSON Body:
        - contract_id: Current contract ID
        - relationships: [
            {
                'related_id': int,
                'type': 'version' | 'amendment' | 'related' | 'child'
            },
            ...
          ]

    Returns:
        Success status
    """
    logger.info("Link contracts request received")

    data = request.get_json()

    if not data:
        return jsonify({'error': 'No data provided'}), 400

    contract_id = data.get('contract_id')
    relationships = data.get('relationships', [])

    if not contract_id:
        return jsonify({'error': 'contract_id required'}), 400

    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        for rel in relationships:
            related_id = rel.get('related_id')
            rel_type = rel.get('type')

            if not related_id or not rel_type:
                logger.warning(f"Skipping invalid relationship: {rel}")
                continue

            # Handle different relationship types
            if rel_type == 'version':
                # This contract is a new version of related_id
                # Set parent_id and increment version_number

                # Get parent's version number
                cursor.execute("SELECT version_number FROM contracts WHERE id = ?", (related_id,))
                row = cursor.fetchone()
                parent_version = row[0] if row and row[0] else 1

                # Update current contract
                cursor.execute("""
                    UPDATE contracts
                    SET parent_id = ?,
                        relationship_type = 'version',
                        version_number = ?
                    WHERE id = ?
                """, (related_id, parent_version + 1, contract_id))

                # Mark parent as not latest version
                cursor.execute("""
                    UPDATE contracts
                    SET is_latest_version = 0
                    WHERE id = ?
                """, (related_id,))

                logger.info(f"Linked contract {contract_id} as version {parent_version + 1} of {related_id}")

            elif rel_type == 'amendment':
                # This contract amends related_id
                cursor.execute("""
                    UPDATE contracts
                    SET parent_id = ?,
                        relationship_type = 'amendment'
                    WHERE id = ?
                """, (related_id, contract_id))

                logger.info(f"Linked contract {contract_id} as amendment to {related_id}")

            elif rel_type == 'child':
                # This contract is a child of related_id (e.g., SOW under MSA)
                cursor.execute("""
                    UPDATE contracts
                    SET parent_id = ?,
                        relationship_type = 'child'
                    WHERE id = ?
                """, (related_id, contract_id))

                logger.info(f"Linked contract {contract_id} as child of {related_id}")

            elif rel_type == 'related':
                # General relationship, no parent_id
                cursor.execute("""
                    UPDATE contracts
                    SET relationship_type = 'related'
                    WHERE id = ?
                """, (contract_id,))

                logger.info(f"Marked contract {contract_id} as related to {related_id}")

        # Update status to active
        cursor.execute("""
            UPDATE contracts
            SET status = 'active'
            WHERE id = ?
        """, (contract_id,))

        conn.commit()
        conn.close()

        logger.info(f"Successfully linked {len(relationships)} relationships for contract {contract_id}")

        return jsonify({
            'success': True,
            'contract_id': contract_id,
            'status': 'active',
            'relationships_created': len(relationships)
        }), 200

    except Exception as e:
        logger.error(f"Link contracts failed: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


# ============================================================================
# PATTERN SELECTION ENDPOINT (Gate 4 - Pattern Hand)
# ============================================================================

@app.route('/api/patterns/select', methods=['POST'])
def select_patterns():
    """
    Select applicable patterns for a contract (Gate 4 - Pattern Hand)
    v1.1: Added stage filtering and escalation tracking

    JSON Body:
        - contract_type: Type of contract (MSA, SOW, NDA, etc.)
        - position: Our position (customer, vendor, integrator, channel)
        - leverage: Leverage level (strong, balanced, weak)
        - contract_stage: Contract stage (MNDA, NDA, COMMERCIAL, EXECUTED) - v1.3
        - include_research_needed: Whether to include RESEARCH_NEEDED patterns - v1.3

    Returns:
        JSON with selected pattern cards, filter metadata, and escalations
    """
    logger.info("Pattern selection request received (v1.1)")

    if not prompt_composer:
        return jsonify({
            'error': 'Prompt Composer not available',
            'details': 'Pattern deck failed to load'
        }), 503

    data = request.get_json()

    if not data:
        return jsonify({'error': 'No data provided'}), 400

    contract_type = data.get('contract_type', 'MSA')
    position = data.get('position', 'customer')
    leverage = data.get('leverage', 'balanced')
    contract_stage = data.get('contract_stage', 'COMMERCIAL')
    include_research_needed = data.get('include_research_needed', False)

    try:
        # Get pattern hand for display with v1.3 filtering
        result = prompt_composer.get_pattern_hand_display(
            contract_type=contract_type,
            position=position,
            leverage=leverage,
            contract_stage=contract_stage,
            include_research_needed=include_research_needed
        )

        # Get category mapping info
        category = prompt_composer._map_category(contract_type)

        logger.info(f"Selected {len(result['cards'])} patterns for {contract_type}/{position}/{leverage} (stage={contract_stage})")

        return jsonify({
            'status': 'success',
            'contract_type': contract_type,
            'category': category,
            'position': position,
            'leverage': leverage,
            'contract_stage': contract_stage,
            'pattern_count': len(result['cards']),
            'patterns': result['cards'],
            'filter_meta': result['filter_meta'],
            'deck_version': result['deck_version'],
            'pattern_library_version': result['pattern_library_version']
        }), 200

    except Exception as e:
        logger.error(f"Pattern selection failed: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


# ============================================================================
# CONTRACT ANALYSIS ENDPOINT
# ============================================================================

@app.route('/api/analyze', methods=['POST'])
def analyze_contract():
    """
    Trigger contract analysis

    JSON Body:
        - contract_id: ID of uploaded contract
        OR
        - file_path: Direct path to contract file
        - position: Our position (vendor, customer, etc.)
        - leverage: Leverage level (strong, moderate, weak)
        - narrative: Specific concerns (optional)

    Returns:
        JSON with risk assessment
    """
    logger.info("Analysis request received")

    # Validate orchestrator availability
    orch_error = validate_orchestrator_available()
    if orch_error:
        return jsonify(orch_error[0]), orch_error[1]

    data = request.get_json()

    if not data:
        return jsonify({'error': 'No data provided'}), 400

    try:
        # Get contract from database or use provided path
        if 'contract_id' in data:
            contract_id = data['contract_id']

            conn, db_error = safe_db_connection()
            if db_error:
                return jsonify(db_error[0]), db_error[1]

            cursor = conn.cursor()
            cursor.execute("SELECT * FROM contracts WHERE id = ?", (contract_id,))
            row = cursor.fetchone()
            conn.close()

            if not row:
                return jsonify({'error': 'Contract not found'}), 404

            contract = dict(row)
            # Handle both filepath and file_path column names
            file_path = contract.get('filepath') or contract.get('file_path')

            # Create context from database
            context = ContractContext(
                position=contract.get('position', 'unknown'),
                leverage=contract.get('leverage', 'unknown'),
                narrative=contract.get('narrative', ''),
                contract_type=contract.get('contract_type'),
                parties=contract.get('parties')
            )

        else:
            # Use provided data
            file_path = data.get('file_path')
            if not file_path:
                return jsonify({'error': 'contract_id or file_path required'}), 400

            contract_id = None
            context = ContractContext(
                position=data.get('position', 'unknown'),
                leverage=data.get('leverage', 'unknown'),
                narrative=data.get('narrative', ''),
                contract_type=data.get('contract_type'),
                parties=data.get('parties')
            )

        logger.info(f"Starting analysis for: {file_path}")

        # Compose prompt with pattern cards if composer is available
        pattern_metadata = None
        composed_prompt = None

        # Get contract_stage from contract record or request data
        contract_stage = contract.get('contract_stage', 'COMMERCIAL') if 'contract_id' in data else data.get('contract_stage', 'COMMERCIAL')

        if prompt_composer and context.contract_type:
            try:
                # Compose full prompt with pattern cards (v1.1 with stage filtering)
                composed_prompt, compose_meta = prompt_composer.compose_prompt(
                    contract_id=contract_id or 0,
                    contract_type=context.contract_type,
                    position=context.position,
                    leverage=context.leverage,
                    contract_text="",  # Contract text extracted by orchestrator
                    contract_stage=contract_stage,
                    include_research_needed=data.get('include_research_needed', False),
                    base_system_prompt=""  # Let composer build full prompt
                )

                # Get display-friendly pattern list for response
                pattern_result = prompt_composer.get_pattern_hand_display(
                    contract_type=context.contract_type,
                    position=context.position,
                    leverage=context.leverage,
                    contract_stage=contract_stage,
                    include_research_needed=data.get('include_research_needed', False)
                )

                pattern_metadata = {
                    'patterns': pattern_result['cards'],
                    'pattern_count': compose_meta.get('pattern_count', len(pattern_result['cards'])),
                    'deck_version': pattern_result['deck_version'],
                    'pattern_library_version': pattern_result['pattern_library_version'],
                    'category': compose_meta.get('category'),
                    'contract_stage': contract_stage,
                    'prompt_id': compose_meta.get('prompt_id'),
                    'token_count': compose_meta.get('token_count'),
                    # v1.3 additions
                    'filter_meta': pattern_result['filter_meta'],
                    'escalations_triggered': compose_meta.get('escalations_triggered', []),
                    'dealbreakers_flagged': compose_meta.get('dealbreakers_flagged', [])
                }

                logger.info(f"Composed prompt with {compose_meta.get('pattern_count')} patterns, {compose_meta.get('token_count')} tokens")
            except Exception as pe:
                logger.warning(f"Pattern composition failed (falling back to knowledge base): {pe}")
                composed_prompt = None

        # Perform analysis - pass composed prompt if available
        assessment = orchestrator.analyze_contract_file(
            file_path=file_path,
            context=context,
            save_to_db=bool(contract_id),
            pattern_prompt=composed_prompt,
            contract_id=contract_id  # Pass existing ID to prevent duplicate INSERT
        )

        logger.info(f"Analysis complete - Risk: {assessment.overall_risk}")

        # Format response
        response = format_risk_assessment(assessment)

        # Include pattern metadata if available
        result = {
            'status': 'completed',
            'analysis': response
        }
        if pattern_metadata:
            result['patterns'] = pattern_metadata

        return jsonify(result), 200

    except FileNotFoundError as e:
        logger.error(f"File not found: {e}")
        return jsonify({'error': 'Contract file not found'}), 404

    except AnalyzeError as e:
        # RA-ERR-10054: Clean error response without raw exception details
        logger.error(
            f"Analysis error: contract_id={data.get('contract_id')}, "
            f"error_category={e.error_category}, "
            f"exception_type={type(e).__name__}"
        )
        return jsonify({
            'success': False,
            'error_category': e.error_category,
            'error_message_key': e.error_message_key
        }), 500

    except (ConnectionResetError, ConnectionError, OSError) as e:
        # RA-ERR-10054: Network errors - classify and return clean response
        error_result = classify_analyze_error(e)
        logger.error(
            f"Analysis network error: contract_id={data.get('contract_id')}, "
            f"error_category={error_result.error_category}, "
            f"exception_type={type(e).__name__}"
        )
        return jsonify({
            'success': False,
            'error_category': error_result.error_category,
            'error_message_key': error_result.error_message_key
        }), 500

    except Exception as e:
        # RA-ERR-10054: Catch-all - classify and return clean response
        error_result = classify_analyze_error(e)
        logger.error(
            f"Analysis failed: contract_id={data.get('contract_id')}, "
            f"error_category={error_result.error_category}, "
            f"exception_type={type(e).__name__}",
            exc_info=True
        )
        return jsonify({
            'success': False,
            'error_category': error_result.error_category,
            'error_message_key': error_result.error_message_key
        }), 500


# ============================================================================
# COMPARE CONTRACTS ENDPOINT
# ============================================================================

@app.route('/api/compare', methods=['POST'])
def compare_contracts():
    """
    Compare two contract versions

    JSON Body:
        - v1_contract_id: ID of first contract version
        - v2_contract_id: ID of second contract version
        - include_recommendations: Include recommendations in report (optional, default: True)

    Returns:
        JSON with comparison results and report path
    """
    logger.info("Contract comparison request received")

    data = request.get_json()

    if not data:
        return jsonify({'error': 'No data provided'}), 400

    v1_id = data.get('v1_contract_id')
    v2_id = data.get('v2_contract_id')
    include_recommendations = data.get('include_recommendations', True)
    force_refresh = data.get('force_refresh', False)

    if not v1_id or not v2_id:
        return jsonify({'error': 'Both v1_contract_id and v2_contract_id required'}), 400

    if v1_id == v2_id:
        return jsonify({'error': 'Contract IDs must be different'}), 400

    try:
        # Import comparison tools
        import sys
        from pathlib import Path as PathLib
        comparison_scripts = PathLib(__file__).parent.parent / "tools" / "comparison" / "scripts"
        if str(comparison_scripts) not in sys.path:
            sys.path.insert(0, str(comparison_scripts))

        # Import Claude-native comparison (replaces Python section detection)
        import sys
        comparison_dir = PathLib(__file__).parent.parent / "tools" / "comparison"
        if str(comparison_dir) not in sys.path:
            sys.path.insert(0, str(comparison_dir))

        from compare_with_claude import compare_contracts_with_claude
        from scripts.generate_report import ReportGenerator

        # Get both contracts from database
        conn = get_db_connection()
        cursor = conn.cursor()

        cursor.execute("SELECT * FROM contracts WHERE id = ?", (v1_id,))
        v1_row = cursor.fetchone()

        if not v1_row:
            conn.close()
            return jsonify({'error': f'Contract V1 (ID: {v1_id}) not found'}), 404

        cursor.execute("SELECT * FROM contracts WHERE id = ?", (v2_id,))
        v2_row = cursor.fetchone()

        if not v2_row:
            conn.close()
            return jsonify({'error': f'Contract V2 (ID: {v2_id}) not found'}), 404

        v1_contract = dict(v1_row)
        v2_contract = dict(v2_row)

        # Handle both filepath and file_path column names
        v1_path = v1_contract.get('filepath') or v1_contract.get('file_path')
        v2_path = v2_contract.get('filepath') or v2_contract.get('file_path')

        # Validate files exist and are .docx
        v1_file = PathLib(v1_path)
        v2_file = PathLib(v2_path)

        if not v1_file.exists():
            conn.close()
            return jsonify({'error': f'Contract V1 file not found: {v1_path}'}), 404

        if not v2_file.exists():
            conn.close()
            return jsonify({'error': f'Contract V2 file not found: {v2_path}'}), 404

        if v1_file.suffix.lower() != '.docx':
            conn.close()
            return jsonify({'error': 'V1 contract must be .docx format'}), 400

        if v2_file.suffix.lower() != '.docx':
            conn.close()
            return jsonify({'error': 'V2 contract must be .docx format'}), 400

        logger.info(f"Comparing: {v1_file.name} → {v2_file.name}")

        # === CACHE LOOKUP ===
        # Compute comparison_hash from content hashes
        import hashlib
        v1_hash = v1_contract.get('content_hash') or ''
        v2_hash = v2_contract.get('content_hash') or ''

        if v1_hash and v2_hash:
            comparison_hash = hashlib.sha256(f"{v1_hash}:{v2_hash}".encode()).hexdigest()

            # Check cache (skip if force_refresh)
            if not force_refresh:
                cursor.execute("""
                    SELECT comparison_id, result_json, created_at
                    FROM comparison_snapshots
                    WHERE comparison_hash = ?
                """, (comparison_hash,))
                cached = cursor.fetchone()

                if cached and cached['result_json']:
                    logger.info(f"Cache HIT for comparison_hash {comparison_hash[:16]}...")
                    cached_results = json.loads(cached['result_json'])
                    conn.close()
                    return jsonify({
                        'status': 'completed',
                        'cached': True,
                        'comparison_id': cached['comparison_id'],
                        'cache_created_at': cached['created_at'],
                        'comparison_hash': comparison_hash,
                        'total_changes': cached_results.get('total_changes', 0),
                        'impact_breakdown': cached_results.get('impact_breakdown', {}),
                        'executive_summary': cached_results.get('executive_summary', ''),
                        'v1_contract': {'id': v1_id, 'filename': v1_contract['filename']},
                        'v2_contract': {'id': v2_id, 'filename': v2_contract['filename']},
                        'v1_risk': v1_contract.get('risk_level') or 'Unknown',
                        'v2_risk': v2_contract.get('risk_level') or 'Unknown'
                    }), 200
                else:
                    logger.info(f"Cache MISS for comparison_hash {comparison_hash[:16]}...")
            else:
                logger.info(f"Cache BYPASS (force_refresh) for comparison_hash {comparison_hash[:16]}...")
        else:
            comparison_hash = None
            logger.warning("Cannot cache: content_hash missing for one or both contracts")

        # Create output paths
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        reports_dir = PathLib(__file__).parent.parent / "data" / "reports"
        reports_dir.mkdir(exist_ok=True)

        comparison_json_path = reports_dir / f"comparison_{v1_id}_{v2_id}_{timestamp}.json"
        report_docx_path = reports_dir / f"comparison_report_{v1_id}_{v2_id}_{timestamp}.docx"

        # Run Claude-native comparison (finds ALL changes with full analysis)
        logger.info("Running Claude-native document comparison...")

        # Build context for Claude
        v1_context = {
            'contract_id': v1_contract['id'],
            'filename': v1_contract['filename'],
            'position': v1_contract.get('position', 'Customer'),
            'leverage': v1_contract.get('leverage', 'Moderate'),
            'narrative': v1_contract.get('narrative', '')
        }

        v2_context = {
            'contract_id': v2_contract['id'],
            'filename': v2_contract['filename'],
            'position': v2_contract.get('position', 'Customer'),
            'leverage': v2_contract.get('leverage', 'Moderate'),
            'narrative': v2_contract.get('narrative', '')
        }

        try:
            comparison_results = compare_contracts_with_claude(
                v1_path=str(v1_path),
                v2_path=str(v2_path),
                output_path=str(comparison_json_path),
                api_key=ANTHROPIC_API_KEY,
                v1_context=v1_context,
                v2_context=v2_context
            )

            logger.info(f"Claude comparison complete: {comparison_results.get('total_changes', 0)} changes detected")

            # Add file paths for report generation
            comparison_results['v1_path'] = str(v1_path)
            comparison_results['v2_path'] = str(v2_path)
            comparison_results['v1_contract'] = v1_contract
            comparison_results['v2_contract'] = v2_contract

        except Exception as e:
            # Phase 4B: Classify error and return AIResult-compatible response
            error_str = str(e).lower()
            if "timeout" in error_str or "connection" in error_str or "network" in error_str:
                error_category = "network_error"
                error_key = "compare.network_failure"
            elif "401" in error_str or "403" in error_str or "auth" in error_str:
                error_category = "auth_error"
                error_key = "compare.auth_failure"
            elif "400" in error_str or "413" in error_str or "payload" in error_str:
                error_category = "payload_error"
                error_key = "compare.payload_failure"
            else:
                error_category = "internal_error"
                error_key = "compare.internal_failure"

            logger.error(
                f"Compare error: v1_id={v1_id}, v2_id={v2_id}, "
                f"error_category={error_category}, exception_type={type(e).__name__}"
            )
            conn.close()
            return jsonify({
                'success': False,
                'error_category': error_category,
                'error_message_key': error_key
            }), 500

        # Generate report (optional - skip if it fails due to missing fields)
        try:
            logger.info("Generating comparison report...")
            generator = ReportGenerator(
                comparison_data=comparison_results,
                user_context={
                    'v1_filename': v1_contract['filename'],
                    'v2_filename': v2_contract['filename'],
                    'role': 'Contract Review Team'
                }
            )

            generator.generate_report(
                output_path=str(report_docx_path),
                include_recommendations=include_recommendations
            )

            logger.info(f"Report generated: {report_docx_path.name}")
        except Exception as report_error:
            logger.warning(f"Report generation failed (non-critical): {report_error}")
            # Continue without DOCX report - JSON is the primary output
            report_docx_path = None

        # Create executive summary
        impact_counts = {}
        for change in comparison_results['changes']:
            impact = change['impact']
            impact_counts[impact] = impact_counts.get(impact, 0) + 1

        executive_summary = f"Comparison identified {comparison_results['total_changes']} substantive changes"
        if impact_counts.get('CRITICAL', 0) > 0:
            executive_summary += f" including {impact_counts['CRITICAL']} CRITICAL items requiring immediate attention"

        # Store in reports database
        reports_conn = get_db_connection(db_path=str(REPORTS_DB))
        reports_cursor = reports_conn.cursor()

        reports_cursor.execute("""
            INSERT INTO comparisons (
                v1_contract_id, v2_contract_id, substantive_changes,
                administrative_changes, executive_summary, report_path, metadata_json
            ) VALUES (?, ?, ?, ?, ?, ?, ?)
        """, (
            v1_id,
            v2_id,
            comparison_results['total_changes'],
            0,  # Administrative changes not tracked separately
            executive_summary,
            str(report_docx_path),
            json.dumps(impact_counts)
        ))

        comparison_id = reports_cursor.lastrowid
        reports_conn.commit()
        reports_conn.close()

        # === CACHE WRITE ===
        # Store in comparison_snapshots for future cache hits
        if comparison_hash:
            try:
                cache_data = {
                    'total_changes': comparison_results['total_changes'],
                    'impact_breakdown': impact_counts,
                    'executive_summary': executive_summary,
                    'changes': comparison_results.get('changes', [])
                }
                cursor.execute("""
                    INSERT INTO comparison_snapshots (
                        v1_contract_id, v2_contract_id, comparison_hash,
                        similarity_score, changed_clauses, risk_delta,
                        result_json, created_at
                    ) VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                """, (
                    v1_id,
                    v2_id,
                    comparison_hash,
                    100.0 - (comparison_results['total_changes'] * 2),  # Rough similarity
                    json.dumps(comparison_results.get('changes', [])),
                    json.dumps(impact_counts),
                    json.dumps(cache_data),
                    datetime.now().isoformat()
                ))
                conn.commit()
                logger.info(f"Cache WRITE for comparison_hash {comparison_hash[:16]}...")
            except Exception as cache_err:
                logger.warning(f"Cache write failed (non-critical): {cache_err}")

        conn.close()

        logger.info(f"Comparison stored with ID: {comparison_id}")

        # Return response
        return jsonify({
            'status': 'completed',
            'cached': False,
            'comparison_id': comparison_id,
            'comparison_hash': comparison_hash,
            'report_path': str(report_docx_path) if report_docx_path else None,
            'json_path': str(comparison_json_path),
            'claude_enhanced': comparison_results.get('claude_enhanced', False),
            'executive_summary': executive_summary,
            'total_changes': comparison_results['total_changes'],
            'impact_breakdown': impact_counts,
            'v1_contract': {
                'id': v1_id,
                'filename': v1_contract['filename']
            },
            'v2_contract': {
                'id': v2_id,
                'filename': v2_contract['filename']
            },
            'v1_risk': v1_contract.get('risk_level') or 'Unknown',
            'v2_risk': v2_contract.get('risk_level') or 'Unknown'
        }), 200

    except ImportError as e:
        logger.error(f"Failed to import comparison tools: {e}")
        return jsonify({
            'success': False,
            'error_category': 'internal_error',
            'error_message_key': 'compare.internal_failure'
        }), 500

    except FileNotFoundError as e:
        logger.error(f"File not found during comparison: {e}")
        return jsonify({
            'success': False,
            'error_category': 'internal_error',
            'error_message_key': 'compare.internal_failure'
        }), 404

    except Exception as e:
        logger.error(f"Comparison failed: {e}", exc_info=True)
        return jsonify({
            'success': False,
            'error_category': 'internal_error',
            'error_message_key': 'compare.internal_failure'
        }), 500


# ============================================================================
# GET RISK ASSESSMENT ENDPOINT
# ============================================================================

@app.route('/api/assessment/<int:contract_id>', methods=['GET'])
def get_risk_assessment(contract_id: int):
    """
    Get risk assessment for a contract

    URL Parameters:
        - contract_id: Contract ID

    Returns:
        JSON with risk assessment details
    """
    logger.info(f"Get assessment request for contract {contract_id}")

    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        # Get latest assessment
        cursor.execute("""
            SELECT * FROM risk_assessments
            WHERE contract_id = ?
            ORDER BY assessment_date DESC
            LIMIT 1
        """, (contract_id,))

        row = cursor.fetchone()

        if not row:
            conn.close()
            return jsonify({'error': 'No assessment found for this contract'}), 404

        assessment = dict(row)

        # Parse JSON fields
        if assessment.get('critical_items'):
            assessment['critical_items'] = json.loads(assessment['critical_items'])
        if assessment.get('dealbreakers'):
            assessment['dealbreakers'] = json.loads(assessment['dealbreakers'])
        if assessment.get('analysis_json'):
            assessment['analysis_json'] = json.loads(assessment['analysis_json'])

        # Get contract info
        cursor.execute("SELECT * FROM contracts WHERE id = ?", (contract_id,))
        contract_row = cursor.fetchone()

        conn.close()

        response = {
            'assessment': assessment,
            'contract': dict(contract_row) if contract_row else None
        }

        logger.info(f"Assessment retrieved for contract {contract_id}")
        return jsonify(response), 200

    except Exception as e:
        logger.error(f"Failed to retrieve assessment: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


# ============================================================================
# LIST CONTRACTS ENDPOINT
# ============================================================================

@app.route('/api/contracts', methods=['GET'])
def list_contracts():
    """
    List all contracts with optional filtering

    Query Parameters:
        - status: Filter by status
        - contract_type: Filter by type
        - limit: Max results (default: 100)
        - offset: Pagination offset (default: 0)

    Returns:
        JSON array of contracts
    """
    logger.info("List contracts request")

    try:
        # Get query parameters
        status = request.args.get('status')
        contract_type = request.args.get('contract_type')
        include_archived = request.args.get('include_archived', 'false').lower() == 'true'
        limit = int(request.args.get('limit', 100))
        offset = int(request.args.get('offset', 0))

        conn = get_db_connection()
        cursor = conn.cursor()

        # Build query - exclude archived by default
        query = "SELECT * FROM contracts WHERE 1=1"
        params = []

        if not include_archived:
            query += " AND (archived = 0 OR archived IS NULL)"

        if status:
            query += " AND status = ?"
            params.append(status)

        if contract_type:
            query += " AND contract_type = ?"
            params.append(contract_type)

        query += " ORDER BY upload_date DESC LIMIT ? OFFSET ?"
        params.extend([limit, offset])

        cursor.execute(query, params)
        rows = cursor.fetchall()

        # Get total count
        count_query = "SELECT COUNT(*) as total FROM contracts WHERE 1=1"
        count_params = []

        if not include_archived:
            count_query += " AND (archived = 0 OR archived IS NULL)"

        if status:
            count_query += " AND status = ?"
            count_params.append(status)

        if contract_type:
            count_query += " AND contract_type = ?"
            count_params.append(contract_type)

        cursor.execute(count_query, count_params)
        total = cursor.fetchone()['total']

        conn.close()

        contracts = [dict(row) for row in rows]

        logger.info(f"Retrieved {len(contracts)} contracts (total: {total})")

        return jsonify({
            'contracts': contracts,
            'total': total,
            'limit': limit,
            'offset': offset
        }), 200

    except Exception as e:
        logger.error(f"Failed to list contracts: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


# ============================================================================
# GET CONTRACT DETAILS ENDPOINT
# ============================================================================

@app.route('/api/contract/<int:contract_id>', methods=['GET'])
def get_contract_details(contract_id: int):
    """
    Get detailed contract information

    URL Parameters:
        - contract_id: Contract ID

    Returns:
        JSON with contract details, clauses, and assessments
    """
    logger.info(f"Get contract details for ID {contract_id}")

    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        # Get contract
        cursor.execute("SELECT * FROM contracts WHERE id = ?", (contract_id,))
        contract_row = cursor.fetchone()

        if not contract_row:
            conn.close()
            return jsonify({'error': 'Contract not found'}), 404

        contract = dict(contract_row)

        # Get clauses
        cursor.execute("""
            SELECT * FROM clauses
            WHERE contract_id = ?
            ORDER BY section_number
        """, (contract_id,))

        clauses = [dict(row) for row in cursor.fetchall()]

        # Get assessments
        cursor.execute("""
            SELECT * FROM risk_assessments
            WHERE contract_id = ?
            ORDER BY assessment_date DESC
        """, (contract_id,))

        assessments = [dict(row) for row in cursor.fetchall()]

        # Negotiations table removed in schema v2.0
        negotiations = []

        conn.close()

        response = {
            'contract': contract,
            'clauses': clauses,
            'assessments': assessments,
            'negotiations': negotiations,
            'statistics': {
                'total_clauses': len(clauses),
                'total_assessments': len(assessments),
                'total_negotiations': len(negotiations),
                'critical_clauses': len([c for c in clauses if c.get('risk_level') == 'CRITICAL']),
                'high_risk_clauses': len([c for c in clauses if c.get('risk_level') in ['CRITICAL', 'HIGH']])
            }
        }

        logger.info(f"Contract details retrieved for ID {contract_id}")
        return jsonify(response), 200

    except Exception as e:
        logger.error(f"Failed to get contract details: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


# ============================================================================
# ARCHIVE/RESTORE ENDPOINTS
# ============================================================================

@app.route('/api/contracts/<int:contract_id>/archive', methods=['POST'])
def archive_contract(contract_id):
    """
    Archive a contract (soft delete)

    Args:
        contract_id: ID of contract to archive

    Returns:
        JSON with success status
    """
    logger.info(f"Archive request for contract {contract_id}")

    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        # Check contract exists
        cursor.execute("SELECT id, filename, archived FROM contracts WHERE id = ?", (contract_id,))
        contract = cursor.fetchone()

        if not contract:
            conn.close()
            return jsonify({'error': 'Contract not found'}), 404

        if contract['archived'] == 1:
            conn.close()
            return jsonify({'error': 'Contract is already archived'}), 400

        # Archive the contract
        cursor.execute(
            "UPDATE contracts SET archived = 1, updated_at = CURRENT_TIMESTAMP WHERE id = ?",
            (contract_id,)
        )
        conn.commit()
        conn.close()

        logger.info(f"Contract {contract_id} archived successfully")
        return jsonify({
            'status': 'success',
            'message': f"Contract '{contract['filename']}' has been archived",
            'contract_id': contract_id
        }), 200

    except Exception as e:
        logger.error(f"Failed to archive contract: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/contracts/<int:contract_id>/restore', methods=['POST'])
def restore_contract(contract_id):
    """
    Restore an archived contract

    Args:
        contract_id: ID of contract to restore

    Returns:
        JSON with success status
    """
    logger.info(f"Restore request for contract {contract_id}")

    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        # Check contract exists
        cursor.execute("SELECT id, filename, archived FROM contracts WHERE id = ?", (contract_id,))
        contract = cursor.fetchone()

        if not contract:
            conn.close()
            return jsonify({'error': 'Contract not found'}), 404

        if contract['archived'] != 1:
            conn.close()
            return jsonify({'error': 'Contract is not archived'}), 400

        # Restore the contract
        cursor.execute(
            "UPDATE contracts SET archived = 0, updated_at = CURRENT_TIMESTAMP WHERE id = ?",
            (contract_id,)
        )
        conn.commit()
        conn.close()

        logger.info(f"Contract {contract_id} restored successfully")
        return jsonify({
            'status': 'success',
            'message': f"Contract '{contract['filename']}' has been restored",
            'contract_id': contract_id
        }), 200

    except Exception as e:
        logger.error(f"Failed to restore contract: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/contracts/<int:contract_id>/delete', methods=['DELETE', 'POST'])
def delete_contract(contract_id):
    """
    Permanently delete a contract (hard delete)

    Args:
        contract_id: ID of contract to delete

    Returns:
        JSON with success status
    """
    logger.info(f"Delete request for contract {contract_id}")

    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        # Check contract exists
        cursor.execute("SELECT id, filename FROM contracts WHERE id = ?", (contract_id,))
        contract = cursor.fetchone()

        if not contract:
            conn.close()
            return jsonify({'error': 'Contract not found'}), 404

        filename = contract['filename']

        # Delete related records first (foreign key constraints)
        cursor.execute("DELETE FROM clauses WHERE contract_id = ?", (contract_id,))
        cursor.execute("DELETE FROM risk_assessments WHERE contract_id = ?", (contract_id,))

        # Delete the contract
        cursor.execute("DELETE FROM contracts WHERE id = ?", (contract_id,))
        conn.commit()
        conn.close()

        logger.info(f"Contract {contract_id} deleted permanently")
        return jsonify({
            'status': 'success',
            'message': f"Contract '{filename}' has been permanently deleted",
            'contract_id': contract_id
        }), 200

    except Exception as e:
        logger.error(f"Failed to delete contract: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/contracts/<int:contract_id>/stage', methods=['POST'])
def update_contract_stage(contract_id):
    """Update workflow stage for contract."""
    data = request.get_json()
    stage = data.get('workflow_stage')

    if stage is None:
        return jsonify({'error': 'workflow_stage required'}), 400

    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("UPDATE contracts SET workflow_stage = ? WHERE id = ?", (stage, contract_id))

        # Auto-update status based on stage
        if stage == 3:
            cursor.execute("UPDATE contracts SET status = 'negotiation' WHERE id = ? AND status IN ('intake', 'review')", (contract_id,))

        conn.commit()
        conn.close()
        logger.info(f"[WORKFLOW] Updated contract {contract_id} to stage {stage}")
        return jsonify({'status': 'success', 'workflow_stage': stage}), 200
    except Exception as e:
        logger.error(f"[WORKFLOW] Stage update failed: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/contracts/delete-batch', methods=['POST'])
def delete_contracts_batch():
    """
    Permanently delete multiple contracts (hard delete)

    Args:
        JSON body with 'contract_ids': list of IDs to delete

    Returns:
        JSON with success status and count
    """
    try:
        data = request.json or {}
        contract_ids = data.get('contract_ids', [])

        if not contract_ids:
            return jsonify({'error': 'No contract IDs provided'}), 400

        logger.info(f"Batch delete request for contracts: {contract_ids}")

        conn = get_db_connection()
        cursor = conn.cursor()

        deleted_count = 0
        errors = []

        for contract_id in contract_ids:
            try:
                # Check contract exists
                cursor.execute("SELECT id, filename FROM contracts WHERE id = ?", (contract_id,))
                contract = cursor.fetchone()

                if not contract:
                    errors.append(f"Contract {contract_id} not found")
                    continue

                # Delete related records first
                cursor.execute("DELETE FROM clauses WHERE contract_id = ?", (contract_id,))
                cursor.execute("DELETE FROM risk_assessments WHERE contract_id = ?", (contract_id,))

                # Delete the contract
                cursor.execute("DELETE FROM contracts WHERE id = ?", (contract_id,))
                deleted_count += 1

            except Exception as e:
                errors.append(f"Failed to delete contract {contract_id}: {str(e)}")

        conn.commit()
        conn.close()

        logger.info(f"Batch delete completed: {deleted_count} contracts deleted")
        return jsonify({
            'status': 'success',
            'deleted_count': deleted_count,
            'errors': errors if errors else None
        }), 200

    except Exception as e:
        logger.error(f"Failed to batch delete contracts: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


# ============================================================================
# STATISTICS ENDPOINT
# ============================================================================

@app.route('/api/statistics', methods=['GET'])
def get_statistics():
    """
    Get platform statistics

    Returns:
        JSON with overall statistics
    """
    logger.info("Statistics request")

    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        # Total contracts (exclude archived)
        cursor.execute("SELECT COUNT(*) as total FROM contracts WHERE (archived = 0 OR archived IS NULL)")
        total_contracts = cursor.fetchone()['total']

        # Active contracts (exclude archived)
        cursor.execute("SELECT COUNT(*) as total FROM contracts WHERE status = 'active' AND (archived = 0 OR archived IS NULL)")
        active_contracts = cursor.fetchone()['total']

        # Total assessments (for non-archived contracts)
        cursor.execute("""
            SELECT COUNT(*) as total FROM risk_assessments ra
            JOIN contracts c ON ra.contract_id = c.id
            WHERE (c.archived = 0 OR c.archived IS NULL)
        """)
        total_assessments = cursor.fetchone()['total']

        # High risk contracts (for non-archived contracts)
        cursor.execute("""
            SELECT COUNT(*) as total FROM risk_assessments ra
            JOIN contracts c ON ra.contract_id = c.id
            WHERE ra.overall_risk = 'HIGH' AND (c.archived = 0 OR c.archived IS NULL)
        """)
        high_risk = cursor.fetchone()['total']

        conn.close()

        # Reports statistics
        reports_conn = sqlite3.connect(str(REPORTS_DB))
        reports_conn.row_factory = sqlite3.Row
        reports_cursor = reports_conn.cursor()

        reports_cursor.execute("SELECT COUNT(*) as total FROM comparisons")
        total_comparisons = reports_cursor.fetchone()['total']

        reports_cursor.execute("SELECT COUNT(*) as total FROM risk_reports")
        total_reports = reports_cursor.fetchone()['total']

        reports_conn.close()

        stats = {
            'contracts': {
                'total': total_contracts,
                'active': active_contracts
            },
            'assessments': {
                'total': total_assessments,
                'high_risk': high_risk
            },
            'reports': {
                'comparisons': total_comparisons,
                'risk_reports': total_reports
            },
            'timestamp': datetime.now().isoformat()
        }

        logger.info("Statistics retrieved successfully")
        return jsonify(stats), 200

    except Exception as e:
        logger.error(f"Failed to get statistics: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/dashboard/stats', methods=['GET'])
def get_dashboard_stats():
    """
    Get comprehensive dashboard statistics for portfolio overview.

    Query Parameters (all optional):
        - type: Filter by contract_type (e.g., MSA, NDA)
        - role: Filter by contract_role (customer_contract, vendor_contract)
        - status: Filter by status (active, expired, etc.)
        - risk: Filter by risk_level (Critical, High, Moderate, Low, Administrative)

    Returns:
        JSON with portfolio metrics, distributions, trends, and contracts list
    """
    logger.info("Dashboard stats request")

    # Get filter parameters
    filter_type = request.args.get('type')
    filter_role = request.args.get('role')
    filter_status = request.args.get('status')
    filter_risk = request.args.get('risk')

    logger.info(f"Filters: type={filter_type}, role={filter_role}, status={filter_status}, risk={filter_risk}")

    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        # Build base WHERE clause for filters (always exclude archived)
        where_clauses = ["(archived = 0 OR archived IS NULL)"]
        params = []

        if filter_type:
            where_clauses.append("contract_type = ?")
            params.append(filter_type)

        if filter_role:
            where_clauses.append("contract_role = ?")
            params.append(filter_role)

        if filter_status:
            where_clauses.append("status = ?")
            params.append(filter_status)

        if filter_risk:
            where_clauses.append("risk_level = ?")
            params.append(filter_risk)

        base_where = " AND ".join(where_clauses)

        # =====================================================================
        # KEY METRICS (filtered)
        # =====================================================================

        # Total Portfolio Value (filtered)
        cursor.execute(f"SELECT COALESCE(SUM(contract_value), 0) as total FROM contracts WHERE {base_where}", params)
        total_portfolio_value = cursor.fetchone()['total'] or 0

        # Active Contracts (filtered + status=active)
        active_where = where_clauses.copy()
        active_params = params.copy()
        if not filter_status:
            active_where.append("status = 'active'")
        cursor.execute(f"""
            SELECT COUNT(*) as total FROM contracts
            WHERE {" AND ".join(active_where)}
        """, active_params)
        active_contracts = cursor.fetchone()['total']

        # Total Contracts (filtered)
        cursor.execute(f"SELECT COUNT(*) as total FROM contracts WHERE {base_where}", params)
        total_contracts = cursor.fetchone()['total']

        # Expiring in 90 Days (filtered)
        expiring_where = where_clauses.copy()
        expiring_params = params.copy()
        expiring_where.append("expiration_date IS NOT NULL")
        expiring_where.append("expiration_date <= date('now', '+90 days')")
        expiring_where.append("expiration_date >= date('now')")
        cursor.execute(f"""
            SELECT COUNT(*) as total FROM contracts
            WHERE {" AND ".join(expiring_where)}
        """, expiring_params)
        expiring_90_days = cursor.fetchone()['total']

        # High-value contracts expiring (filtered)
        hv_where = expiring_where.copy()
        hv_params = expiring_params.copy()
        hv_where.append("contract_value > 100000")
        cursor.execute(f"""
            SELECT COUNT(*) as total FROM contracts
            WHERE {" AND ".join(hv_where)}
        """, hv_params)
        high_value_expiring = cursor.fetchone()['total']

        # High-Risk Items (filtered by risk if specified, otherwise Critical+High)
        if filter_risk:
            # If filtering by specific risk, count those
            cursor.execute(f"""
                SELECT COUNT(*) as total FROM contracts
                WHERE {base_where} AND risk_level IS NOT NULL
            """, params)
        else:
            # Default: count Critical + High
            risk_where = where_clauses.copy()
            risk_params = params.copy()
            risk_where.append("risk_level IN ('Critical', 'High', 'CRITICAL', 'HIGH')")
            cursor.execute(f"""
                SELECT COUNT(*) as total FROM contracts
                WHERE {" AND ".join(risk_where)}
            """, risk_params)
        high_risk_items = cursor.fetchone()['total']

        # =====================================================================
        # STATUS DISTRIBUTION (filtered)
        # =====================================================================
        cursor.execute(f"""
            SELECT
                COALESCE(status, 'Unknown') as status,
                COUNT(*) as count,
                COALESCE(SUM(contract_value), 0) as value
            FROM contracts
            WHERE {base_where}
            GROUP BY status
            ORDER BY count DESC
        """, params)
        status_rows = cursor.fetchall()
        status_distribution = [
            {'status': row['status'], 'count': row['count'], 'value': row['value']}
            for row in status_rows
        ]

        # =====================================================================
        # RISK DISTRIBUTION (filtered - from contracts.risk_level)
        # =====================================================================
        cursor.execute(f"""
            SELECT
                COALESCE(risk_level, 'Unknown') as risk,
                COUNT(*) as count
            FROM contracts
            WHERE {base_where} AND risk_level IS NOT NULL
            GROUP BY risk_level
            ORDER BY
                CASE risk_level
                    WHEN 'Low' THEN 1 WHEN 'LOW' THEN 1
                    WHEN 'Moderate' THEN 2 WHEN 'Medium' THEN 2 WHEN 'MEDIUM' THEN 2
                    WHEN 'High' THEN 3 WHEN 'HIGH' THEN 3
                    WHEN 'Critical' THEN 4 WHEN 'CRITICAL' THEN 4
                    WHEN 'Administrative' THEN 5
                    ELSE 6
                END
        """, params)
        risk_rows = cursor.fetchall()
        risk_distribution = [
            {'risk': row['risk'], 'count': row['count']}
            for row in risk_rows
        ]

        # =====================================================================
        # CONTRACT TYPE DISTRIBUTION (filtered)
        # =====================================================================
        cursor.execute(f"""
            SELECT
                COALESCE(contract_type, 'Unknown') as type,
                COUNT(*) as count,
                COALESCE(AVG(contract_value), 0) as avg_value
            FROM contracts
            WHERE {base_where}
            GROUP BY contract_type
            ORDER BY count DESC
            LIMIT 10
        """, params)
        type_rows = cursor.fetchall()
        type_distribution = [
            {'type': row['type'], 'count': row['count'], 'avg_value': row['avg_value']}
            for row in type_rows
        ]

        # =====================================================================
        # TOP COUNTERPARTIES (filtered)
        # =====================================================================
        counterparty_where = where_clauses.copy()
        counterparty_params = params.copy()
        counterparty_where.append("counterparty IS NOT NULL")
        counterparty_where.append("counterparty != ''")
        cursor.execute(f"""
            SELECT
                COALESCE(counterparty, 'Unknown') as name,
                COUNT(*) as contracts,
                COALESCE(SUM(contract_value), 0) as value
            FROM contracts
            WHERE {" AND ".join(counterparty_where)}
            GROUP BY counterparty
            ORDER BY value DESC
            LIMIT 10
        """, counterparty_params)
        counterparty_rows = cursor.fetchall()
        top_counterparties = [
            {'name': row['name'], 'contracts': row['contracts'], 'value': row['value']}
            for row in counterparty_rows
        ]

        # =====================================================================
        # EXPIRATION CALENDAR (Next 90 days, filtered)
        # =====================================================================
        exp_cal_where = where_clauses.copy()
        exp_cal_params = params.copy()
        exp_cal_where.append("expiration_date IS NOT NULL")
        exp_cal_where.append("expiration_date >= date('now')")
        exp_cal_where.append("expiration_date <= date('now', '+90 days')")
        cursor.execute(f"""
            SELECT
                expiration_date as date,
                COUNT(*) as count
            FROM contracts
            WHERE {" AND ".join(exp_cal_where)}
            GROUP BY expiration_date
            ORDER BY expiration_date
        """, exp_cal_params)
        expiration_rows = cursor.fetchall()
        expiration_calendar = [
            {'date': row['date'], 'count': row['count']}
            for row in expiration_rows
        ]

        # =====================================================================
        # VALUE TREND (by upload month, filtered)
        # =====================================================================
        trend_where = where_clauses.copy()
        trend_params = params.copy()
        trend_where.append("upload_date IS NOT NULL")
        cursor.execute(f"""
            SELECT
                strftime('%Y-%m', upload_date) as month,
                COUNT(*) as contracts,
                COALESCE(SUM(contract_value), 0) as value
            FROM contracts
            WHERE {" AND ".join(trend_where)}
            GROUP BY strftime('%Y-%m', upload_date)
            ORDER BY month DESC
            LIMIT 12
        """, trend_params)
        trend_rows = cursor.fetchall()
        value_trend = [
            {'month': row['month'], 'contracts': row['contracts'], 'value': row['value']}
            for row in trend_rows
        ][::-1]  # Reverse to chronological order

        # =====================================================================
        # RECENT ACTIVITY (last 10 contracts, filtered)
        # =====================================================================
        cursor.execute(f"""
            SELECT
                id,
                filename,
                contract_type,
                contract_value,
                upload_date,
                status
            FROM contracts
            WHERE {base_where}
            ORDER BY upload_date DESC
            LIMIT 10
        """, params)
        recent_rows = cursor.fetchall()
        recent_activity = [
            {
                'id': row['id'],
                'filename': row['filename'],
                'contract_type': row['contract_type'],
                'value': row['contract_value'],
                'date': row['upload_date'],
                'status': row['status']
            }
            for row in recent_rows
        ]

        # =====================================================================
        # CONTRACTS LIST (paginated, filtered) - for table display
        # =====================================================================
        page = int(request.args.get('page', 0))
        limit = int(request.args.get('limit', 25))
        offset = page * limit

        cursor.execute(f"""
            SELECT
                id, title, counterparty, contract_type, contract_role,
                status, contract_value, risk_level, expiration_date
            FROM contracts
            WHERE {base_where}
            ORDER BY id DESC
            LIMIT ? OFFSET ?
        """, params + [limit, offset])
        list_rows = cursor.fetchall()
        contracts_list = [
            {
                'id': row['id'],
                'title': row['title'] or 'Untitled',
                'counterparty': row['counterparty'] or '-',
                'contract_type': row['contract_type'] or '-',
                'contract_role': row['contract_role'] or '-',
                'status': row['status'] or '-',
                'contract_value': row['contract_value'],
                'risk_level': row['risk_level'] or '-',
                'expiration_date': row['expiration_date'] or '-'
            }
            for row in list_rows
        ]

        conn.close()

        # Build response
        dashboard_stats = {
            'key_metrics': {
                'total_portfolio_value': total_portfolio_value,
                'active_contracts': active_contracts,
                'total_contracts': total_contracts,
                'expiring_90_days': expiring_90_days,
                'high_value_expiring': high_value_expiring,
                'high_risk_items': high_risk_items
            },
            'status_distribution': status_distribution,
            'risk_distribution': risk_distribution,
            'type_distribution': type_distribution,
            'top_counterparties': top_counterparties,
            'expiration_calendar': expiration_calendar,
            'value_trend': value_trend,
            'recent_activity': recent_activity,
            'contracts_list': contracts_list,
            'total_filtered': total_contracts,
            'timestamp': datetime.now().isoformat()
        }

        logger.info("Dashboard stats retrieved successfully")
        return jsonify(dashboard_stats), 200

    except Exception as e:
        logger.error(f"Failed to get dashboard stats: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


# ============================================================================
# PORTFOLIO FILTER ENDPOINTS (Phase 4A.2)
# ============================================================================

@app.route('/api/portfolio/filters', methods=['GET'])
def get_filter_options():
    """Return available filter values from database"""
    try:
        conn = get_db_connection()
        return jsonify({
            'types': query_distinct(conn, 'contracts', 'contract_type'),
            'statuses': query_distinct(conn, 'contracts', 'status'),
            'risk_levels': ['Critical', 'High', 'Moderate', 'Low', 'Administrative'],
            'counterparties': query_distinct(conn, 'contracts', 'counterparty')
        }), 200
    except Exception as e:
        logger.error(f"Failed to get filter options: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/portfolio/contracts', methods=['POST'])
def get_filtered_contracts():
    """Return contracts matching filter criteria"""
    try:
        filters = request.json or {}
        include_archived = filters.get('include_archived', False)

        # Start with archived filter (exclude by default) and exclude incomplete intake
        if include_archived:
            query = "SELECT * FROM contracts WHERE status != 'intake'"
        else:
            query = "SELECT * FROM contracts WHERE (archived = 0 OR archived IS NULL) AND status != 'intake'"
        params = []

        if filters.get('type'):
            query += " AND contract_type = ?"
            params.append(filters['type'])
        if filters.get('status'):
            query += " AND status = ?"
            params.append(filters['status'])
        if filters.get('risk'):
            query += " AND risk_level = ?"
            params.append(filters['risk'])
        if filters.get('risk_high'):
            query += " AND risk_level IN ('Critical', 'High')"
        if filters.get('counterparty'):
            query += " AND counterparty LIKE ?"
            params.append(f"%{filters['counterparty']}%")
        if filters.get('expiring_days'):
            query += " AND expiration_date IS NOT NULL AND expiration_date BETWEEN date('now') AND date('now', '+' || ? || ' days')"
            params.append(filters['expiring_days'])

        query += " ORDER BY created_at DESC"

        conn = get_db_connection()
        results = conn.execute(query, params).fetchall()
        return jsonify([dict(row) for row in results]), 200
    except Exception as e:
        logger.error(f"Failed to get filtered contracts: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/portfolio/kpis', methods=['POST'])
def get_portfolio_kpis():
    """Return KPI values based on current filters (always excludes archived)"""
    try:
        filters = request.json or {}
        conn = get_db_connection()

        # Build base WHERE clause (always exclude archived and incomplete intake)
        where = "(archived = 0 OR archived IS NULL) AND status != 'intake'"
        params = []
        if filters.get('type'):
            where += " AND contract_type = ?"
            params.append(filters['type'])
        if filters.get('status'):
            where += " AND status = ?"
            params.append(filters['status'])
        if filters.get('risk'):
            where += " AND risk_level = ?"
            params.append(filters['risk'])

        total = conn.execute(f"SELECT COALESCE(SUM(contract_value),0) FROM contracts WHERE {where}", params).fetchone()[0]
        active = conn.execute(f"SELECT COUNT(*) FROM contracts WHERE {where} AND status='active'", params).fetchone()[0]
        high_risk = conn.execute(f"SELECT COUNT(*) FROM contracts WHERE {where} AND risk_level IN ('Critical','High')", params).fetchone()[0]

        # Expiring in 90 days (handle null expiration_date gracefully)
        expiring = conn.execute(f"""
            SELECT COUNT(*) FROM contracts
            WHERE {where} AND expiration_date IS NOT NULL
            AND expiration_date BETWEEN date('now') AND date('now', '+90 days')
        """, params).fetchone()[0]

        return jsonify({
            'total_value': total,
            'active_count': active,
            'expiring_90d': expiring,
            'high_risk': high_risk
        }), 200
    except Exception as e:
        logger.error(f"Failed to get portfolio KPIs: {e}")
        return jsonify({'error': str(e)}), 500


# ============================================================================
# CONTRACT DETAIL ENDPOINTS (Phase 4C.2)
# ============================================================================

@app.route('/api/contract/<int:contract_id>/versions', methods=['GET'])
def get_contract_versions(contract_id):
    """Get version history for a contract"""
    try:
        conn = get_db_connection()
        versions = conn.execute(
            """SELECT id, version_number, title, status, last_amended_date, created_at
               FROM contracts
               WHERE id = ? OR parent_id = ?
               ORDER BY version_number DESC""",
            (contract_id, contract_id)
        ).fetchall()
        return jsonify([dict(v) for v in versions]), 200
    except Exception as e:
        logger.error(f"Failed to get contract versions: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/contract/<int:contract_id>/relationships', methods=['GET'])
def get_contract_relationships(contract_id):
    """Get related contracts (parent, children, amendments)"""
    try:
        conn = get_db_connection()
        contract = conn.execute("SELECT * FROM contracts WHERE id = ?", (contract_id,)).fetchone()

        result = {'parent': None, 'children': [], 'amendments': []}

        if contract and contract['parent_id']:
            parent = conn.execute(
                "SELECT id, title, contract_type, status FROM contracts WHERE id = ?",
                (contract['parent_id'],)
            ).fetchone()
            result['parent'] = dict(parent) if parent else None

        children = conn.execute(
            """SELECT id, title, relationship_type, contract_type, status
               FROM contracts
               WHERE parent_id = ?""",
            (contract_id,)
        ).fetchall()
        result['children'] = [dict(c) for c in children]

        amendments = conn.execute(
            """SELECT id, title, status, last_amended_date
               FROM contracts
               WHERE parent_id = ? AND contract_type = 'AMENDMENT'""",
            (contract_id,)
        ).fetchall()
        result['amendments'] = [dict(a) for a in amendments]

        return jsonify(result), 200
    except Exception as e:
        logger.error(f"Failed to get contract relationships: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/contract/<int:contract_id>/history', methods=['GET'])
def get_contract_history(contract_id):
    """Get activity history for a contract"""
    try:
        conn = get_db_connection()
        # Basic history from contract updates
        history = [
            {
                'timestamp': contract['created_at'],
                'action': 'Contract Created',
                'details': f"Uploaded as {contract['filename']}"
            }
            for contract in [conn.execute("SELECT * FROM contracts WHERE id = ?", (contract_id,)).fetchone()]
            if contract
        ]
        return jsonify(history), 200
    except Exception as e:
        logger.error(f"Failed to get contract history: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/redline-review', methods=['POST'])
def redline_review():
    """
    Generate redline suggestions for a single contract

    Request body:
    {
        "contract_id": int,
        "context": {
            "position": "Vendor|Customer",
            "leverage": "Strong|Moderate|Weak",
            "contract_type": "Services Agreement|..."
        }
    }

    Returns:
    {
        "clauses": [
            {
                "section_number": "1.1",
                "section_title": "Definitions",
                "clause_text": "Original text...",
                "risk_level": "HIGH",
                "suggested_revision": "Revised text...",
                "html_redline": "HTML with <span>...",
                "change_metrics": {
                    "change_ratio": 0.15,
                    "word_retention": 0.87,
                    "is_minimal": true
                },
                "pattern_applied": "Pattern name",
                "status": "pending"
            }
        ]
    }
    """
    try:
        data = request.get_json()
        contract_id = data.get('contract_id')
        context = data.get('context', {})

        logger.info(f"[REDLINE] Starting redline review for contract {contract_id}")
        # log_user_action requires logger as first param - skipping for now
        # log_user_action(logger, "redline_review_start", contract_id=contract_id)

        # Get contract from database
        contracts_conn = sqlite3.connect(CONTRACTS_DB)
        contracts_conn.row_factory = sqlite3.Row
        cursor = contracts_conn.cursor()

        cursor.execute("""
            SELECT id, filename, filepath
            FROM contracts
            WHERE id = ?
        """, (contract_id,))

        contract = cursor.fetchone()
        contracts_conn.close()

        if not contract:
            logger.warning(f"[REDLINE] Contract {contract_id} not found")
            return jsonify({'error': 'Contract not found'}), 404

        file_path = contract['filepath']

        # Import and initialize RedlineAnalyzer
        try:
            from redline_analyzer import RedlineAnalyzer

            # Use orchestrator to read contract text from file
            if not orchestrator:
                logger.error("[REDLINE] Orchestrator not initialized")
                return jsonify({'error': 'Document reading service not available'}), 503

            # Read contract using orchestrator's document handler
            from docx import Document as DocxDocument
            import PyPDF2

            contract_text = ""

            try:
                if file_path.lower().endswith('.pdf'):
                    with open(file_path, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        for page in pdf_reader.pages:
                            contract_text += page.extract_text() + "\n"
                elif file_path.lower().endswith('.docx'):
                    doc = DocxDocument(file_path)
                    contract_text = "\n".join([para.text for para in doc.paragraphs])
                else:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        contract_text = f.read()
            except Exception as read_error:
                logger.error(f"[REDLINE] Failed to read contract file: {read_error}")
                return jsonify({'error': f'Failed to read contract file: {str(read_error)}'}), 500

            if not contract_text or len(contract_text.strip()) == 0:
                logger.error(f"[REDLINE] Contract file is empty: {file_path}")
                return jsonify({'error': 'Contract file is empty'}), 500

            analyzer = RedlineAnalyzer()
            logger.info(f"[REDLINE] Analyzer initialized, contract length: {len(contract_text)} chars")

            # Analyze document
            clauses = analyzer.analyze_document(contract_text, context)

            logger.info(f"[REDLINE] Generated {len(clauses)} clause analyses")
            logger.info(f"[REDLINE] {sum(1 for c in clauses if c.get('suggested_revision'))} suggestions created")

            # Log completion (log_user_action requires logger as first param)
            try:
                log_user_action(logger, "redline_review_complete", contract_id=contract_id, details={
                    "clauses_analyzed": len(clauses),
                    "suggestions_generated": sum(1 for c in clauses if c.get('suggested_revision'))
                })
            except Exception as log_err:
                logger.warning(f"[REDLINE] Logging failed: {log_err}")

            # Ensure all clause data is JSON serializable
            try:
                response_data = {
                    'contract_id': contract_id,
                    'filename': contract['filename'],
                    'clauses': clauses,
                    'timestamp': datetime.now().isoformat()
                }

                # Test JSON serialization
                import json
                json_test = json.dumps(response_data)
                logger.info(f"[REDLINE] Response is JSON serializable ({len(json_test)} bytes)")

                return jsonify(response_data), 200
            except Exception as json_err:
                logger.error(f"[REDLINE] JSON serialization failed: {json_err}", exc_info=True)
                return jsonify({'error': f'JSON serialization failed: {str(json_err)}'}), 500

        except Exception as e:
            import traceback
            # Phase 4B: Classify error for AIResult-compatible response
            error_str = str(e).lower()
            if "timeout" in error_str or "connection" in error_str or "network" in error_str:
                error_category = "network_error"
                error_key = "redline.network_failure"
            elif "401" in error_str or "403" in error_str or "auth" in error_str:
                error_category = "auth_error"
                error_key = "redline.auth_failure"
            elif "400" in error_str or "413" in error_str or "payload" in error_str:
                error_category = "payload_error"
                error_key = "redline.payload_failure"
            else:
                error_category = "internal_error"
                error_key = "redline.internal_failure"

            logger.error(
                f"[REDLINE] Analysis failed: contract_id={contract_id}, "
                f"error_category={error_category}, exception_type={type(e).__name__}"
            )
            logger.error(f"[REDLINE] Traceback: {traceback.format_exc()}")
            return jsonify({
                'success': False,
                'error_category': error_category,
                'error_message_key': error_key
            }), 500

    except Exception as e:
        import traceback
        logger.error(f"[REDLINE] Request failed: {e}")
        logger.error(f"[REDLINE] Traceback: {traceback.format_exc()}")
        return jsonify({
            'success': False,
            'error_category': 'internal_error',
            'error_message_key': 'redline.internal_failure'
        }), 500


@app.route('/api/export-redlines', methods=['POST'])
def export_redlines():
    """
    Export approved redlines to Word document

    Request body:
    {
        "contract_id": int,
        "clauses": [...],
        "decisions": {0: "approved", 1: "modified", ...},
        "modifications": {1: "modified text", ...},
        "context": {...}
    }

    Returns:
    File download or error
    """
    try:
        data = request.get_json()
        contract_id = data.get('contract_id')
        clauses = data.get('clauses', [])
        decisions = data.get('decisions', {})
        modifications = data.get('modifications', {})
        context = data.get('context', {})

        logger.info(f"[EXPORT] Starting redline export for contract {contract_id}")

        # Get contract info
        contracts_conn = sqlite3.connect(CONTRACTS_DB)
        contracts_conn.row_factory = sqlite3.Row
        cursor = contracts_conn.cursor()

        cursor.execute("""
            SELECT id, filename
            FROM contracts
            WHERE id = ?
        """, (contract_id,))

        contract = cursor.fetchone()
        contracts_conn.close()

        if not contract:
            logger.warning(f"[EXPORT] Contract {contract_id} not found")
            return jsonify({'error': 'Contract not found'}), 404

        # Import exporter
        try:
            from redline_exporter import RedlineExporter

            exporter = RedlineExporter()

            # Prepare output path
            output_filename = f"redline_{contract_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            output_path = Path(UPLOAD_DIRECTORY) / output_filename

            # Prepare contract info
            contract_info = {
                'contract_id': contract_id,
                'filename': contract['filename'],
                'context': context
            }

            # Convert decision keys from strings to ints (JSON converts dict keys to strings)
            decisions_int = {int(k): v for k, v in decisions.items()}
            modifications_int = {int(k): v for k, v in modifications.items()}

            # Export to .docx
            result = exporter.export_to_docx(
                clauses=clauses,
                decisions=decisions_int,
                modifications=modifications_int,
                contract_info=contract_info,
                output_path=str(output_path)
            )

            logger.info(f"[EXPORT] Generated {output_path}")
            logger.info(f"[EXPORT] Exported {result['clauses_exported']} clauses")

            # Log export action (requires logger as first param)
            try:
                log_user_action(logger, "redline_export", contract_id=contract_id, details={
                    'clauses_exported': result['clauses_exported'],
                    'file': output_filename
                })
            except Exception as log_err:
                logger.warning(f"[EXPORT] Logging failed: {log_err}")

            # Return file
            return send_file(
                str(output_path),
                as_attachment=True,
                download_name=output_filename,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

        except Exception as e:
            logger.error(f"[EXPORT] Export failed: {e}", exc_info=True)
            return jsonify({'error': f'Export failed: {str(e)}'}), 500

    except Exception as e:
        logger.error(f"[EXPORT] Request failed: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/redline/save', methods=['POST'])
def save_redline():
    """
    Save redline review decisions.

    JSON Body:
        - contract_id: int
        - decisions: list of {clause_index, decision, modified_text, note}
        - overall_risk_before: str
        - overall_risk_after: str (optional)
    """
    data = request.get_json()

    if not data or 'contract_id' not in data:
        return jsonify({'error': 'contract_id required'}), 400

    contract_id = data['contract_id']
    decisions = data.get('decisions', [])

    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        # Save to redline_snapshots
        cursor.execute("""
            INSERT INTO redline_snapshots (
                contract_id, base_version_contract_id, source_mode,
                created_at, clauses_json, status,
                overall_risk_before, overall_risk_after, dealbreakers_detected
            ) VALUES (?, ?, 'manual', datetime('now'), ?, 'complete', ?, ?, ?)
        """, (
            contract_id,
            data.get('base_version_contract_id', contract_id),  # Default to same contract
            json.dumps(decisions),
            data.get('overall_risk_before', ''),
            data.get('overall_risk_after', ''),
            data.get('dealbreakers_detected', 0)
        ))

        # Update workflow_stage = 2 (Reviewed)
        cursor.execute("UPDATE contracts SET workflow_stage = 2 WHERE id = ?", (contract_id,))

        conn.commit()
        redline_id = cursor.lastrowid
        conn.close()

        logger.info(f"[REDLINE] Saved redline snapshot {redline_id} for contract {contract_id}")
        return jsonify({'status': 'success', 'redline_id': redline_id}), 201

    except Exception as e:
        logger.error(f"[REDLINE] Save failed: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


# ============================================================================
# CONVERSATIONAL INTAKE ENDPOINTS (v1.4)
# ============================================================================

# CLM Stage definitions
CLM_STAGES = [
    ("NEW CONTRACT", "Received, entering system"),
    ("REVIEWING", "Internal review"),
    ("NEGOTIATING", "Back-and-forth with counterparty"),
    ("APPROVING", "Internal approval process"),
    ("EXECUTING", "Signature process"),
    ("IN_EFFECT", "Active, being performed"),
    ("AMENDING", "Modifications in progress"),
    ("RENEWAL", "Renewal negotiation"),
    ("EXPIRED", "Term ended"),
]

# Taxonomy definitions
PARTY_RELATIONSHIPS = [
    ("CUSTOMER", "They pay us"),
    ("VENDOR", "We pay them"),
    ("PARTNER", "Mutual value exchange"),
    ("RESELLER", "They sell our products"),
    ("CONSULTANT", "They advise us"),
]

CONTRACT_PURPOSES = [
    ("SERVICES", "Professional services delivery"),
    ("PRODUCTS", "Goods purchase/sale"),
    ("LICENSING", "IP/software rights"),
    ("DISTRIBUTION", "Channel/resale arrangements"),
    ("PARTNERSHIP", "Joint ventures, alliances"),
    ("CONFIDENTIALITY", "Information protection only"),
    ("OTHER", "Custom - specify"),
]

CONTRACT_TYPES = [
    ("NDA", "Non-Disclosure Agreement"),
    ("MNDA", "Mutual NDA"),
    ("MSA", "Master Service Agreement"),
    ("SOW", "Statement of Work"),
    ("SLA", "Service Level Agreement"),
    ("MOU", "Memorandum of Understanding"),
    ("LOI", "Letter of Intent"),
    ("AMENDMENT", "Contract Amendment"),
    ("ADDENDUM", "Contract Addendum"),
    ("ORDER", "Purchase/Sales Order"),
    ("LICENSE", "License Agreement"),
    ("LEASE", "Lease Agreement"),
    ("EMPLOYMENT", "Employment Contract"),
    ("OTHER", "Custom - specify"),
]

# Question templates for gap-filling
QUESTION_TEMPLATES = {
    "entity_unknown": {
        "question": "Which party is your company?",
        "field": "our_entity",
        "options_from": "extracted_parties"
    },
    "relationship_unknown": {
        "question": "What is your relationship with the counterparty?",
        "field": "party_relationship",
        "options": [r[0] for r in PARTY_RELATIONSHIPS]
    },
    "stage_unknown": {
        "question": "What stage is this contract in?",
        "field": "clm_stage",
        "options": [s[0] for s in CLM_STAGES]
    },
    "purpose_unknown": {
        "question": "What is the primary purpose of this contract?",
        "field": "contract_purpose",
        "options": [p[0] for p in CONTRACT_PURPOSES]
    },
    "type_uncertain": {
        "question": "What type of agreement is this?",
        "field": "contract_type",
        "options": [t[0] for t in CONTRACT_TYPES]
    }
}


@app.route('/api/intake/pending', methods=['GET'])
def get_pending_intakes():
    """
    Get contracts with status='intake' (incomplete uploads).
    These are contracts that were uploaded but intake was not completed.

    Returns:
        JSON list of pending intake contracts
    """
    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        cursor.execute("""
            SELECT id, filename, filepath, upload_date, created_at
            FROM contracts
            WHERE status = 'intake'
            ORDER BY created_at DESC
        """)

        contracts = cursor.fetchall()
        conn.close()

        return jsonify({
            'pending': [dict(row) for row in contracts],
            'count': len(contracts)
        }), 200

    except Exception as e:
        logger.error(f"Failed to get pending intakes: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/intake/analyze', methods=['POST'])
def intake_analyze():
    """
    Analyze uploaded contract and return findings with confidence scores.
    Simplified conversational intake - returns structured findings and questions.

    JSON Body:
        - contract_id: Contract ID to analyze

    Returns:
        {
            findings: { type, parties, dates, values, ... },
            confidence: { type: 0.9, our_entity: 0.0, ... },
            overall_confidence: 0.72,
            questions: [ { field, question, options }, ... ],
            ready_to_confirm: false
        }
    """
    logger.info("[INTAKE] Analyze request received")

    data = request.get_json()
    if not data or not data.get('contract_id'):
        return jsonify({'error': 'contract_id required'}), 400

    contract_id = data['contract_id']

    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        # Get contract
        cursor.execute("SELECT * FROM contracts WHERE id = ?", (contract_id,))
        row = cursor.fetchone()

        if not row:
            conn.close()
            return jsonify({'error': 'Contract not found'}), 404

        contract = dict(row)
        conn.close()

        # Get file path
        file_path = contract.get('filepath') or contract.get('file_path')
        if not file_path or not Path(file_path).exists():
            return jsonify({'error': 'Contract file not found'}), 404

        # Use orchestrator to extract metadata
        extracted = orchestrator.extract_metadata(file_path, contract.get('filename', ''))

        # Build findings from extraction + existing data
        findings = {
            'contract_type': extracted.get('type') or contract.get('contract_type'),
            'parties': extracted.get('parties', []),
            'our_entity': contract.get('our_entity'),
            'counterparty': contract.get('counterparty') or (extracted.get('parties', [None])[0] if extracted.get('parties') else None),
            'clm_stage': contract.get('clm_stage', 'INTAKE'),
            'party_relationship': contract.get('party_relationship'),
            'contract_purpose': contract.get('contract_purpose'),
            'effective_date': extracted.get('dates', {}).get('effective_date') or contract.get('effective_date'),
            'expiration_date': extracted.get('dates', {}).get('expiration_date') or contract.get('expiration_date'),
            'contract_value': extracted.get('amounts', {}).get('total_value') or contract.get('contract_value'),
        }

        # Calculate confidence per dimension
        confidence = {}

        # Type confidence
        type_val = findings.get('contract_type')
        if type_val and type_val != 'Unknown' and type_val != 'Other':
            confidence['contract_type'] = 0.90
        elif type_val == 'Other':
            confidence['contract_type'] = 0.50
        else:
            confidence['contract_type'] = 0.0

        # Our entity confidence
        if findings.get('our_entity'):
            confidence['our_entity'] = 0.95
        elif len(findings.get('parties', [])) == 2:
            confidence['our_entity'] = 0.0  # Need to ask
        else:
            confidence['our_entity'] = 0.0

        # Relationship confidence
        if findings.get('party_relationship'):
            confidence['party_relationship'] = 0.95
        else:
            confidence['party_relationship'] = 0.0

        # Stage confidence - default INTAKE is reasonable
        if findings.get('clm_stage') and findings['clm_stage'] != 'INTAKE':
            confidence['clm_stage'] = 0.90
        else:
            confidence['clm_stage'] = 0.60  # Default is okay but not confirmed

        # Purpose confidence
        if findings.get('contract_purpose'):
            confidence['contract_purpose'] = 0.90
        else:
            # Infer from type
            type_to_purpose = {
                'NDA': 'CONFIDENTIALITY', 'MNDA': 'CONFIDENTIALITY',
                'MSA': 'SERVICES', 'SOW': 'SERVICES', 'SLA': 'SERVICES',
                'LICENSE': 'LICENSING',
            }
            inferred = type_to_purpose.get(findings.get('contract_type'))
            if inferred:
                findings['contract_purpose'] = inferred
                confidence['contract_purpose'] = 0.70
            else:
                confidence['contract_purpose'] = 0.0

        # Calculate overall confidence (weighted)
        weights = {
            'contract_type': 0.25,
            'our_entity': 0.20,
            'party_relationship': 0.20,
            'clm_stage': 0.20,
            'contract_purpose': 0.15
        }

        overall_confidence = sum(
            confidence.get(field, 0) * weight
            for field, weight in weights.items()
        )

        # Generate questions for gaps (confidence < 0.80)
        questions = []

        if confidence.get('our_entity', 0) < 0.80 and len(findings.get('parties', [])) >= 2:
            questions.append({
                'field': 'our_entity',
                'question': "Which party is your company?",
                'options': findings['parties'] + ['Other']
            })

        if confidence.get('party_relationship', 0) < 0.80:
            questions.append({
                'field': 'party_relationship',
                'question': "What is your relationship with the counterparty?",
                'options': [r[0] for r in PARTY_RELATIONSHIPS]
            })

        if confidence.get('clm_stage', 0) < 0.80:
            questions.append({
                'field': 'clm_stage',
                'question': "What stage is this contract in?",
                'options': [s[0] for s in CLM_STAGES[:6]]  # Most common stages
            })

        if confidence.get('contract_purpose', 0) < 0.80:
            questions.append({
                'field': 'contract_purpose',
                'question': "What is the primary purpose of this contract?",
                'options': [p[0] for p in CONTRACT_PURPOSES]
            })

        if confidence.get('contract_type', 0) < 0.80:
            questions.append({
                'field': 'contract_type',
                'question': "What type of agreement is this?",
                'options': [t[0] for t in CONTRACT_TYPES]
            })

        # Determine threshold based on analysis mode
        threshold = 0.90  # Standard

        ready_to_confirm = overall_confidence >= threshold

        logger.info(f"[INTAKE] Analysis complete: confidence={overall_confidence:.2f}, questions={len(questions)}")

        return jsonify({
            'findings': findings,
            'confidence': confidence,
            'overall_confidence': round(overall_confidence, 2),
            'questions': questions,
            'ready_to_confirm': ready_to_confirm,
            'threshold': threshold
        }), 200

    except Exception as e:
        logger.error(f"[INTAKE] Analysis failed: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/intake/update', methods=['POST'])
def intake_update():
    """
    Update findings based on user response (button click).
    Recalculates confidence and returns updated state.

    JSON Body:
        - contract_id: Contract ID
        - field: Field being updated
        - value: Selected value
        - current_findings: Current findings dict

    Returns:
        Updated findings, confidence, questions, ready_to_confirm
    """
    logger.info("[INTAKE] Update request received")

    data = request.get_json()
    if not data:
        return jsonify({'error': 'No data provided'}), 400

    contract_id = data.get('contract_id')
    field = data.get('field')
    value = data.get('value')
    current_findings = data.get('current_findings', {})

    if not all([contract_id, field, value]):
        return jsonify({'error': 'contract_id, field, and value required'}), 400

    try:
        # Update findings
        current_findings[field] = value

        # Special handling: if our_entity selected, set counterparty
        if field == 'our_entity' and current_findings.get('parties'):
            parties = current_findings['parties']
            if value in parties:
                other_parties = [p for p in parties if p != value]
                if other_parties:
                    current_findings['counterparty'] = other_parties[0]

        # Recalculate confidence
        confidence = {}

        # Type
        type_val = current_findings.get('contract_type')
        confidence['contract_type'] = 0.95 if type_val and type_val != 'Unknown' else 0.0

        # Entity
        confidence['our_entity'] = 0.95 if current_findings.get('our_entity') else 0.0

        # Relationship
        confidence['party_relationship'] = 0.95 if current_findings.get('party_relationship') else 0.0

        # Stage
        stage = current_findings.get('clm_stage')
        confidence['clm_stage'] = 0.95 if stage and stage != 'INTAKE' else 0.60

        # Purpose
        confidence['contract_purpose'] = 0.95 if current_findings.get('contract_purpose') else 0.0

        # Overall
        weights = {
            'contract_type': 0.25,
            'our_entity': 0.20,
            'party_relationship': 0.20,
            'clm_stage': 0.20,
            'contract_purpose': 0.15
        }

        overall_confidence = sum(
            confidence.get(f, 0) * w for f, w in weights.items()
        )

        # Generate remaining questions
        questions = []

        if confidence.get('our_entity', 0) < 0.80 and current_findings.get('parties'):
            questions.append({
                'field': 'our_entity',
                'question': "Which party is your company?",
                'options': current_findings['parties'] + ['Other']
            })

        if confidence.get('party_relationship', 0) < 0.80:
            questions.append({
                'field': 'party_relationship',
                'question': "What is your relationship?",
                'options': [r[0] for r in PARTY_RELATIONSHIPS]
            })

        if confidence.get('clm_stage', 0) < 0.80:
            questions.append({
                'field': 'clm_stage',
                'question': "What stage is this contract in?",
                'options': [s[0] for s in CLM_STAGES[:6]]
            })

        if confidence.get('contract_purpose', 0) < 0.80:
            questions.append({
                'field': 'contract_purpose',
                'question': "What is the purpose?",
                'options': [p[0] for p in CONTRACT_PURPOSES]
            })

        if confidence.get('contract_type', 0) < 0.80:
            questions.append({
                'field': 'contract_type',
                'question': "What type of agreement?",
                'options': [t[0] for t in CONTRACT_TYPES]
            })

        ready_to_confirm = overall_confidence >= 0.90

        logger.info(f"[INTAKE] Update complete: field={field}, confidence={overall_confidence:.2f}")

        return jsonify({
            'findings': current_findings,
            'confidence': confidence,
            'overall_confidence': round(overall_confidence, 2),
            'questions': questions,
            'ready_to_confirm': ready_to_confirm
        }), 200

    except Exception as e:
        logger.error(f"[INTAKE] Update failed: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/intake/confirm', methods=['POST'])
def intake_confirm():
    """
    Confirm and save intake findings to database.

    JSON Body:
        - contract_id: Contract ID
        - confirmed_metadata OR findings: Confirmed data dict

    Returns:
        Success status
    """
    logger.info("[INTAKE] Confirm request received")

    data = request.get_json()
    if not data:
        return jsonify({'error': 'No data provided'}), 400

    contract_id = data.get('contract_id')
    # Accept either 'confirmed_metadata' or 'findings' for flexibility
    findings = data.get('confirmed_metadata') or data.get('findings', {})

    if not contract_id:
        return jsonify({'error': 'contract_id required'}), 400

    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        # Update contract with confirmed findings
        cursor.execute("""
            UPDATE contracts
            SET title = COALESCE(?, title),
                contract_type = ?,
                our_entity = ?,
                counterparty = ?,
                clm_stage = ?,
                party_relationship = ?,
                contract_purpose = ?,
                position = ?,
                leverage = ?,
                analysis_mode = ?,
                effective_date = ?,
                expiration_date = ?,
                contract_value = ?,
                metadata_verified = 1,
                status = 'active',
                updated_at = CURRENT_TIMESTAMP
            WHERE id = ?
        """, (
            findings.get('title'),
            findings.get('contract_type'),
            findings.get('our_entity'),
            findings.get('counterparty'),
            findings.get('clm_stage', 'INTAKE'),
            findings.get('party_relationship'),
            findings.get('contract_purpose'),
            findings.get('position'),
            findings.get('leverage'),
            findings.get('analysis_mode'),
            findings.get('effective_date'),
            findings.get('expiration_date'),
            findings.get('contract_value'),
            contract_id
        ))

        conn.commit()
        conn.close()

        logger.info(f"[INTAKE] Contract {contract_id} confirmed and saved")

        return jsonify({
            'status': 'success',
            'contract_id': contract_id,
            'message': 'Contract intake complete'
        }), 200

    except Exception as e:
        logger.error(f"[INTAKE] Confirm failed: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/intake/taxonomy', methods=['GET'])
def get_taxonomy():
    """
    Return taxonomy definitions for frontend dropdowns.
    Returns just the codes (first element of tuples) for easy dropdown use.
    """
    return jsonify({
        'clm_stages': [s[0] for s in CLM_STAGES],
        'party_relationships': [r[0] for r in PARTY_RELATIONSHIPS],
        'contract_purposes': [p[0] for p in CONTRACT_PURPOSES],
        'contract_types': [t[0] for t in CONTRACT_TYPES],
        'analysis_modes': ['FULL', 'NDA_ONLY', 'COMPARISON']
    }), 200


# ============================================================================
# REPORT BUILDER ENDPOINTS (v1)
# ============================================================================

import re
from docx import Document as DocxDocument
from docx.shared import Inches, Pt

# Token definitions for report generation
REPORT_TOKENS = {
    '{{OUR_PARTY_NAME}}': 'our_entity',
    '{{OUR_PARTY_ROLE}}': 'position',
    '{{OTHER_PARTY_NAME}}': 'counterparty',
    '{{OTHER_PARTY_ROLE}}': 'other_role',  # Derived
    '{{CONTRACT_TYPE}}': 'contract_type',
    '{{EFFECTIVE_DATE}}': 'effective_date',
    '{{EXPIRATION_DATE}}': 'expiration_date',
    '{{CLM_STAGE}}': 'clm_stage',
    '{{RELATIONSHIP_TYPE}}': 'party_relationship',
    '{{PURPOSE}}': 'contract_purpose',
    '{{TOTAL_VALUE}}': 'contract_value',
    '{{SUMMARY}}': 'summary',
    '{{CONTRACT_TITLE}}': 'title',
    '{{FILENAME}}': 'filename',
}

# Ensure reports and templates directories exist
REPORTS_DIR = Path(__file__).parent.parent / 'data' / 'reports'
REPORTS_DIR.mkdir(parents=True, exist_ok=True)
TEMPLATES_DIR = Path(__file__).parent.parent / 'data' / 'templates'
TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)


@app.route('/api/templates/extract', methods=['POST'])
def extract_template():
    """
    Upload a .docx template and extract placeholders.

    Returns:
        {
            filename: str,
            placeholders: [str],
            suggested_mappings: {placeholder: suggested_field}
        }
    """
    logger.info("[TEMPLATES] Extract request received")

    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400

    file = request.files['file']
    if not file.filename.endswith('.docx'):
        return jsonify({'error': 'Only .docx files supported'}), 400

    try:
        # Save uploaded template
        safe_filename = secure_filename(file.filename)
        template_path = TEMPLATES_DIR / safe_filename
        file.save(str(template_path))

        # Parse document
        doc = DocxDocument(str(template_path))

        # Extract all text
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)

        combined_text = '\n'.join(full_text)

        # Find placeholders using regex: {{ANYTHING}}
        placeholder_pattern = r'\{\{[A-Z_]+\}\}'
        found_placeholders = list(set(re.findall(placeholder_pattern, combined_text)))
        found_placeholders.sort()

        # Suggest mappings based on known tokens
        suggested_mappings = {}
        for placeholder in found_placeholders:
            if placeholder in REPORT_TOKENS:
                suggested_mappings[placeholder] = REPORT_TOKENS[placeholder]
            else:
                suggested_mappings[placeholder] = None  # Unknown, needs manual mapping

        logger.info(f"[TEMPLATES] Extracted {len(found_placeholders)} placeholders from {safe_filename}")

        return jsonify({
            'filename': safe_filename,
            'filepath': str(template_path),
            'placeholders': found_placeholders,
            'suggested_mappings': suggested_mappings,
            'known_tokens': list(REPORT_TOKENS.keys())
        }), 200

    except Exception as e:
        logger.error(f"[TEMPLATES] Extract failed: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/templates', methods=['POST'])
def save_template():
    """
    Save a template with field mappings.

    JSON Body:
        - name: Template name
        - description: Optional description
        - source_filename: Original filename
        - schema_json: JSON string of placeholders
        - field_mappings: JSON string of placeholder->field mappings
    """
    logger.info("[TEMPLATES] Save request received")

    data = request.get_json()
    if not data:
        return jsonify({'error': 'No data provided'}), 400

    name = data.get('name')
    if not name:
        return jsonify({'error': 'Template name required'}), 400

    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        cursor.execute("""
            INSERT INTO report_templates (name, description, source_filename, schema_json, field_mappings)
            VALUES (?, ?, ?, ?, ?)
        """, (
            name,
            data.get('description', ''),
            data.get('source_filename', ''),
            json.dumps(data.get('placeholders', [])),
            json.dumps(data.get('field_mappings', {}))
        ))

        template_id = cursor.lastrowid
        conn.commit()
        conn.close()

        logger.info(f"[TEMPLATES] Saved template {template_id}: {name}")

        return jsonify({
            'status': 'success',
            'template_id': template_id,
            'name': name
        }), 200

    except Exception as e:
        logger.error(f"[TEMPLATES] Save failed: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/templates', methods=['GET'])
def list_templates():
    """List all active templates."""
    logger.info("[TEMPLATES] List request received")

    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        cursor.execute("""
            SELECT id, name, description, source_filename, created_at, is_active
            FROM report_templates
            WHERE is_active = 1
            ORDER BY created_at DESC
        """)

        templates = []
        for row in cursor.fetchall():
            templates.append({
                'id': row[0],
                'name': row[1],
                'description': row[2],
                'source_filename': row[3],
                'created_at': row[4],
                'is_active': row[5]
            })

        conn.close()

        return jsonify({'templates': templates}), 200

    except Exception as e:
        logger.error(f"[TEMPLATES] List failed: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/templates/<int:template_id>', methods=['GET'])
def get_template(template_id):
    """Get a specific template with its mappings."""
    logger.info(f"[TEMPLATES] Get template {template_id}")

    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        cursor.execute("""
            SELECT id, name, description, source_filename, schema_json, field_mappings, created_at
            FROM report_templates
            WHERE id = ?
        """, (template_id,))

        row = cursor.fetchone()
        conn.close()

        if not row:
            return jsonify({'error': 'Template not found'}), 404

        return jsonify({
            'id': row[0],
            'name': row[1],
            'description': row[2],
            'source_filename': row[3],
            'placeholders': json.loads(row[4]) if row[4] else [],
            'field_mappings': json.loads(row[5]) if row[5] else {},
            'created_at': row[6]
        }), 200

    except Exception as e:
        logger.error(f"[TEMPLATES] Get failed: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


def generate_skill_report(data):
    """
    Generate a report using document generation skill scripts.

    Args:
        data: Request data with report_type, contract_id, parameters

    Returns:
        JSON response with report_id and filename
    """
    import sys
    sys.path.append(r'C:\Users\jrudy\CIP\.claude\skills\cip-document-generation\scripts')

    try:
        contract_id = data['contract_id']
        report_type = data['report_type']
        parameters = data.get('parameters', {})

        logger.info(f"[REPORTS-SKILL] Generating {report_type} report for contract {contract_id}")

        # Get contract data from database
        conn = get_db_connection()
        cursor = conn.cursor()

        cursor.execute("SELECT * FROM contracts WHERE id = ?", (contract_id,))
        contract_row = cursor.fetchone()

        if not contract_row:
            conn.close()
            return jsonify({'error': 'Contract not found'}), 404

        contract = dict(contract_row)

        # Get contract text (if filepath exists)
        contract_text = ""
        if contract.get('filepath'):
            try:
                with open(contract['filepath'], 'r', encoding='utf-8', errors='ignore') as f:
                    contract_text = f.read()
            except:
                pass

        # Extract clauses
        from extract_clauses import extract_clauses, get_clause_summary

        clauses = extract_clauses(contract_text) if contract_text else []
        clause_summary = get_clause_summary(clauses) if clauses else {}

        # Prepare report data based on type
        report_data = {
            'contract_name': contract.get('title', 'Contract'),
            'our_entity': parameters.get('our_entity', contract.get('our_entity', 'Our Company')),
            'counterparty': parameters.get('counterparty', contract.get('counterparty', 'Counterparty')),
            'position': parameters.get('position', contract.get('position', 'Party A')),
            'clauses': clauses
        }

        # Generate report based on type
        if report_type == 'risk_review':
            from generate_risk_review import generate_risk_review_report

            # Add risk review specific data
            report_data.update({
                'top_concerns': [
                    {'clause_type': c['type'], 'concern': f"{c['type']} needs review"}
                    for c in clauses[:3]
                ],
                'distribution': clause_summary.get('by_weight', {'Critical': 0, 'High': 0, 'Standard': 0}),
                'negotiation_playbook': {
                    'your_leverage': ['Market competition', 'Long-term value'],
                    'counterparty_leverage': ['Specialized expertise'],
                    'sequence': [],
                    'tradeoffs': []
                }
            })

            # Create output path in REPORTS_DIR
            from datetime import datetime
            contract_name = contract.get('title', 'Contract').replace(' ', '_')
            filename = f"{contract_name}_Risk_Review_{datetime.now().strftime('%Y%m%d')}.docx"
            output_path = str(REPORTS_DIR / filename)

            output_path = generate_risk_review_report(report_data, output_path)

        elif report_type == 'redline':
            from generate_redline import generate_redline_report
            from datetime import datetime

            # Add redline specific data
            report_data.update({
                'revision_impact': {
                    'before': {'CRITICAL': 2, 'HIGH': 3, 'MODERATE': 2, 'LOW': 1},
                    'after': {'CRITICAL': 0, 'HIGH': 1, 'MODERATE': 4, 'LOW': 3}
                },
                'changes_proposed': {
                    'total': len(clauses),
                    'dealbreaker': 0,
                    'industry_standard': len(clauses),
                    'nice_to_have': 0
                },
                'key_revisions': [],
                'redlines': [],
                'implementation_notes': {},
                'negotiation_guide': {}
            })

            # Create output path in REPORTS_DIR
            contract_name = contract.get('title', 'Contract').replace(' ', '_')
            filename = f"{contract_name}_Redlines_{datetime.now().strftime('%Y%m%d')}.docx"
            output_path = str(REPORTS_DIR / filename)

            output_path = generate_redline_report(report_data, output_path)

        elif report_type == 'comparison':
            from generate_comparison import generate_comparison_report
            from datetime import datetime

            # Add comparison specific data
            report_data.update({
                'v1_label': 'Version 1',
                'v2_label': 'Version 2',
                'version_delta': {
                    'v1': {'CRITICAL': 2, 'HIGH': 3, 'MODERATE': 2, 'LOW': 1},
                    'v2': {'CRITICAL': 1, 'HIGH': 1, 'MODERATE': 4, 'LOW': 2}
                },
                'changes_detected': {
                    'total': 0,
                    'additions': 0,
                    'modifications': 0,
                    'deletions': 0
                },
                'key_themes': {},
                'comparisons': []
            })

            # Create output path in REPORTS_DIR
            contract_name = contract.get('title', 'Contract').replace(' ', '_')
            v1_label = parameters.get('v1_label', 'V1').replace(' ', '_')
            v2_label = parameters.get('v2_label', 'V2').replace(' ', '_')
            filename = f"{contract_name}_{v1_label}_to_{v2_label}_Comparison_{datetime.now().strftime('%Y%m%d')}.docx"
            output_path = str(REPORTS_DIR / filename)

            output_path = generate_comparison_report(report_data, output_path)

        else:
            conn.close()
            return jsonify({'error': f'Invalid report type: {report_type}'}), 400

        # Save report metadata to database
        import os
        filename = os.path.basename(output_path)

        cursor.execute("""
            INSERT INTO generated_reports (contract_id, output_filename, perspective)
            VALUES (?, ?, ?)
        """, (contract_id, filename, 'A'))

        report_id = cursor.lastrowid
        conn.commit()
        conn.close()

        logger.info(f"[REPORTS-SKILL] Generated report {report_id}: {filename}")

        return jsonify({
            'status': 'success',
            'report_id': report_id,
            'filename': filename
        }), 200

    except ImportError as e:
        logger.error(f"[REPORTS-SKILL] Import failed: {e}", exc_info=True)
        return jsonify({'error': f'Document generation skill not found: {str(e)}'}), 500
    except Exception as e:
        logger.error(f"[REPORTS-SKILL] Generation failed: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/reports/generate', methods=['POST'])
def generate_report():
    """
    Generate a report from a template and contract OR using document generation skill.

    JSON Body (Template-based):
        - template_id: Template to use
        - contract_id: Contract to pull data from
        - perspective: 'A' or 'B' (determines party resolution)

    JSON Body (Skill-based):
        - report_type: 'risk_review', 'redline', or 'comparison'
        - contract_id: Contract to analyze
        - parameters: Dict with our_entity, counterparty, position, etc.

    Returns:
        {
            status: 'success',
            report_id: int,
            filename: str
        }
    """
    logger.info("[REPORTS] Generate request received")

    data = request.get_json()
    if not data:
        return jsonify({'error': 'No data provided'}), 400

    # Check if this is skill-based (report_type) or template-based (template_id)
    report_type = data.get('report_type')
    template_id = data.get('template_id')
    contract_id = data.get('contract_id')

    if not contract_id:
        return jsonify({'error': 'contract_id required'}), 400

    # Route to skill-based generation if report_type is present
    if report_type:
        return generate_skill_report(data)

    # Otherwise, use template-based generation (existing logic)
    perspective = data.get('perspective', 'A')

    if not template_id:
        return jsonify({'error': 'template_id required for template-based reports'}), 400

    if perspective not in ['A', 'B']:
        return jsonify({'error': 'perspective must be A or B'}), 400

    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        # Get template
        cursor.execute("""
            SELECT source_filename, field_mappings
            FROM report_templates
            WHERE id = ?
        """, (template_id,))
        template_row = cursor.fetchone()

        if not template_row:
            conn.close()
            return jsonify({'error': 'Template not found'}), 404

        template_filename = template_row[0]
        field_mappings = json.loads(template_row[1]) if template_row[1] else {}

        # Get contract data
        cursor.execute("SELECT * FROM contracts WHERE id = ?", (contract_id,))
        contract_row = cursor.fetchone()

        if not contract_row:
            conn.close()
            return jsonify({'error': 'Contract not found'}), 404

        contract = dict(contract_row)

        # Build token values based on perspective
        our_entity = contract.get('our_entity', '')
        counterparty = contract.get('counterparty', '')
        our_role = contract.get('position', '')

        # Perspective swap
        if perspective == 'A':
            our_party_name = our_entity
            other_party_name = counterparty
            our_party_role = our_role
            other_party_role = 'counterparty'
        else:  # B
            our_party_name = counterparty
            other_party_name = our_entity
            our_party_role = 'counterparty'
            other_party_role = our_role

        # Build replacement dict
        replacements = {
            '{{OUR_PARTY_NAME}}': our_party_name or '',
            '{{OTHER_PARTY_NAME}}': other_party_name or '',
            '{{OUR_PARTY_ROLE}}': our_party_role or '',
            '{{OTHER_PARTY_ROLE}}': other_party_role or '',
            '{{CONTRACT_TYPE}}': contract.get('contract_type', '') or '',
            '{{EFFECTIVE_DATE}}': str(contract.get('effective_date', '')) or '',
            '{{EXPIRATION_DATE}}': str(contract.get('expiration_date', '')) or '',
            '{{CLM_STAGE}}': contract.get('clm_stage', '') or '',
            '{{RELATIONSHIP_TYPE}}': contract.get('party_relationship', '') or '',
            '{{PURPOSE}}': contract.get('contract_purpose', '') or '',
            '{{TOTAL_VALUE}}': f"${contract.get('contract_value', 0):,.2f}" if contract.get('contract_value') else '',
            '{{SUMMARY}}': contract.get('summary', '') or '',
            '{{CONTRACT_TITLE}}': contract.get('title', '') or '',
            '{{FILENAME}}': contract.get('filename', '') or '',
        }

        # Add any custom mappings from field_mappings
        for placeholder, field in field_mappings.items():
            if placeholder not in replacements and field:
                replacements[placeholder] = str(contract.get(field, '')) or ''

        # Load and process template
        template_path = TEMPLATES_DIR / template_filename
        if not template_path.exists():
            conn.close()
            return jsonify({'error': f'Template file not found: {template_filename}'}), 404

        doc = DocxDocument(str(template_path))

        # Replace in paragraphs
        for para in doc.paragraphs:
            for placeholder, value in replacements.items():
                if placeholder in para.text:
                    # Handle runs to preserve formatting
                    for run in para.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, value)

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for placeholder, value in replacements.items():
                            if placeholder in para.text:
                                for run in para.runs:
                                    if placeholder in run.text:
                                        run.text = run.text.replace(placeholder, value)

        # Generate output filename
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        safe_title = re.sub(r'[^\w\s-]', '', contract.get('title', 'report'))[:30]
        output_filename = f"report_{contract_id}_{safe_title}_{timestamp}.docx"
        output_path = REPORTS_DIR / output_filename

        # Save generated document
        doc.save(str(output_path))

        # Record in database
        cursor.execute("""
            INSERT INTO generated_reports (template_id, contract_id, perspective, output_filename)
            VALUES (?, ?, ?, ?)
        """, (template_id, contract_id, perspective, output_filename))

        report_id = cursor.lastrowid
        conn.commit()
        conn.close()

        logger.info(f"[REPORTS] Generated report {report_id}: {output_filename}")

        return jsonify({
            'status': 'success',
            'report_id': report_id,
            'output_filename': output_filename,
            'download_path': str(output_path)
        }), 200

    except Exception as e:
        logger.error(f"[REPORTS] Generate failed: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/reports', methods=['GET'])
def list_reports():
    """List generated reports."""
    logger.info("[REPORTS] List request received")

    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        cursor.execute("""
            SELECT gr.id, gr.template_id, gr.contract_id, gr.perspective,
                   gr.output_filename, gr.generated_at,
                   rt.name as template_name,
                   c.title as contract_title
            FROM generated_reports gr
            LEFT JOIN report_templates rt ON gr.template_id = rt.id
            LEFT JOIN contracts c ON gr.contract_id = c.id
            ORDER BY gr.generated_at DESC
            LIMIT 50
        """)

        reports = []
        for row in cursor.fetchall():
            reports.append({
                'id': row[0],
                'template_id': row[1],
                'contract_id': row[2],
                'perspective': row[3],
                'output_filename': row[4],
                'generated_at': row[5],
                'template_name': row[6],
                'contract_title': row[7]
            })

        conn.close()

        return jsonify({'reports': reports}), 200

    except Exception as e:
        logger.error(f"[REPORTS] List failed: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/reports/<int:report_id>/download', methods=['GET'])
def download_report(report_id):
    """Download a generated report."""
    logger.info(f"[REPORTS] Download request for report {report_id}")

    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        cursor.execute("SELECT output_filename FROM generated_reports WHERE id = ?", (report_id,))
        row = cursor.fetchone()
        conn.close()

        if not row:
            return jsonify({'error': 'Report not found'}), 404

        filename = row[0]
        filepath = REPORTS_DIR / filename

        if not filepath.exists():
            return jsonify({'error': 'Report file not found'}), 404

        return send_file(
            str(filepath),
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        logger.error(f"[REPORTS] Download failed: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/reports/tokens', methods=['GET'])
def get_report_tokens():
    """Return available tokens for template mapping."""
    return jsonify({
        'tokens': list(REPORT_TOKENS.keys()),
        'token_mappings': REPORT_TOKENS
    }), 200


# ============================================================================
# COMPARE V3 ENDPOINT (Phase 4E)
# ============================================================================

@app.route('/api/compare-v3', methods=['POST'])
def compare_v3_endpoint():
    """
    Compare v3 Engine - Phase 4F Implementation

    Phase 4F: Real intelligence with SAE/ERCE/BIRL/FAR pipeline.
    Falls back to placeholders on failure.

    JSON Body:
        - v1_text: First contract version text (or v1_contract_id)
        - v2_text: Second contract version text (or v2_contract_id)

    Returns:
        JSON with Compare v3 pipeline results
    """
    logger.info("Compare v3 request received (Phase 4F)")

    try:
        from compare_v3_engine import compare_v3_api

        data = request.get_json()

        if not data:
            return jsonify({
                'success': False,
                'error': 'No data provided',
                'error_message_key': 'compare.payload_failure'
            }), 400

        # Get text directly or fetch from contract IDs
        v1_text = data.get('v1_text', '')
        v2_text = data.get('v2_text', '')

        # Get contract IDs for clause lookup
        v1_id = data.get('v1_contract_id')
        v2_id = data.get('v2_contract_id')

        if v1_id and not v1_text:
            contract = get_contract_by_id(v1_id)
            if contract and contract.get('full_text'):
                v1_text = contract['full_text']

        if v2_id and not v2_text:
            contract = get_contract_by_id(v2_id)
            if contract and contract.get('full_text'):
                v2_text = contract['full_text']

        if not v1_text or not v2_text:
            return jsonify({
                'success': False,
                'error': 'Both v1_text and v2_text required (or valid contract IDs)',
                'error_message_key': 'compare.payload_failure'
            }), 400

        # Run Compare v3 pipeline (Phase 4F: pass contract IDs for clause lookup)
        result = compare_v3_api(v1_text, v2_text, v1_id, v2_id)

        if result.get('success'):
            logger.info("Compare v3 pipeline completed successfully")
            return jsonify(result), 200
        else:
            logger.warning(f"Compare v3 pipeline failed: {result.get('error_detail')}")
            return jsonify(result), 500

    except ImportError as e:
        logger.error(f"Compare v3 import error: {e}")
        return jsonify({
            'success': False,
            'error': 'Compare v3 engine not available',
            'error_message_key': 'compare.internal_failure',
            'error_detail': str(e)
        }), 500

    except Exception as e:
        logger.error(f"Compare v3 error: {e}")
        return jsonify({
            'success': False,
            'error': str(e),
            'error_message_key': 'compare.internal_failure'
        }), 500


# ============================================================================
# INTELLIGENT INTAKE ENDPOINTS
# ============================================================================

@app.route('/api/intake/upload', methods=['POST'])
def intake_upload():
    """
    Upload a contract file for intelligent intake.

    Accepts: .docx files only
    Returns: {file_id, filename, status}
    """
    logger.info("[INTAKE] Upload request received")

    try:
        # Validate file present
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400

        file = request.files['file']

        if not file.filename:
            return jsonify({'error': 'No file selected'}), 400

        # Validate file type (.docx only for now)
        if not file.filename.lower().endswith('.docx'):
            return jsonify({
                'error': 'Unsupported file type',
                'details': 'Only .docx files are currently supported',
                'filename': file.filename
            }), 400

        # Save file
        file_path = save_uploaded_file(file)

        # Generate a simple file_id (timestamp-based)
        import hashlib
        file_id = hashlib.md5(str(file_path).encode()).hexdigest()[:12]

        logger.info(f"[INTAKE] File uploaded: {file_path} (ID: {file_id})")

        return jsonify({
            'status': 'success',
            'file_id': file_id,
            'filename': file.filename,
            'file_path': str(file_path)
        }), 200

    except Exception as e:
        logger.error(f"[INTAKE] Upload failed: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/intake/extract', methods=['POST'])
def intake_extract():
    """
    Extract metadata from uploaded contract.

    Body: {file_path: string}
    Returns: {extraction: {...}, related_contracts: [...]}
    """
    logger.info("[INTAKE] Extract request received")

    try:
        # Check dependencies
        deps = check_dependencies()
        if not deps.get('python-docx'):
            return jsonify({
                'error': 'python-docx not installed',
                'action': 'Run: pip install python-docx'
            }), 503

        if not deps.get('api_key_configured'):
            return jsonify({
                'error': 'ANTHROPIC_API_KEY not configured',
                'action': 'Set ANTHROPIC_API_KEY in .env file'
            }), 503

        # Get file path from request
        data = request.get_json()
        if not data or 'file_path' not in data:
            return jsonify({'error': 'file_path required'}), 400

        file_path = Path(data['file_path'])

        if not file_path.exists():
            return jsonify({'error': f'File not found: {file_path}'}), 404

        # Run extraction
        logger.info(f"[INTAKE] Extracting metadata from: {file_path}")
        result = extract_contract_metadata(file_path)

        # Find related contracts if we have party names
        related = []
        if result.success and result.parties:
            related = find_related_contracts(result.parties)
            logger.info(f"[INTAKE] Found {len(related)} related contracts")

        return jsonify({
            'status': 'success' if result.success else 'partial',
            'extraction': result.to_dict(),
            'related_contracts': related
        }), 200

    except ImportError as e:
        logger.error(f"[INTAKE] Import error: {e}")
        return jsonify({
            'error': 'Extraction service not available',
            'details': str(e)
        }), 503

    except Exception as e:
        logger.error(f"[INTAKE] Extract failed: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/intake/store', methods=['POST'])
def intake_store():
    """
    Store confirmed contract metadata in database.

    Body: {
        file_path: string,
        title: string,
        contract_type: string,
        purpose: string,
        party: string,
        counterparty: string,
        orientation: string,
        leverage: string,
        related_contract_id: int|null,
        full_text: string (optional)
    }

    Returns: {contract_id: int, suggested_next: string}
    """
    logger.info("[INTAKE] Store request received")

    try:
        data = request.get_json()

        if not data:
            return jsonify({'error': 'No data provided'}), 400

        # Required fields
        required = ['file_path', 'title', 'contract_type', 'party', 'counterparty', 'orientation']
        missing = [f for f in required if not data.get(f)]

        if missing:
            return jsonify({
                'error': 'Missing required fields',
                'missing': missing
            }), 400

        # Extract filename from path
        file_path = Path(data['file_path'])
        filename = file_path.name

        # Get full text if not provided
        full_text = data.get('full_text', '')
        if not full_text and file_path.exists():
            try:
                full_text, _ = extract_text_from_docx(file_path)
            except:
                pass

        # Store in database
        conn = get_db_connection()
        cursor = conn.cursor()

        # Column names matched to actual schema (verified 12/15/25)
        cursor.execute("""
            INSERT INTO contracts (
                filename, filepath, title, contract_type, contract_purpose,
                our_entity, counterparty, party_relationship, leverage,
                parent_id, status, parsed_metadata, created_at
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, datetime('now'))
        """, (
            filename,
            str(file_path),
            data['title'],
            data['contract_type'],
            data.get('purpose'),
            data['party'],           # Maps to 'our_entity' column
            data['counterparty'],
            data['orientation'],     # Maps to 'party_relationship' (CUSTOMER/VENDOR)
            data.get('leverage'),
            data.get('related_contract_id'),  # Maps to 'parent_id'
            'intake',                # Initial status per schema default
            full_text[:50000] if full_text else None  # Store in parsed_metadata
        ))

        contract_id = cursor.lastrowid
        conn.commit()
        conn.close()

        logger.info(f"[INTAKE] Contract stored: ID={contract_id}, Title={data['title']}")

        # Determine suggested next step
        suggested_next = "Risk Analysis"
        if data['contract_type'] == 'Amendment':
            suggested_next = "Compare Versions"

        return jsonify({
            'status': 'success',
            'contract_id': contract_id,
            'suggested_next': suggested_next
        }), 201

    except sqlite3.IntegrityError as e:
        logger.error(f"[INTAKE] Database integrity error: {e}")
        return jsonify({'error': 'Database constraint violation', 'details': str(e)}), 409

    except Exception as e:
        logger.error(f"[INTAKE] Store failed: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500


@app.route('/api/intake/status', methods=['GET'])
def intake_status():
    """
    Check intake service status and dependencies.

    Returns: {ready: bool, dependencies: {...}}
    """
    try:
        deps = check_dependencies()

        ready = all(deps.values())

        return jsonify({
            'ready': ready,
            'dependencies': deps
        }), 200 if ready else 503

    except ImportError:
        return jsonify({
            'ready': False,
            'error': 'extraction_service not found'
        }), 503


# ============================================================================
# REGISTER BLUEPRINTS
# ============================================================================

try:
    from diagnostics_api import register_diagnostics
    register_diagnostics(app)
    logger.info("Diagnostics API endpoints registered")
except ImportError as e:
    logger.warning(f"Diagnostics API not available: {e}")

try:
    from dashboard_api import register_dashboard
    register_dashboard(app)
    logger.info("Dashboard API endpoints registered")
except ImportError as e:
    logger.warning(f"Dashboard API not available: {e}")

try:
    from export_api import register_export
    register_export(app)
    logger.info("Export API endpoints registered")
except ImportError as e:
    logger.warning(f"Export API not available: {e}")

try:
    from health_api import register_health
    register_health(app)
    logger.info("Health Score API endpoints registered")
except ImportError as e:
    logger.warning(f"Health Score API not available: {e}")

try:
    from qa_api import register_qa
    register_qa(app)
    logger.info("QA/QC API endpoints registered")
except ImportError as e:
    logger.warning(f"QA/QC API not available: {e}")


# ============================================================================
# MAIN
# ============================================================================

if __name__ == '__main__':
    logger.info("="*60)
    logger.info("CIP Backend API Starting")
    logger.info("="*60)
    logger.info(f"Host: {API_HOST}")
    logger.info(f"Port: {API_PORT}")
    logger.info(f"Debug: {DEBUG_MODE}")
    logger.info(f"Orchestrator: {'Initialized' if orchestrator else 'Failed'}")
    logger.info(f"API Key: {'Configured' if ANTHROPIC_API_KEY else 'Missing'}")
    logger.info("="*60)

    app.run(
        host=API_HOST,
        port=API_PORT,
        debug=DEBUG_MODE
    )
