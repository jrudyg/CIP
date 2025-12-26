#!/usr/bin/env python3
"""
Generate Version Comparison report (.docx)

Components:
1. Title Page
2. Executive Summary (version delta, changes detected, key themes by clause priority)
3. Combined Risk Matrix (V1 | V2 | Delta)
4. Detailed Comparison Table (5-column with related clauses)
5. Disclaimers

Usage:
    from generate_comparison import generate_comparison_report

    report_data = {
        'contract_name': 'Master Services Agreement',
        'v1_label': 'October Draft',
        'v2_label': 'November Final',
        'comparisons': [...],  # Clause-by-clause comparisons
        ...
    }

    docx_path = generate_comparison_report(report_data)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
from typing import Dict, List

# Risk level colors
COLORS = {
    'CRITICAL': RGBColor(192, 0, 0),
    'HIGH': RGBColor(237, 125, 49),
    'MODERATE': RGBColor(46, 117, 182),
    'LOW': RGBColor(0, 176, 80),
    'HEADER_BG': '1F4E79',
    'HEADER_TEXT': RGBColor(255, 255, 255),
    'DELETION': RGBColor(255, 0, 0),
    'ADDITION': RGBColor(0, 176, 80)
}

RISK_SYMBOL = '●'
DELTA_SYMBOLS = {'increased': '▲', 'decreased': '▼', 'unchanged': '●'}


def set_cell_shading(cell, color_hex):
    """Add background color to table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_risk_indicator(paragraph, risk_level: str):
    """Add colored risk indicator symbol to paragraph."""
    run = paragraph.add_run(f'{RISK_SYMBOL} ')
    run.font.size = Pt(12)
    run.font.color.rgb = COLORS[risk_level]
    run.bold = True


def add_redline_text(paragraph, v1_text: str, v2_text: str):
    """Add V2 text with redline markings showing changes from V1."""
    # Simplified implementation - in production, use proper diff algorithm
    if v1_text != v2_text:
        # Show deletion
        if v1_text:
            run = paragraph.add_run(v1_text)
            run.font.strike = True
            run.font.color.rgb = COLORS['DELETION']
            run.font.size = Pt(10)
            paragraph.add_run(' ')

        # Show addition
        if v2_text:
            run = paragraph.add_run(v2_text)
            run.bold = True
            run.font.color.rgb = COLORS['ADDITION']
            run.font.size = Pt(10)
    else:
        # No change
        run = paragraph.add_run(v2_text)
        run.font.size = Pt(10)


def add_title_page(doc: Document, data: Dict):
    """Add title page to document."""
    # Contract name
    p = doc.add_paragraph()
    run = p.add_run(data['contract_name'])
    run.font.size = Pt(24)
    run.bold = True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.space_after = Pt(24)

    # Report type
    p = doc.add_paragraph()
    run = p.add_run('VERSION COMPARISON REPORT')
    run.font.size = Pt(16)
    run.bold = True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.space_after = Pt(12)

    # Version info
    p = doc.add_paragraph()
    run = p.add_run(f"{data['v1_label']} → {data['v2_label']}")
    run.font.size = Pt(14)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.space_after = Pt(48)

    # Date
    p = doc.add_paragraph()
    run = p.add_run(f"Date: {datetime.now().strftime('%B %d, %Y')}")
    run.font.size = Pt(12)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.space_after = Pt(24)

    # Entity information
    p = doc.add_paragraph()
    run = p.add_run(f"Our Entity: {data['our_entity']}\n")
    run.font.size = Pt(12)
    run = p.add_run(f"Counterparty: {data['counterparty']}\n")
    run.font.size = Pt(12)
    run = p.add_run(f"Position: {data['position']}")
    run.font.size = Pt(12)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_page_break()


def add_executive_summary(doc: Document, data: Dict):
    """Add executive summary with version delta and key themes."""
    # Section header
    p = doc.add_paragraph()
    run = p.add_run('EXECUTIVE SUMMARY')
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = COLORS['MODERATE']
    p.space_after = Pt(12)

    # Version delta table
    p = doc.add_paragraph()
    run = p.add_run('VERSION DELTA')
    run.font.size = Pt(11)
    run.bold = True

    delta = data.get('version_delta', {})
    v1 = delta.get('v1', {})
    v2 = delta.get('v2', {})

    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'

    # Header
    headers = ['', 'V1', 'V2', 'Delta']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, COLORS['HEADER_BG'])
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.color.rgb = COLORS['HEADER_TEXT']
        run.bold = True
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Data rows
    risk_levels = ['CRITICAL', 'HIGH', 'MODERATE', 'LOW']
    for i, risk_level in enumerate(risk_levels, 1):
        row = table.rows[i]

        # Risk level with indicator
        p = row.cells[0].paragraphs[0]
        add_risk_indicator(p, risk_level)
        run = p.add_run(risk_level)
        run.font.size = Pt(10)

        # V1 count
        v1_count = v1.get(risk_level, 0)
        row.cells[1].text = str(v1_count)
        row.cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # V2 count
        v2_count = v2.get(risk_level, 0)
        row.cells[2].text = str(v2_count)
        row.cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Delta
        delta_val = v2_count - v1_count
        row.cells[3].text = f"{delta_val:+d}"
        row.cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph()

    # Changes detected
    p = doc.add_paragraph()
    run = p.add_run('CHANGES DETECTED')
    run.font.size = Pt(11)
    run.bold = True

    changes = data.get('changes_detected', {})
    p = doc.add_paragraph(
        f"Total: {changes.get('total', 0)}  |  "
        f"Additions: {changes.get('additions', 0)}  |  "
        f"Modifications: {changes.get('modifications', 0)}  |  "
        f"Deletions: {changes.get('deletions', 0)}"
    )
    p.paragraph_format.left_indent = Inches(0.25)

    doc.add_paragraph()

    # Key themes by clause priority
    p = doc.add_paragraph()
    run = p.add_run('KEY THEMES (by clause priority)')
    run.font.size = Pt(11)
    run.bold = True

    themes = data.get('key_themes', {})

    # Critical weight
    p = doc.add_paragraph()
    run = p.add_run('Critical Weight:')
    run.font.size = Pt(11)
    run.bold = True

    for clause_type in ['Indemnification', 'Limitation of Liability', 'IP/Work Ownership']:
        theme = themes.get('critical', {}).get(clause_type, 'No changes')
        p = doc.add_paragraph(f"• {clause_type}: {theme}", style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)

    # High weight
    p = doc.add_paragraph()
    run = p.add_run('High Weight:')
    run.font.size = Pt(11)
    run.bold = True

    for clause_type in ['Termination', 'Insurance', 'Vendor Displacement', 'Non-Solicitation']:
        theme = themes.get('high', {}).get(clause_type, 'No changes')
        p = doc.add_paragraph(f"• {clause_type}: {theme}", style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)

    # Standard weight summary
    standard_summary = themes.get('standard', {})
    p = doc.add_paragraph()
    run = p.add_run('Standard Weight: ')
    run.font.size = Pt(11)
    run.bold = True
    run = p.add_run(f"{standard_summary.get('count', 0)} changes across {standard_summary.get('types', 0)} clause types")
    run.font.size = Pt(11)

    doc.add_paragraph()


def add_risk_matrix(doc: Document, data: Dict):
    """Add combined risk matrix (V1 | V2 | Delta)."""
    p = doc.add_paragraph()
    run = p.add_run('COMBINED RISK MATRIX')
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = COLORS['MODERATE']
    p.space_after = Pt(12)

    # Create table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    # Header row
    header_cells = table.rows[0].cells
    headers = ['Clause Type', 'V1', 'V2', 'Delta']

    for i, header in enumerate(headers):
        cell = header_cells[i]
        set_cell_shading(cell, COLORS['HEADER_BG'])
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.color.rgb = COLORS['HEADER_TEXT']
        run.bold = True
        run.font.size = Pt(10)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Data rows (clauses with changes only)
    for comparison in data.get('comparisons', []):
        if comparison.get('v1_risk') != comparison.get('v2_risk'):
            row = table.add_row()
            cells = row.cells

            # Clause type
            cells[0].text = comparison['clause_type']

            # V1 risk
            p = cells[1].paragraphs[0]
            add_risk_indicator(p, comparison['v1_risk'])
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # V2 risk
            p = cells[2].paragraphs[0]
            add_risk_indicator(p, comparison['v2_risk'])
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Delta
            delta = comparison.get('delta', 'unchanged')
            p = cells[3].paragraphs[0]
            run = p.add_run(DELTA_SYMBOLS[delta])
            run.font.size = Pt(12)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # No changes section
    p = doc.add_paragraph()
    run = p.add_run('NO CHANGES REQUIRED')
    run.font.size = Pt(11)
    run.bold = True

    no_changes = [c for c in data.get('comparisons', []) if c.get('v1_risk') == c.get('v2_risk')]
    for comparison in no_changes[:6]:  # Show first 6
        p = doc.add_paragraph(f"• {comparison.get('section', '')} {comparison['clause_type']}", style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)

    doc.add_paragraph()


def add_comparison_table(doc: Document, data: Dict):
    """Add detailed comparison table (5-column with related clauses)."""
    p = doc.add_paragraph()
    run = p.add_run('DETAILED COMPARISON')
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = COLORS['MODERATE']
    p.space_after = Pt(12)

    # Create table
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'

    # Header row
    header_cells = table.rows[0].cells
    headers = ['#', 'Section / Category', f"V1 ({data['v1_label']})", f"V2 ({data['v2_label']}) - Redlined", 'Business Impact']

    for i, header in enumerate(headers):
        cell = header_cells[i]
        set_cell_shading(cell, COLORS['HEADER_BG'])
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.color.rgb = COLORS['HEADER_TEXT']
        run.bold = True
        run.font.size = Pt(10)

    # Data rows
    for i, comparison in enumerate(data.get('comparisons', []), 1):
        row = table.add_row()
        cells = row.cells

        # Number
        cells[0].text = str(i)
        cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Section / Category
        section_text = f"{comparison.get('section', '')} - {comparison['clause_type']}"
        cells[1].text = section_text

        # Related clauses (if any)
        if comparison.get('related_sections'):
            p = cells[1].add_paragraph()
            run = p.add_run(f"Related: {', '.join(comparison['related_sections'])}")
            run.font.size = Pt(9)
            run.font.italic = True

        # V1 text
        cells[2].text = comparison.get('v1_text', '')[:200]  # Truncate long text
        if len(comparison.get('v1_text', '')) > 200:
            cells[2].text += '...'

        # V2 redlined text
        p = cells[3].paragraphs[0]
        add_redline_text(p, comparison.get('v1_text', ''), comparison.get('v2_text', ''))

        # Business impact
        cells[4].text = comparison.get('business_impact', '')

    doc.add_paragraph()


def add_disclaimers(doc: Document):
    """Add disclaimer section."""
    doc.add_page_break()

    p = doc.add_paragraph()
    run = p.add_run('DISCLAIMER')
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = COLORS['MODERATE']
    p.space_after = Pt(12)

    disclaimers = [
        ("Risk Advisory", "This report applies generally accepted contract risk categories and contract management best practices. Stakeholders must use this analysis to make informed business decisions regarding risk acceptance and mitigation. Failure to address identified risks may result in material financial, operational, or legal exposure."),
        ("Legal", "This analysis is for informational purposes only and does not constitute legal advice. Consult qualified legal counsel before making decisions based on this report."),
        ("AI-Generated", "This report was generated with AI assistance. All findings should be verified against source documents."),
        ("Confidential", "This document contains confidential analysis. Do not distribute without authorization.")
    ]

    for title, text in disclaimers:
        p = doc.add_paragraph()
        run = p.add_run(f"{title}: ")
        run.bold = True
        run.font.size = Pt(11)
        run = p.add_run(text)
        run.font.size = Pt(11)
        p.space_after = Pt(6)


def generate_comparison_report(data: Dict, output_path: str = None) -> str:
    """
    Generate Version Comparison report as .docx file.

    Args:
        data: Report data dictionary with keys:
            - contract_name: str
            - v1_label: str (e.g., "October Draft")
            - v2_label: str (e.g., "November Final")
            - our_entity: str
            - counterparty: str
            - position: str
            - version_delta: Dict (V1/V2 risk counts)
            - changes_detected: Dict (total, additions, modifications, deletions)
            - key_themes: Dict (by weight tier)
            - comparisons: List[Dict] (clause-by-clause comparisons)
        output_path: Path to save .docx file (optional)

    Returns:
        Path to generated .docx file
    """
    # Create document
    doc = Document()

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Add components
    add_title_page(doc, data)
    add_executive_summary(doc, data)
    add_risk_matrix(doc, data)
    add_comparison_table(doc, data)
    add_disclaimers(doc)

    # Save document
    if output_path is None:
        contract_name = data['contract_name'].replace(' ', '_')
        output_path = f"{contract_name}_V1_to_V2_Comparison_{datetime.now().strftime('%Y%m%d')}.docx"

    doc.save(output_path)
    return output_path


# Example usage
if __name__ == '__main__':
    sample_data = {
        'contract_name': 'Master Services Agreement',
        'v1_label': 'October Draft',
        'v2_label': 'November Final',
        'our_entity': 'ACME Corporation',
        'counterparty': 'Vendor Inc',
        'position': 'Buyer',
        'version_delta': {
            'v1': {'CRITICAL': 2, 'HIGH': 3, 'MODERATE': 2, 'LOW': 1},
            'v2': {'CRITICAL': 1, 'HIGH': 1, 'MODERATE': 4, 'LOW': 2}
        },
        'changes_detected': {
            'total': 11,
            'additions': 2,
            'modifications': 8,
            'deletions': 1
        },
        'key_themes': {
            'critical': {
                'Indemnification': 'Added $5M liability cap',
                'Limitation of Liability': 'Increased cap from 3x to 12x monthly fees',
                'IP/Work Ownership': 'No changes'
            },
            'high': {
                'Termination': 'Extended cure period from 15 to 30 days',
                'Insurance': 'No changes',
                'Vendor Displacement': 'No changes',
                'Non-Solicitation': 'No changes'
            },
            'standard': {
                'count': 3,
                'types': 2
            }
        },
        'comparisons': [
            {
                'clause_type': 'Indemnification',
                'section': '8.1',
                'v1_risk': 'CRITICAL',
                'v2_risk': 'MODERATE',
                'delta': 'decreased',
                'v1_text': 'Vendor shall indemnify Client for all losses.',
                'v2_text': 'Vendor shall indemnify Client for all losses up to $5,000,000.',
                'business_impact': 'Significantly reduces financial exposure',
                'related_sections': ['9.2']
            },
            {
                'clause_type': 'Limitation of Liability',
                'section': '9.2',
                'v1_risk': 'HIGH',
                'v2_risk': 'MODERATE',
                'delta': 'decreased',
                'v1_text': 'Maximum liability shall not exceed 3x monthly fees.',
                'v2_text': 'Maximum liability shall not exceed 12x monthly fees.',
                'business_impact': 'Better protection for larger incidents',
                'related_sections': ['8.1']
            }
        ]
    }

    output_file = generate_comparison_report(sample_data, 'test_comparison.docx')
    print(f"Report generated: {output_file}")
