#!/usr/bin/env python3
"""
Generate Suggested Redlines and Revisions report (.docx)

Components:
1. Title Page
2. Executive Summary (Before/After impact, changes proposed)
3. Combined Risk Matrix (Before | After | Delta)
4. Redline Table (grouped by clause type)
5. Implementation Notes
6. Negotiation Guide
7. Disclaimers

Redline Notation:
- Deletion: Red #FF0000 strikethrough
- Addition: Green #00B050 bold
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
from typing import Dict, List

# Risk level colors
COLORS = {
    'CRITICAL': RGBColor(192, 0, 0),
    'HIGH': RGBColor(237, 125, 49),
    'MODERATE': RGBColor(46, 117, 182),
    'LOW': RGBColor(0, 176, 80),
    'HEADER_BG': '1F4E79',
    'HEADER_TEXT': RGBColor(255, 255, 255),
    'DELETION': RGBColor(255, 0, 0),
    'ADDITION': RGBColor(0, 176, 80)
}

RISK_SYMBOL = '●'
DELTA_SYMBOLS = {'increased': '▲', 'decreased': '▼', 'unchanged': '●'}


def set_cell_shading(cell, color_hex):
    """Add background color to table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_risk_indicator(paragraph, risk_level: str):
    """Add colored risk indicator symbol to paragraph."""
    run = paragraph.add_run(f'{RISK_SYMBOL} ')
    run.font.size = Pt(12)
    run.font.color.rgb = COLORS[risk_level]
    run.bold = True


def add_redline_text(paragraph, text: str, is_deletion: bool = False, is_addition: bool = False):
    """
    Add redlined text with proper formatting.

    Args:
        paragraph: Document paragraph
        text: Text to add
        is_deletion: If True, format as deletion (red strikethrough)
        is_addition: If True, format as addition (green bold)
    """
    run = paragraph.add_run(text)
    run.font.size = Pt(11)

    if is_deletion:
        run.font.strike = True
        run.font.color.rgb = COLORS['DELETION']
    elif is_addition:
        run.bold = True
        run.font.color.rgb = COLORS['ADDITION']


def add_title_page(doc: Document, data: Dict):
    """Add title page to document."""
    # Contract name
    p = doc.add_paragraph()
    run = p.add_run(data['contract_name'])
    run.font.size = Pt(24)
    run.bold = True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.space_after = Pt(24)

    # Report type
    p = doc.add_paragraph()
    run = p.add_run('SUGGESTED REDLINES AND REVISIONS')
    run.font.size = Pt(16)
    run.bold = True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.space_after = Pt(48)

    # Date
    p = doc.add_paragraph()
    run = p.add_run(f"Date: {datetime.now().strftime('%B %d, %Y')}")
    run.font.size = Pt(12)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.space_after = Pt(24)

    # Entity information
    p = doc.add_paragraph()
    run = p.add_run(f"Our Entity: {data['our_entity']}\n")
    run.font.size = Pt(12)
    run = p.add_run(f"Counterparty: {data['counterparty']}\n")
    run.font.size = Pt(12)
    run = p.add_run(f"Position: {data['position']}")
    run.font.size = Pt(12)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_page_break()


def add_executive_summary(doc: Document, data: Dict):
    """Add executive summary with revision impact."""
    # Section header
    p = doc.add_paragraph()
    run = p.add_run('EXECUTIVE SUMMARY')
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = COLORS['MODERATE']
    p.space_after = Pt(12)

    # Revision impact table
    p = doc.add_paragraph()
    run = p.add_run('REVISION IMPACT')
    run.font.size = Pt(11)
    run.bold = True

    impact = data.get('revision_impact', {})
    before = impact.get('before', {})
    after = impact.get('after', {})

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'

    # Header
    headers = ['', 'Before', 'After']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, COLORS['HEADER_BG'])
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.color.rgb = COLORS['HEADER_TEXT']
        run.bold = True
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Data rows
    risk_levels = ['CRITICAL', 'HIGH', 'MODERATE', 'LOW']
    for i, risk_level in enumerate(risk_levels, 1):
        row = table.rows[i]

        # Risk level with indicator
        p = row.cells[0].paragraphs[0]
        add_risk_indicator(p, risk_level)
        run = p.add_run(risk_level)
        run.font.size = Pt(10)

        # Before count
        row.cells[1].text = str(before.get(risk_level, 0))
        row.cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # After count
        row.cells[2].text = str(after.get(risk_level, 0))
        row.cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph()

    # Changes proposed
    p = doc.add_paragraph()
    run = p.add_run('CHANGES PROPOSED')
    run.font.size = Pt(11)
    run.bold = True

    changes = data.get('changes_proposed', {})
    p = doc.add_paragraph(
        f"Total: {changes.get('total', 0)}  |  "
        f"Dealbreaker: {changes.get('dealbreaker', 0)}  |  "
        f"Industry Standard: {changes.get('industry_standard', 0)}  |  "
        f"Nice-to-Have: {changes.get('nice_to_have', 0)}"
    )
    p.paragraph_format.left_indent = Inches(0.25)

    # Key revisions
    p = doc.add_paragraph()
    run = p.add_run('KEY REVISIONS')
    run.font.size = Pt(11)
    run.bold = True

    for i, revision in enumerate(data.get('key_revisions', [])[:3], 1):
        p = doc.add_paragraph(f"{i}. {revision['clause_type']}: {revision['summary']}", style='List Number')
        p.paragraph_format.left_indent = Inches(0.25)

    doc.add_paragraph()


def add_risk_matrix(doc: Document, data: Dict):
    """Add combined risk matrix (Before | After | Delta)."""
    p = doc.add_paragraph()
    run = p.add_run('COMBINED RISK MATRIX')
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = COLORS['MODERATE']
    p.space_after = Pt(12)

    # Create table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    # Header row
    header_cells = table.rows[0].cells
    headers = ['Clause Type', 'Before', 'After', 'Delta']

    for i, header in enumerate(headers):
        cell = header_cells[i]
        set_cell_shading(cell, COLORS['HEADER_BG'])
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.color.rgb = COLORS['HEADER_TEXT']
        run.bold = True
        run.font.size = Pt(10)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Data rows (clauses with changes)
    for clause in data.get('redlines', []):
        if clause.get('before_risk') != clause.get('after_risk'):
            row = table.add_row()
            cells = row.cells

            # Clause type
            cells[0].text = clause['type']

            # Before risk
            p = cells[1].paragraphs[0]
            add_risk_indicator(p, clause['before_risk'])
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # After risk
            p = cells[2].paragraphs[0]
            add_risk_indicator(p, clause['after_risk'])
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Delta
            delta = clause.get('delta', 'unchanged')
            p = cells[3].paragraphs[0]
            run = p.add_run(DELTA_SYMBOLS[delta])
            run.font.size = Pt(12)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # No changes section
    p = doc.add_paragraph()
    run = p.add_run('NO CHANGES REQUIRED')
    run.font.size = Pt(11)
    run.bold = True

    no_changes = [c for c in data.get('redlines', []) if c.get('before_risk') == c.get('after_risk')]
    for clause in no_changes[:5]:  # Show first 5
        p = doc.add_paragraph(f"• {clause.get('section', '')} {clause['type']}", style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)

    doc.add_paragraph()


def add_redline_table(doc: Document, data: Dict):
    """Add redline table grouped by clause type."""
    p = doc.add_paragraph()
    run = p.add_run('REDLINE DETAILS')
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = COLORS['MODERATE']
    p.space_after = Pt(12)

    # Group redlines by type
    grouped = {}
    for redline in data.get('redlines', []):
        clause_type = redline['type']
        if clause_type not in grouped:
            grouped[clause_type] = []
        grouped[clause_type].append(redline)

    # For each clause type
    for clause_type, redlines in grouped.items():
        # Clause type header
        p = doc.add_paragraph()
        p.add_run(clause_type.upper())
        p.runs[0].font.size = Pt(12)
        p.runs[0].bold = True

        # Create table
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'

        # Header row
        header_cells = table.rows[0].cells
        headers = ['Section', 'Risk', 'Original', 'Proposed Change', 'Rationale']

        for i, header in enumerate(headers):
            cell = header_cells[i]
            set_cell_shading(cell, COLORS['HEADER_BG'])
            p = cell.paragraphs[0]
            run = p.add_run(header)
            run.font.color.rgb = COLORS['HEADER_TEXT']
            run.bold = True
            run.font.size = Pt(10)

        # Data rows
        for redline in redlines:
            row = table.add_row()
            cells = row.cells

            # Section
            cells[0].text = redline.get('section', '')

            # Risk indicator (after)
            p = cells[1].paragraphs[0]
            add_risk_indicator(p, redline.get('after_risk', 'LOW'))

            # Original text
            cells[2].text = redline.get('original_text', '')

            # Proposed change with redline formatting
            p = cells[3].paragraphs[0]
            if 'proposed_change' in redline:
                # Parse redline format: ~~deletion~~ and **addition**
                text = redline['proposed_change']
                # Simplified parsing (in real implementation, use regex)
                add_redline_text(p, text)

            # Rationale
            cells[4].text = redline.get('rationale', '')

        doc.add_paragraph()


def add_implementation_notes(doc: Document, data: Dict):
    """Add implementation notes section."""
    p = doc.add_paragraph()
    run = p.add_run('IMPLEMENTATION NOTES')
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = COLORS['MODERATE']
    p.space_after = Pt(12)

    notes = data.get('implementation_notes', {})

    # Sequencing
    p = doc.add_paragraph()
    run = p.add_run('SEQUENCING')
    run.font.size = Pt(11)
    run.bold = True

    for i, item in enumerate(notes.get('sequencing', []), 1):
        p = doc.add_paragraph(f"{i}. ", style='List Number')
        add_risk_indicator(p, item['risk_level'])
        run = p.add_run(f"{item['section']} - {item['reason']}")
        run.font.size = Pt(11)

    # Dependencies
    p = doc.add_paragraph()
    run = p.add_run('DEPENDENCIES')
    run.font.size = Pt(11)
    run.bold = True

    for dep in notes.get('dependencies', []):
        doc.add_paragraph(f"• {dep}", style='List Bullet')

    # Notes
    p = doc.add_paragraph()
    run = p.add_run('NOTES')
    run.font.size = Pt(11)
    run.bold = True

    for note in notes.get('notes', []):
        doc.add_paragraph(f"• {note}", style='List Bullet')

    doc.add_paragraph()


def add_negotiation_guide(doc: Document, data: Dict):
    """Add negotiation guide with talking points and strategy."""
    p = doc.add_paragraph()
    run = p.add_run('NEGOTIATION GUIDE')
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = COLORS['MODERATE']
    p.space_after = Pt(12)

    guide = data.get('negotiation_guide', {})

    # Talking points
    p = doc.add_paragraph()
    run = p.add_run('TALKING POINTS')
    run.font.size = Pt(11)
    run.bold = True

    for point in guide.get('talking_points', []):
        p = doc.add_paragraph()
        add_risk_indicator(p, point['risk_level'])
        run = p.add_run(f"{point['clause']} ({point['section']})")
        run.font.size = Pt(11)
        run.bold = True

        for arg in point.get('arguments', []):
            doc.add_paragraph(f"  • {arg}", style='List Bullet')

    # Concession strategy
    p = doc.add_paragraph()
    run = p.add_run('CONCESSION STRATEGY')
    run.font.size = Pt(11)
    run.bold = True

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    # Header
    header_cells = table.rows[0].cells
    for i, header in enumerate(['GIVE', 'GET']):
        cell = header_cells[i]
        set_cell_shading(cell, COLORS['HEADER_BG'])
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.color.rgb = COLORS['HEADER_TEXT']
        run.bold = True
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Concessions
    for concession in guide.get('concessions', []):
        row = table.add_row()
        cells = row.cells

        # Give
        p = cells[0].paragraphs[0]
        add_risk_indicator(p, concession['give_risk'])
        run = p.add_run(concession['give'])
        run.font.size = Pt(10)

        # Get
        p = cells[1].paragraphs[0]
        add_risk_indicator(p, concession['get_risk'])
        run = p.add_run(concession['get'])
        run.font.size = Pt(10)

    doc.add_paragraph()

    # Walk-away triggers
    p = doc.add_paragraph()
    run = p.add_run('WALK-AWAY TRIGGERS')
    run.font.size = Pt(11)
    run.bold = True

    for trigger in guide.get('walkaways', []):
        p = doc.add_paragraph()
        add_risk_indicator(p, 'CRITICAL')
        run = p.add_run(trigger)
        run.font.size = Pt(11)

    doc.add_paragraph()


def add_disclaimers(doc: Document):
    """Add disclaimer section."""
    doc.add_page_break()

    p = doc.add_paragraph()
    run = p.add_run('DISCLAIMER')
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = COLORS['MODERATE']
    p.space_after = Pt(12)

    disclaimers = [
        ("Risk Advisory", "This report applies generally accepted contract risk categories and contract management best practices. Stakeholders must use this analysis to make informed business decisions regarding risk acceptance and mitigation. Failure to address identified risks may result in material financial, operational, or legal exposure."),
        ("Legal", "This analysis is for informational purposes only and does not constitute legal advice. Consult qualified legal counsel before making decisions based on this report."),
        ("AI-Generated", "This report was generated with AI assistance. All findings should be verified against source documents."),
        ("Confidential", "This document contains confidential analysis. Do not distribute without authorization.")
    ]

    for title, text in disclaimers:
        p = doc.add_paragraph()
        run = p.add_run(f"{title}: ")
        run.bold = True
        run.font.size = Pt(11)
        run = p.add_run(text)
        run.font.size = Pt(11)
        p.space_after = Pt(6)


def generate_redline_report(data: Dict, output_path: str = None) -> str:
    """
    Generate Suggested Redlines and Revisions report as .docx file.

    Args:
        data: Report data dictionary
        output_path: Path to save .docx file (optional)

    Returns:
        Path to generated .docx file
    """
    # Create document
    doc = Document()

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Add components
    add_title_page(doc, data)
    add_executive_summary(doc, data)
    add_risk_matrix(doc, data)
    add_redline_table(doc, data)
    add_implementation_notes(doc, data)
    add_negotiation_guide(doc, data)
    add_disclaimers(doc)

    # Save document
    if output_path is None:
        contract_name = data['contract_name'].replace(' ', '_')
        output_path = f"{contract_name}_Redlines_{datetime.now().strftime('%Y%m%d')}.docx"

    doc.save(output_path)
    return output_path


# Example usage
if __name__ == '__main__':
    sample_data = {
        'contract_name': 'Master Services Agreement',
        'our_entity': 'ACME Corporation',
        'counterparty': 'Vendor Inc',
        'position': 'Buyer',
        'revision_impact': {
            'before': {'CRITICAL': 2, 'HIGH': 3, 'MODERATE': 2, 'LOW': 1},
            'after': {'CRITICAL': 0, 'HIGH': 1, 'MODERATE': 4, 'LOW': 3}
        },
        'changes_proposed': {
            'total': 8,
            'dealbreaker': 2,
            'industry_standard': 4,
            'nice_to_have': 2
        },
        'key_revisions': [
            {'clause_type': 'Indemnification', 'summary': 'Add $5M liability cap'},
            {'clause_type': 'Limitation of Liability', 'summary': 'Increase cap to 12x monthly fees'},
            {'clause_type': 'Termination', 'summary': 'Extend cure period to 30 days'}
        ],
        'redlines': [
            {
                'type': 'Indemnification',
                'section': '8.1',
                'before_risk': 'CRITICAL',
                'after_risk': 'MODERATE',
                'delta': 'decreased',
                'original_text': 'Vendor shall indemnify Client for all losses.',
                'proposed_change': 'Vendor shall indemnify Client for all losses up to $5,000,000.',
                'rationale': 'Adds financial cap to limit exposure'
            }
        ],
        'implementation_notes': {
            'sequencing': [
                {'risk_level': 'CRITICAL', 'section': '8.1', 'reason': 'Highest financial exposure'}
            ],
            'dependencies': ['Section 8.1 must be agreed before 9.2'],
            'notes': ['Consider bundling indemnification and liability cap']
        },
        'negotiation_guide': {
            'talking_points': [
                {
                    'risk_level': 'CRITICAL',
                    'clause': 'Indemnification',
                    'section': '8.1',
                    'arguments': ['Industry standard includes caps', 'Unlimited exposure is uninsurable']
                }
            ],
            'concessions': [
                {'give': 'Shorter payment terms', 'give_risk': 'MODERATE', 'get': 'Liability cap', 'get_risk': 'CRITICAL'}
            ],
            'walkaways': ['No indemnification cap']
        }
    }

    output_file = generate_redline_report(sample_data, 'test_redline.docx')
    print(f"Report generated: {output_file}")
