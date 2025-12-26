#!/usr/bin/env python3
"""
Generate Contract Risk Review report (.docx)

Components:
1. Title Page
2. Executive Summary
3. Risk Heat Map
4. Clause Analysis Table (grouped by clause type)
5. Negotiation Playbook
6. Disclaimers

Usage:
    from generate_risk_review import generate_risk_review_report

    report_data = {
        'contract_name': 'Master Services Agreement',
        'our_entity': 'ACME Corp',
        'counterparty': 'Vendor Inc',
        'position': 'Buyer',
        'clauses': [...],  # From extract_clauses()
        'analysis': {...}  # Risk analysis results
    }

    docx_path = generate_risk_review_report(report_data, output_path='output.docx')
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
from typing import Dict, List
import json

# Risk level colors (RGB)
COLORS = {
    'CRITICAL': RGBColor(192, 0, 0),
    'HIGH': RGBColor(237, 125, 49),
    'MODERATE': RGBColor(46, 117, 182),
    'LOW': RGBColor(0, 176, 80),
    'HEADER_BG': '1F4E79',  # Navy (hex for cell shading)
    'HEADER_TEXT': RGBColor(255, 255, 255),
    'BORDER': RGBColor(208, 206, 206)
}

# Risk symbol (filled circle)
RISK_SYMBOL = '●'


def set_cell_shading(cell, color_hex):
    """Add background color to table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_risk_indicator(paragraph, risk_level: str):
    """Add colored risk indicator symbol to paragraph."""
    run = paragraph.add_run(f'{RISK_SYMBOL} ')
    run.font.size = Pt(12)
    run.font.color.rgb = COLORS[risk_level]
    run.bold = True


def calculate_overall_risk(clauses: List[Dict]) -> str:
    """
    Calculate overall risk using clause-weighted scoring logic.

    Rules:
    - IF any Critical-weight clause has CRITICAL risk → Overall: CRITICAL
    - ELSE IF any Critical-weight clause has HIGH risk → Overall: HIGH
    - ELSE IF any High-weight clause has CRITICAL risk → Overall: HIGH
    - ELSE IF multiple High-weight clauses have HIGH risk → Overall: HIGH
    - ELSE IF any clause has HIGH risk → Overall: MODERATE
    - ELSE → Overall: LOW
    """
    critical_weight_clauses = [c for c in clauses if c.get('weight') == 'Critical']
    high_weight_clauses = [c for c in clauses if c.get('weight') == 'High']

    # Check Critical-weight clauses
    for clause in critical_weight_clauses:
        if clause.get('risk_level') == 'CRITICAL':
            return 'CRITICAL'

    for clause in critical_weight_clauses:
        if clause.get('risk_level') == 'HIGH':
            return 'HIGH'

    # Check High-weight clauses
    for clause in high_weight_clauses:
        if clause.get('risk_level') == 'CRITICAL':
            return 'HIGH'

    high_risk_high_weight = [c for c in high_weight_clauses if c.get('risk_level') == 'HIGH']
    if len(high_risk_high_weight) >= 2:
        return 'HIGH'

    # Check any HIGH risk
    if any(c.get('risk_level') == 'HIGH' for c in clauses):
        return 'MODERATE'

    return 'LOW'


def add_title_page(doc: Document, data: Dict):
    """Add title page to document."""
    # Contract name
    p = doc.add_paragraph()
    run = p.add_run(data['contract_name'])
    run.font.size = Pt(24)
    run.bold = True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.space_after = Pt(24)

    # Report type
    p = doc.add_paragraph()
    run = p.add_run('CONTRACT RISK REVIEW')
    run.font.size = Pt(16)
    run.bold = True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.space_after = Pt(48)

    # Date
    p = doc.add_paragraph()
    run = p.add_run(f"Date: {datetime.now().strftime('%B %d, %Y')}")
    run.font.size = Pt(12)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.space_after = Pt(24)

    # Entity information
    p = doc.add_paragraph()
    run = p.add_run(f"Our Entity: {data['our_entity']}\n")
    run.font.size = Pt(12)
    run = p.add_run(f"Counterparty: {data['counterparty']}\n")
    run.font.size = Pt(12)
    run = p.add_run(f"Position: {data['position']}")
    run.font.size = Pt(12)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_page_break()


def add_executive_summary(doc: Document, data: Dict):
    """Add executive summary section."""
    # Section header
    p = doc.add_paragraph()
    run = p.add_run('EXECUTIVE SUMMARY')
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = COLORS['MODERATE']
    p.space_after = Pt(12)

    # Overall risk assessment
    overall_risk = data.get('overall_risk', calculate_overall_risk(data['clauses']))
    p = doc.add_paragraph()
    run = p.add_run('OVERALL RISK ASSESSMENT: ')
    run.font.size = Pt(12)
    run.bold = True
    run = p.add_run(overall_risk)
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = COLORS[overall_risk]
    p.space_after = Pt(12)

    # Top concerns
    p = doc.add_paragraph()
    run = p.add_run('TOP CONCERNS')
    run.font.size = Pt(11)
    run.bold = True
    p.space_after = Pt(6)

    for i, concern in enumerate(data.get('top_concerns', [])[:3], 1):
        p = doc.add_paragraph(f"{i}. {concern['clause_type']}: {concern['concern']}", style='List Number')
        p.paragraph_format.left_indent = Inches(0.25)

    # Risk distribution
    p = doc.add_paragraph()
    run = p.add_run('RISK DISTRIBUTION')
    run.font.size = Pt(11)
    run.bold = True
    p.space_after = Pt(6)

    distribution = data.get('distribution', {'CRITICAL': 0, 'HIGH': 0, 'MODERATE': 0, 'LOW': 0})
    p = doc.add_paragraph()
    for risk_level in ['CRITICAL', 'HIGH', 'MODERATE', 'LOW']:
        count = distribution.get(risk_level, 0)
        add_risk_indicator(p, risk_level)
        run = p.add_run(f'{risk_level}: {count}  ')
        run.font.size = Pt(11)

    doc.add_paragraph()


def add_risk_heat_map(doc: Document, data: Dict):
    """Add risk heat map table."""
    p = doc.add_paragraph()
    run = p.add_run('RISK HEAT MAP')
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = COLORS['MODERATE']
    p.space_after = Pt(12)

    # Create table
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'

    # Header row
    header_cells = table.rows[0].cells
    headers = ['Clause Type', 'CRITICAL', 'HIGH', 'MODERATE', 'LOW']

    for i, header in enumerate(headers):
        cell = header_cells[i]
        set_cell_shading(cell, COLORS['HEADER_BG'])
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.color.rgb = COLORS['HEADER_TEXT']
        run.bold = True
        run.font.size = Pt(10)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Data rows
    for clause in data['clauses']:
        row = table.add_row()
        cells = row.cells

        # Clause type
        cells[0].text = clause['type']

        # Add indicator in appropriate risk column
        risk_level = clause.get('risk_level', 'LOW')
        risk_col_idx = {'CRITICAL': 1, 'HIGH': 2, 'MODERATE': 3, 'LOW': 4}[risk_level]

        p = cells[risk_col_idx].paragraphs[0]
        add_risk_indicator(p, risk_level)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph()


def add_clause_analysis_table(doc: Document, data: Dict):
    """Add clause analysis table grouped by clause type."""
    p = doc.add_paragraph()
    run = p.add_run('CLAUSE ANALYSIS')
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = COLORS['MODERATE']
    p.space_after = Pt(12)

    # Group clauses by type
    grouped = {}
    for clause in data['clauses']:
        clause_type = clause['type']
        if clause_type not in grouped:
            grouped[clause_type] = []
        grouped[clause_type].append(clause)

    # For each clause type
    for clause_type, clauses_of_type in grouped.items():
        # Clause type header
        p = doc.add_paragraph()
        p.add_run(clause_type.upper())
        p.runs[0].font.size = Pt(12)
        p.runs[0].bold = True
        p.space_after = Pt(6)

        # Create table for this clause type
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'

        # Header row
        header_cells = table.rows[0].cells
        headers = ['Section', 'Risk', 'Concern', 'Recommendation']

        for i, header in enumerate(headers):
            cell = header_cells[i]
            set_cell_shading(cell, COLORS['HEADER_BG'])
            p = cell.paragraphs[0]
            run = p.add_run(header)
            run.font.color.rgb = COLORS['HEADER_TEXT']
            run.bold = True
            run.font.size = Pt(10)

        # Data rows
        for clause in clauses_of_type:
            row = table.add_row()
            cells = row.cells

            cells[0].text = clause.get('section', '')

            # Risk indicator
            risk_level = clause.get('risk_level', 'LOW')
            p = cells[1].paragraphs[0]
            add_risk_indicator(p, risk_level)

            cells[2].text = clause.get('concern', '')
            cells[3].text = clause.get('recommendation', '')

        doc.add_paragraph()


def add_negotiation_playbook(doc: Document, data: Dict):
    """Add negotiation playbook section."""
    p = doc.add_paragraph()
    run = p.add_run('NEGOTIATION PLAYBOOK')
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = COLORS['MODERATE']
    p.space_after = Pt(12)

    playbook = data.get('negotiation_playbook', {})

    # Your leverage
    p = doc.add_paragraph()
    run = p.add_run('YOUR LEVERAGE')
    run.font.size = Pt(11)
    run.bold = True

    for item in playbook.get('your_leverage', []):
        doc.add_paragraph(item, style='List Bullet')

    # Counterparty leverage
    p = doc.add_paragraph()
    run = p.add_run('COUNTERPARTY LEVERAGE')
    run.font.size = Pt(11)
    run.bold = True

    for item in playbook.get('counterparty_leverage', []):
        doc.add_paragraph(item, style='List Bullet')

    # Recommended sequence
    p = doc.add_paragraph()
    run = p.add_run('RECOMMENDED SEQUENCE')
    run.font.size = Pt(11)
    run.bold = True

    for i, item in enumerate(playbook.get('sequence', []), 1):
        p = doc.add_paragraph(f"{i}. ", style='List Number')
        add_risk_indicator(p, item['risk_level'])
        run = p.add_run(item['description'])
        run.font.size = Pt(11)

    # Trade-offs
    p = doc.add_paragraph()
    run = p.add_run('POTENTIAL TRADE-OFFS')
    run.font.size = Pt(11)
    run.bold = True

    for tradeoff in playbook.get('tradeoffs', []):
        p = doc.add_paragraph()
        run = p.add_run('Give ')
        run.font.size = Pt(11)
        add_risk_indicator(p, tradeoff['give_risk'])
        run = p.add_run(f"{tradeoff['give']} → Get ")
        run.font.size = Pt(11)
        add_risk_indicator(p, tradeoff['get_risk'])
        run = p.add_run(tradeoff['get'])
        run.font.size = Pt(11)

    doc.add_paragraph()


def add_disclaimers(doc: Document):
    """Add disclaimer section."""
    doc.add_page_break()

    p = doc.add_paragraph()
    run = p.add_run('DISCLAIMER')
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = COLORS['MODERATE']
    p.space_after = Pt(12)

    disclaimers = [
        ("Risk Advisory", "This report applies generally accepted contract risk categories and contract management best practices. Stakeholders must use this analysis to make informed business decisions regarding risk acceptance and mitigation. Failure to address identified risks may result in material financial, operational, or legal exposure."),
        ("Legal", "This analysis is for informational purposes only and does not constitute legal advice. Consult qualified legal counsel before making decisions based on this report."),
        ("AI-Generated", "This report was generated with AI assistance. All findings should be verified against source documents."),
        ("Confidential", "This document contains confidential analysis. Do not distribute without authorization.")
    ]

    for title, text in disclaimers:
        p = doc.add_paragraph()
        run = p.add_run(f"{title}: ")
        run.bold = True
        run.font.size = Pt(11)
        run = p.add_run(text)
        run.font.size = Pt(11)
        p.space_after = Pt(6)


def generate_risk_review_report(data: Dict, output_path: str = None) -> str:
    """
    Generate Contract Risk Review report as .docx file.

    Args:
        data: Report data dictionary with keys:
            - contract_name: str
            - our_entity: str
            - counterparty: str
            - position: str
            - clauses: List[Dict] (from extract_clauses)
            - overall_risk: str (optional, will calculate if not provided)
            - top_concerns: List[Dict]
            - distribution: Dict[str, int]
            - negotiation_playbook: Dict
        output_path: Path to save .docx file (optional)

    Returns:
        Path to generated .docx file
    """
    # Create document
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Add components
    add_title_page(doc, data)
    add_executive_summary(doc, data)
    add_risk_heat_map(doc, data)
    add_clause_analysis_table(doc, data)
    add_negotiation_playbook(doc, data)
    add_disclaimers(doc)

    # Save document
    if output_path is None:
        contract_name = data['contract_name'].replace(' ', '_')
        output_path = f"{contract_name}_Risk_Review_{datetime.now().strftime('%Y%m%d')}.docx"

    doc.save(output_path)
    return output_path


# Example usage
if __name__ == '__main__':
    # Sample report data
    sample_data = {
        'contract_name': 'Master Services Agreement',
        'our_entity': 'ACME Corporation',
        'counterparty': 'Vendor Inc',
        'position': 'Buyer',
        'clauses': [
            {
                'type': 'Indemnification',
                'section': '8.1',
                'weight': 'Critical',
                'risk_level': 'CRITICAL',
                'concern': 'Unlimited indemnification scope without cap',
                'recommendation': 'Add liability cap and scope limitations'
            },
            {
                'type': 'Limitation of Liability',
                'section': '9.2',
                'weight': 'Critical',
                'risk_level': 'HIGH',
                'concern': 'Cap is only 3x monthly fees',
                'recommendation': 'Negotiate for 12x monthly fees minimum'
            },
            {
                'type': 'Termination',
                'section': '11.1',
                'weight': 'High',
                'risk_level': 'MODERATE',
                'concern': 'Short 15-day cure period',
                'recommendation': 'Request 30-day cure period'
            }
        ],
        'top_concerns': [
            {'clause_type': 'Indemnification', 'concern': 'Unlimited scope without cap'},
            {'clause_type': 'Limitation of Liability', 'concern': 'Cap too low at 3x monthly fees'},
            {'clause_type': 'Termination', 'concern': 'Short cure period'}
        ],
        'distribution': {
            'CRITICAL': 1,
            'HIGH': 1,
            'MODERATE': 1,
            'LOW': 0
        },
        'negotiation_playbook': {
            'your_leverage': [
                'Long-term contract value',
                'Market competition for vendor'
            ],
            'counterparty_leverage': [
                'Specialized expertise',
                'Incumbent advantage'
            ],
            'sequence': [
                {'risk_level': 'CRITICAL', 'description': 'Indemnification cap first'},
                {'risk_level': 'HIGH', 'description': 'Liability limitations second'}
            ],
            'tradeoffs': [
                {
                    'give': 'Shorter payment terms',
                    'give_risk': 'MODERATE',
                    'get': 'Higher liability cap',
                    'get_risk': 'CRITICAL'
                }
            ]
        }
    }

    output_file = generate_risk_review_report(sample_data, 'test_risk_review.docx')
    print(f"Report generated: {output_file}")
