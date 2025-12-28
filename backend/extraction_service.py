"""
CIP Extraction Service v1.0
Pre-processes contract documents and extracts metadata via Claude Haiku.

Features:
- Smart text zone extraction (header, parties, signature)
- Token-optimized package (~3000 chars → ~750 tokens)
- Haiku API call with structured JSON response
- Graceful fallback on failure

Usage:
    from extraction_service import extract_contract_metadata
    result = extract_contract_metadata(file_path)
"""

import json
import re
import logging
from pathlib import Path
from typing import Dict, Optional, Tuple, List
from dataclasses import dataclass, asdict

# Document processing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not installed. Run: pip install python-docx")

# Claude API
try:
    from anthropic import Anthropic
    ANTHROPIC_AVAILABLE = True
except ImportError:
    ANTHROPIC_AVAILABLE = False
    logging.warning("anthropic not installed. Run: pip install anthropic")

from config import ANTHROPIC_API_KEY, UPLOAD_DIRECTORY

logger = logging.getLogger(__name__)

# ============================================================================
# CONFIGURATION
# ============================================================================

EXTRACTION_MODEL = "claude-3-haiku-20240307"
EXTRACTION_MAX_TOKENS = 500
EXTRACTION_TIMEOUT = 5  # seconds
EXTRACTION_TEMPERATURE = 0.1  # Low for consistency

# Zone sizes (chars)
HEADER_ZONE_SIZE = 2000
SIGNATURE_ZONE_SIZE = 1500
MAX_PACKAGE_SIZE = 4000  # Total chars to send to Claude

# Contract type keywords for pre-scan
TYPE_KEYWORDS = {
    "NDA": ["non-disclosure", "nda", "confidentiality agreement"],
    "MNDA": ["mutual non-disclosure", "mnda", "mutual nda"],
    "MOU": ["memorandum of understanding", "mou"],
    "MSA": ["master services agreement", "msa", "master agreement"],
    "PSA": ["professional services agreement", "psa"],
    "MPA": ["master purchase agreement", "mpa"],
    "SOW": ["statement of work", "sow", "work order"],
    "PO": ["purchase order", "po number"],
    "EULA": ["end user license", "eula", "license agreement"],
    "Amendment": ["amendment", "addendum", "modification"],
    "Proposal": ["proposal", "quotation", "quote"],
}

# ============================================================================
# DATA STRUCTURES
# ============================================================================

@dataclass
class ExtractionResult:
    """Result of metadata extraction."""
    success: bool
    title: Optional[str] = None
    title_confidence: float = 0.0
    parties: List[str] = None
    parties_confidence: float = 0.0
    contract_type: Optional[str] = None
    contract_type_confidence: float = 0.0
    purpose: Optional[str] = None
    purpose_confidence: float = 0.0
    orientation_hint: Optional[str] = None
    extraction_notes: Optional[str] = None
    error: Optional[str] = None
    raw_text_length: int = 0
    package_length: int = 0
    
    def __post_init__(self):
        if self.parties is None:
            self.parties = []
    
    def to_dict(self) -> Dict:
        return asdict(self)


# ============================================================================
# TEXT EXTRACTION
# ============================================================================

def extract_text_from_docx(file_path: Path) -> Tuple[str, Optional[str]]:
    """
    Extract full text from a .docx file.
    
    Returns:
        Tuple of (text, error_message)
    """
    if not DOCX_AVAILABLE:
        return "", "python-docx not installed"
    
    try:
        doc = Document(str(file_path))
        paragraphs = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        
        # Extract from tables (often contain party info)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    paragraphs.append(" | ".join(row_text))
        
        full_text = "\n".join(paragraphs)
        return full_text, None
        
    except Exception as e:
        logger.error(f"Failed to extract text from {file_path}: {e}")
        return "", str(e)


# ============================================================================
# PRE-PROCESSOR
# ============================================================================

def extract_party_zone(text: str) -> str:
    """
    Extract text sections likely to contain party names.
    Uses regex patterns common in contracts.
    """
    party_patterns = [
        # "between X and Y"
        r"(?:between|by and between)\s+([^,]+(?:,\s*(?:a|an)\s+[^,]+)?)\s+(?:and|AND)\s+([^,]+(?:,\s*(?:a|an)\s+[^,]+)?)",
        # "Party A" / "Party B" definitions
        r'["\']([^"\']+)["\']\s*\(["\'](?:Party [AB]|Client|Vendor|Provider|Customer|Contractor|Company)["\']',
        # "hereinafter referred to as"
        r"([A-Z][A-Za-z\s,\.]+(?:Inc|LLC|Corp|Ltd|Company|Co)\.?)\s*(?:\(|,)\s*(?:hereinafter|referred to as)",
    ]
    
    found_sections = []
    for pattern in party_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        for match in matches:
            if isinstance(match, tuple):
                found_sections.extend(match)
            else:
                found_sections.append(match)
    
    # Also grab context around matches
    party_zone = "\n".join(found_sections[:5])  # Limit to avoid bloat
    return party_zone[:800] if party_zone else ""


def prescan_contract_type(text: str) -> Optional[str]:
    """
    Quick keyword scan to suggest contract type.
    Returns best match or None.
    """
    text_lower = text.lower()
    
    for contract_type, keywords in TYPE_KEYWORDS.items():
        for keyword in keywords:
            if keyword in text_lower:
                return contract_type
    
    return None


def build_extraction_package(full_text: str) -> str:
    """
    Build token-optimized package for Claude.
    Combines key zones into ~3000 char package.
    
    Zones:
    - HEADER: First 2000 chars (title, preamble, date)
    - PARTIES: Regex-extracted party mentions
    - SIGNATURE: Last 1500 chars (signature blocks)
    """
    if not full_text:
        return ""
    
    text_length = len(full_text)
    
    # Extract zones
    header_zone = full_text[:HEADER_ZONE_SIZE].strip()
    signature_zone = full_text[-SIGNATURE_ZONE_SIZE:].strip() if text_length > SIGNATURE_ZONE_SIZE else ""
    party_zone = extract_party_zone(full_text)
    
    # Pre-scan for type hint
    type_hint = prescan_contract_type(full_text)
    
    # Build package
    package_parts = [
        "=== HEADER ZONE ===",
        header_zone,
    ]
    
    if party_zone:
        package_parts.extend([
            "\n=== PARTY ZONE ===",
            party_zone,
        ])
    
    if signature_zone and signature_zone != header_zone:
        package_parts.extend([
            "\n=== SIGNATURE ZONE ===",
            signature_zone,
        ])
    
    if type_hint:
        package_parts.extend([
            f"\n=== PRE-SCAN HINT ===",
            f"Keyword match suggests: {type_hint}",
        ])
    
    package = "\n".join(package_parts)
    
    # Truncate if still too long
    if len(package) > MAX_PACKAGE_SIZE:
        package = package[:MAX_PACKAGE_SIZE] + "\n... [truncated]"
    
    return package


# ============================================================================
# CLAUDE EXTRACTION
# ============================================================================

def load_extraction_prompt() -> str:
    """Load the extraction prompt template."""
    prompt_path = Path(__file__).parent / "prompts" / "extraction_prompt.md"
    
    try:
        with open(prompt_path, 'r', encoding='utf-8') as f:
            return f.read()
    except FileNotFoundError:
        logger.warning(f"Prompt file not found: {prompt_path}, using inline prompt")
        # Inline fallback prompt
        return """You are a contract metadata extractor. Extract from the contract text:
1. title - The formal agreement name
2. parties - Array of 2 company names
3. contract_type - One of: NDA, MNDA, MOU, MSA, PSA, MPA, SOW, PO, EULA, Proposal, Amendment, Other
4. purpose - One of: Professional Services, Consulting Services, Equipment Purchase, Construction, Staffing, Software License, Maintenance, Other
5. orientation_hint - Which party provides vs receives services

Return ONLY valid JSON:
{"title": "string", "title_confidence": 0.0-1.0, "parties": ["A", "B"], "parties_confidence": 0.0-1.0, "contract_type": "string", "contract_type_confidence": 0.0-1.0, "purpose": "string", "purpose_confidence": 0.0-1.0, "orientation_hint": "string or null", "extraction_notes": "string or null"}

CONTRACT TEXT:
"""


def call_claude_extraction(package: str) -> Tuple[Dict, Optional[str]]:
    """
    Call Claude Haiku to extract metadata.
    
    Returns:
        Tuple of (extraction_dict, error_message)
    """
    if not ANTHROPIC_AVAILABLE:
        return {}, "anthropic not installed"
    
    if not ANTHROPIC_API_KEY:
        return {}, "ANTHROPIC_API_KEY not configured"
    
    try:
        client = Anthropic(api_key=ANTHROPIC_API_KEY)
        
        prompt_template = load_extraction_prompt()
        full_prompt = prompt_template + "\n" + package
        
        response = client.messages.create(
            model=EXTRACTION_MODEL,
            max_tokens=EXTRACTION_MAX_TOKENS,
            temperature=EXTRACTION_TEMPERATURE,
            messages=[
                {"role": "user", "content": full_prompt}
            ]
        )
        
        # Extract text from response
        response_text = response.content[0].text.strip()
        
        # Parse JSON (handle potential markdown wrapping)
        json_text = response_text
        if "```json" in json_text:
            json_text = json_text.split("```json")[1].split("```")[0].strip()
        elif "```" in json_text:
            json_text = json_text.split("```")[1].split("```")[0].strip()
        
        extraction = json.loads(json_text)
        return extraction, None
        
    except json.JSONDecodeError as e:
        logger.error(f"Failed to parse Claude response as JSON: {e}")
        return {}, f"JSON parse error: {e}"
        
    except Exception as e:
        logger.error(f"Claude API call failed: {e}")
        return {}, str(e)


# ============================================================================
# MAIN EXTRACTION FUNCTION
# ============================================================================

def extract_contract_metadata(file_path: Path) -> ExtractionResult:
    """
    Main extraction function. Orchestrates the full pipeline.
    
    Args:
        file_path: Path to .docx file
        
    Returns:
        ExtractionResult with extracted metadata or error
    """
    logger.info(f"Starting extraction for: {file_path}")
    
    # Validate file exists
    if not file_path.exists():
        return ExtractionResult(
            success=False,
            error=f"File not found: {file_path}"
        )
    
    # Validate file type
    if file_path.suffix.lower() != '.docx':
        return ExtractionResult(
            success=False,
            error=f"Unsupported file type: {file_path.suffix}. Only .docx supported."
        )
    
    # Step 1: Extract text from docx
    full_text, error = extract_text_from_docx(file_path)
    if error:
        return ExtractionResult(
            success=False,
            error=f"Text extraction failed: {error}"
        )
    
    if not full_text or len(full_text) < 100:
        return ExtractionResult(
            success=False,
            error="Document appears to be empty or too short"
        )
    
    # Step 2: Build extraction package
    package = build_extraction_package(full_text)
    logger.info(f"Built extraction package: {len(full_text)} chars → {len(package)} chars")
    
    # Step 3: Call Claude
    extraction, error = call_claude_extraction(package)
    
    if error:
        # Partial failure - return what we can from pre-scan
        type_hint = prescan_contract_type(full_text)
        return ExtractionResult(
            success=False,
            contract_type=type_hint,
            contract_type_confidence=0.6 if type_hint else 0.0,
            error=error,
            raw_text_length=len(full_text),
            package_length=len(package),
            extraction_notes="Claude extraction failed. Pre-scan type hint provided."
        )
    
    # Step 4: Build result
    return ExtractionResult(
        success=True,
        title=extraction.get("title"),
        title_confidence=extraction.get("title_confidence", 0.0),
        parties=extraction.get("parties", []),
        parties_confidence=extraction.get("parties_confidence", 0.0),
        contract_type=extraction.get("contract_type"),
        contract_type_confidence=extraction.get("contract_type_confidence", 0.0),
        purpose=extraction.get("purpose"),
        purpose_confidence=extraction.get("purpose_confidence", 0.0),
        orientation_hint=extraction.get("orientation_hint"),
        extraction_notes=extraction.get("extraction_notes"),
        raw_text_length=len(full_text),
        package_length=len(package)
    )


# ============================================================================
# RELATED CONTRACTS LOOKUP
# ============================================================================

def find_related_contracts(parties: List[str], exclude_id: Optional[int] = None) -> List[Dict]:
    """
    Find contracts with matching party names.
    
    Args:
        parties: List of party names to search for
        exclude_id: Contract ID to exclude (if updating existing)
        
    Returns:
        List of matching contracts with basic metadata
    """
    if not parties:
        return []
    
    try:
        import sqlite3
        from config import CONTRACTS_DB
        
        conn = sqlite3.connect(str(CONTRACTS_DB))
        conn.row_factory = sqlite3.Row
        conn.execute("PRAGMA foreign_keys = ON")
        cursor = conn.cursor()
        
        # Build query for party matching
        party_conditions = []
        params = []
        
        for party in parties:
            if party:
                party_conditions.append("(our_entity LIKE ? OR counterparty LIKE ?)")
                params.extend([f"%{party}%", f"%{party}%"])
        
        if not party_conditions:
            return []
        
        query = f"""
            SELECT id, title, contract_type, our_entity, counterparty, created_at
            FROM contracts
            WHERE ({" OR ".join(party_conditions)})
        """
        
        if exclude_id:
            query += " AND id != ?"
            params.append(exclude_id)
        
        query += " ORDER BY created_at DESC LIMIT 10"
        
        cursor.execute(query, params)
        rows = cursor.fetchall()
        conn.close()
        
        return [dict(row) for row in rows]
        
    except Exception as e:
        logger.error(f"Related contracts lookup failed: {e}")
        return []


# ============================================================================
# DEPENDENCY CHECK
# ============================================================================

def check_dependencies() -> Dict[str, bool]:
    """Check if all required dependencies are available."""
    return {
        "python-docx": DOCX_AVAILABLE,
        "anthropic": ANTHROPIC_AVAILABLE,
        "api_key_configured": bool(ANTHROPIC_API_KEY),
    }


if __name__ == "__main__":
    # Test dependency check
    deps = check_dependencies()
    print("Dependency check:")
    for dep, available in deps.items():
        status = "✅" if available else "❌"
        print(f"  {status} {dep}")
