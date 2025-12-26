#!/usr/bin/env python3
"""
Document comparison script for contract version analysis v2.0.
Extracts, matches, and classifies changes between two contract versions.

v2.0 Changes:
- Extracts Word list numbering (Roman numerals, etc.) from numbering.xml
- Reconstructs rendered section numbers (I., II., IV.A, IV.B, etc.)
- Improved section detection for contracts without text-based numbers
"""

import sys
import json
import re
import argparse
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from datetime import datetime
from difflib import SequenceMatcher
from typing import Dict, List, Tuple, Optional

SKILL_VERSION = '2.0.0'

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("[FAIL] python-docx not installed. Install with: pip install python-docx")
    sys.exit(1)


def to_roman(num: int) -> str:
    """Convert integer to Roman numeral."""
    val = [(1000, 'M'), (900, 'CM'), (500, 'D'), (400, 'CD'),
           (100, 'C'), (90, 'XC'), (50, 'L'), (40, 'XL'),
           (10, 'X'), (9, 'IX'), (5, 'V'), (4, 'IV'), (1, 'I')]
    result = ''
    for (n, r) in val:
        while num >= n:
            result += r
            num -= n
    return result


def to_lower_roman(num: int) -> str:
    """Convert integer to lowercase Roman numeral."""
    return to_roman(num).lower()


def to_upper_letter(num: int) -> str:
    """Convert integer to uppercase letter (1=A, 2=B, etc.)."""
    return chr(64 + num) if 1 <= num <= 26 else str(num)


def to_lower_letter(num: int) -> str:
    """Convert integer to lowercase letter (1=a, 2=b, etc.)."""
    return chr(96 + num) if 1 <= num <= 26 else str(num)


class NumberingExtractor:
    """Extracts Word list numbering definitions from numbering.xml."""
    
    def __init__(self, docx_path: str):
        self.docx_path = docx_path
        self.abstract_nums = {}  # abstractNumId -> level definitions
        self.num_to_abstract = {}  # numId -> abstractNumId
        self._parse_numbering_xml()
    
    def _parse_numbering_xml(self):
        """Parse numbering.xml to extract numbering definitions."""
        try:
            with zipfile.ZipFile(self.docx_path, 'r') as z:
                if 'word/numbering.xml' not in z.namelist():
                    return
                
                numbering_xml = z.read('word/numbering.xml')
                root = ET.fromstring(numbering_xml)
                ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                
                # Parse abstractNum definitions
                for abstractNum in root.findall('.//w:abstractNum', ns):
                    abstractNumId = abstractNum.get(qn('w:abstractNumId'))
                    levels = {}
                    
                    for lvl in abstractNum.findall('.//w:lvl', ns):
                        ilvl = int(lvl.get(qn('w:ilvl')))
                        numFmt = lvl.find('.//w:numFmt', ns)
                        lvlText = lvl.find('.//w:lvlText', ns)
                        
                        levels[ilvl] = {
                            'format': numFmt.get(qn('w:val')) if numFmt is not None else 'decimal',
                            'text': lvlText.get(qn('w:val')) if lvlText is not None else '%1.'
                        }
                    
                    self.abstract_nums[abstractNumId] = levels
                
                # Parse num -> abstractNum mapping
                for num in root.findall('.//w:num', ns):
                    numId = num.get(qn('w:numId'))
                    abstractNumIdRef = num.find('.//w:abstractNumId', ns)
                    if abstractNumIdRef is not None:
                        self.num_to_abstract[numId] = abstractNumIdRef.get(qn('w:val'))
                        
        except Exception as e:
            print(f"[WARN] Could not parse numbering.xml: {e}")
    
    def get_format(self, numId: str, ilvl: int) -> Dict:
        """Get numbering format for a given numId and level."""
        abstractNumId = self.num_to_abstract.get(numId)
        if not abstractNumId:
            return {'format': 'decimal', 'text': '%1.'}
        
        levels = self.abstract_nums.get(abstractNumId, {})
        return levels.get(ilvl, {'format': 'decimal', 'text': '%1.'})


class SectionNumberRenderer:
    """Renders section numbers based on Word numbering definitions."""
    
    def __init__(self, numbering_extractor: NumberingExtractor):
        self.numbering = numbering_extractor
        self.counters = {}  # numId -> list of counters per level
    
    def render_number(self, numId: str, ilvl: int) -> str:
        """Render the section number for a paragraph with numbering."""
        # Initialize counter for this numId if needed
        if numId not in self.counters:
            self.counters[numId] = [0] * 10
        
        # Increment current level
        self.counters[numId][ilvl] += 1
        
        # Reset deeper levels
        for j in range(ilvl + 1, 10):
            self.counters[numId][j] = 0
        
        # Get format info
        fmt_info = self.numbering.get_format(numId, ilvl)
        fmt = fmt_info['format']
        
        # Render based on format
        val = self.counters[numId][ilvl]
        
        if fmt == 'upperRoman':
            rendered = to_roman(val)
        elif fmt == 'lowerRoman':
            rendered = to_lower_roman(val)
        elif fmt == 'upperLetter':
            rendered = to_upper_letter(val)
        elif fmt == 'lowerLetter':
            rendered = to_lower_letter(val)
        elif fmt == 'decimal':
            rendered = str(val)
        elif fmt == 'bullet':
            return ''  # Bullets don't have numbers
        else:
            rendered = str(val)
        
        # Handle multi-level numbers (e.g., IV.B)
        if ilvl > 0:
            parent_parts = []
            for level in range(ilvl):
                parent_fmt = self.numbering.get_format(numId, level)['format']
                parent_val = self.counters[numId][level]
                if parent_val > 0:
                    if parent_fmt == 'upperRoman':
                        parent_parts.append(to_roman(parent_val))
                    elif parent_fmt == 'decimal':
                        parent_parts.append(str(parent_val))
                    else:
                        parent_parts.append(str(parent_val))
            
            if parent_parts:
                return f"{'.'.join(parent_parts)}.{rendered}"
        
        return f"{rendered}."


class ContractSection:
    """Represents a section of a contract document."""

    def __init__(self, number: str, title: str, content: str, original_index: int):
        self.number = number
        self.title = title
        self.content = content
        self.original_index = original_index

    def __repr__(self):
        return f"Section {self.number}: {self.title[:50]}"


class DocumentExtractor:
    """Extracts structured content from .docx contract files with Word numbering support."""

    def __init__(self, docx_path: str):
        self.path = Path(docx_path)
        if not self.path.exists():
            raise FileNotFoundError(f"Document not found: {docx_path}")
        self.doc = Document(str(self.path))
        self.numbering = NumberingExtractor(str(self.path))
        self.renderer = SectionNumberRenderer(self.numbering)

    def extract_sections(self) -> List[ContractSection]:
        """Extract sections from document with Word numbering support."""
        sections = []
        current_section = None
        section_index = 0

        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Check for Word numbering
            numPr = para._p.find(qn('w:pPr'))
            rendered_number = None
            
            if numPr is not None:
                numPr_el = numPr.find(qn('w:numPr'))
                if numPr_el is not None:
                    ilvl_el = numPr_el.find(qn('w:ilvl'))
                    numId_el = numPr_el.find(qn('w:numId'))
                    
                    if ilvl_el is not None and numId_el is not None:
                        ilvl = int(ilvl_el.get(qn('w:val')))
                        numId = numId_el.get(qn('w:val'))
                        rendered_number = self.renderer.render_number(numId, ilvl)

            # Detect section headers
            is_header = False
            
            # Check if this is a numbered header (short text with numbering)
            if rendered_number and len(text) < 100:
                # This is likely a section header
                is_header = True
                
                # Save previous section
                if current_section:
                    sections.append(current_section)
                
                # Start new section
                current_section = ContractSection(
                    rendered_number, 
                    text, 
                    "", 
                    section_index
                )
                section_index += 1
                continue
            
            # Fallback: Check for text-based section numbers
            section_match = re.match(r'^(\d+\.?\d*)\s+(.+)$', text)
            article_match = re.match(r'^(Article|Section|Clause)\s+(\d+\.?\d*)[:\-\s]+(.+)$', text, re.IGNORECASE)

            if section_match and len(text) < 200:
                if current_section:
                    sections.append(current_section)
                number = section_match.group(1)
                title = section_match.group(2)
                current_section = ContractSection(number, title, "", section_index)
                section_index += 1

            elif article_match and len(text) < 200:
                if current_section:
                    sections.append(current_section)
                number = article_match.group(2)
                title = article_match.group(3)
                current_section = ContractSection(number, title, "", section_index)
                section_index += 1

            else:
                # Add to current section content
                if current_section:
                    current_section.content += text + "\n"
                else:
                    # Create initial section for preamble/intro
                    current_section = ContractSection("0", "Preamble", text + "\n", section_index)
                    section_index += 1

        # Add final section
        if current_section:
            sections.append(current_section)

        return sections

    def get_full_text(self) -> str:
        """Extract all text from document."""
        return "\n".join([para.text for para in self.doc.paragraphs])


class SectionMatcher:
    """Matches sections between two document versions using content-based matching."""

    def __init__(self, confidence_threshold: float = 0.85):
        self.confidence_threshold = confidence_threshold

    def match_sections(self, sections_v1: List[ContractSection],
                      sections_v2: List[ContractSection]) -> List[Dict]:
        """Match sections from V1 to V2 using content similarity."""
        matches = []
        used_v2_indices = set()

        for v1_section in sections_v1:
            best_match = None
            best_score = 0.0
            best_v2_index = -1

            for i, v2_section in enumerate(sections_v2):
                if i in used_v2_indices:
                    continue

                score = self._calculate_similarity(v1_section, v2_section)

                if score > best_score:
                    best_score = score
                    best_match = v2_section
                    best_v2_index = i

            if best_score >= self.confidence_threshold:
                used_v2_indices.add(best_v2_index)
                matches.append({
                    'v1_section': v1_section,
                    'v2_section': best_match,
                    'confidence': best_score,
                    'tie_breaker': False
                })
            else:
                if best_match:
                    used_v2_indices.add(best_v2_index)
                matches.append({
                    'v1_section': v1_section,
                    'v2_section': best_match,
                    'confidence': best_score,
                    'tie_breaker': True
                })

        # Handle unmatched V2 sections (new sections)
        for i, v2_section in enumerate(sections_v2):
            if i not in used_v2_indices:
                matches.append({
                    'v1_section': None,
                    'v2_section': v2_section,
                    'confidence': 0.0,
                    'tie_breaker': False,
                    'new_section': True
                })

        return matches

    def _calculate_similarity(self, section1: ContractSection,
                            section2: ContractSection) -> float:
        """Calculate similarity score between two sections."""
        title_sim = SequenceMatcher(None,
                                   section1.title.lower(),
                                   section2.title.lower()).ratio()
        content_sim = SequenceMatcher(None,
                                     section1.content[:500].lower(),
                                     section2.content[:500].lower()).ratio()
        return (title_sim * 0.4) + (content_sim * 0.6)


class ChangeDetector:
    """Detects substantive changes between matched sections."""

    IGNORE_PATTERNS = [
        r'\b\d{4}\b',
        r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b',
    ]

    def __init__(self):
        self.anonymization_tokens = [
            r'\[COMPANY[_\s]?[A-Z]?\]',
            r'\[CLIENT\]',
            r'\[VENDOR\]',
            r'\[CUSTOMER\]',
        ]

    def detect_changes(self, v1_content: str, v2_content: str) -> Dict:
        """Detect substantive differences between two section contents."""
        v1_normalized = self._normalize_text(v1_content)
        v2_normalized = self._normalize_text(v2_content)

        similarity = SequenceMatcher(None, v1_normalized, v2_normalized).ratio()

        if similarity > 0.95:
            return {
                'has_changes': False,
                'similarity': similarity,
                'change_type': 'none'
            }

        change_type = self._classify_change_type(v1_content, v2_content)

        return {
            'has_changes': True,
            'similarity': similarity,
            'change_type': change_type,
            'v1_text': v1_content,
            'v2_text': v2_content
        }

    def _normalize_text(self, text: str) -> str:
        """Normalize text for comparison."""
        text = text.lower()
        text = re.sub(r'\s+', ' ', text)
        for pattern in self.IGNORE_PATTERNS:
            text = re.sub(pattern, '', text)
        return text.strip()

    def _classify_change_type(self, v1: str, v2: str) -> str:
        """Classify the type of change."""
        if not v1.strip():
            return 'addition'
        if not v2.strip():
            return 'deletion'
        
        v1_words = set(v1.lower().split())
        v2_words = set(v2.lower().split())
        
        common = v1_words & v2_words
        total = v1_words | v2_words
        
        if len(common) / len(total) < 0.2:
            return 'wholesale_replacement'
        
        return 'modification'


def load_classification_rules() -> Dict:
    """Load impact classification rules."""
    return {
        'CRITICAL': [
            'limitation of liability', 'indemnif', 'intellectual property',
            'ip ownership', 'insurance', 'compliance', 'regulatory'
        ],
        'HIGH_PRIORITY': [
            'terminat', 'warrant', 'acceptance', 'sla', 'service level',
            'fee', 'price', 'pricing'
        ],
        'IMPORTANT': [
            'payment', 'invoice', 'milestone'
        ],
        'OPERATIONAL': [
            'notice', 'approval', 'report'
        ]
    }


def classify_impact(section_title: str, section_content: str, rules: Dict) -> str:
    """Classify the impact level of a section."""
    combined = (section_title + ' ' + section_content).lower()
    
    for impact, keywords in rules.items():
        for keyword in keywords:
            if keyword in combined:
                return impact
    
    return 'ADMINISTRATIVE'


def add_metadata(comparison_data: Dict) -> Dict:
    """Add metadata fields to comparison output for QA/QC workflow."""
    comparison_data['generated_date'] = datetime.now(tz=None).astimezone().isoformat()
    comparison_data['skill_version'] = SKILL_VERSION
    comparison_data['metadata'] = {
        'validated': False,
        'qa_qc_complete': False,
        'review_date': None,
        'reviewed_by': None,
        'flagged_count': 0,
        'approved_count': 0,
        'modified_count': 0,
        'rejected_count': 0
    }
    return comparison_data


def load_expected_changes(filepath: str) -> Dict:
    """Load expected changes from risk report JSON."""
    if not filepath:
        return {}
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return json.load(f)
    except (FileNotFoundError, json.JSONDecodeError) as e:
        print(f"[WARN] Could not load expected changes: {e}")
        return {}


def compare_documents(v1_path: str, v2_path: str, output_path: str = None, expected_changes_path: str = None) -> Dict:
    """Main comparison function."""
    print(f"[OK] Extracting V1: {v1_path}")
    extractor_v1 = DocumentExtractor(v1_path)
    sections_v1 = extractor_v1.extract_sections()
    print(f"  Found {len(sections_v1)} sections in V1")
    
    # Debug: show section numbers
    for s in sections_v1[:5]:
        print(f"    {s.number} - {s.title[:40]}")

    print(f"[OK] Extracting V2: {v2_path}")
    extractor_v2 = DocumentExtractor(v2_path)
    sections_v2 = extractor_v2.extract_sections()
    print(f"  Found {len(sections_v2)} sections in V2")

    print("[OK] Matching sections...")
    matcher = SectionMatcher()
    matches = matcher.match_sections(sections_v1, sections_v2)
    print(f"  Matched {len(matches)} section pairs")

    print("[OK] Detecting changes...")
    detector = ChangeDetector()
    rules = load_classification_rules()

    results = {
        'v1_path': str(v1_path),
        'v2_path': str(v2_path),
        'sections_analyzed': len(matches),
        'changes': []
    }

    changes_count = 0
    for match in matches:
        v1_sec = match.get('v1_section')
        v2_sec = match.get('v2_section')

        if not v1_sec or not v2_sec:
            changes_count += 1
            results['changes'].append({
                'section_number': v2_sec.number if v2_sec else v1_sec.number,
                'section_title': v2_sec.title if v2_sec else v1_sec.title,
                'v1_content': v1_sec.content if v1_sec else "",
                'v2_content': v2_sec.content if v2_sec else "",
                'impact': 'HIGH_PRIORITY' if not v1_sec else 'OPERATIONAL',
                'tie_breaker': False,
                'new_section': match.get('new_section', False)
            })
        else:
            change_info = detector.detect_changes(v1_sec.content, v2_sec.content)

            if change_info['has_changes']:
                changes_count += 1
                impact = classify_impact(v2_sec.title, v2_sec.content, rules)

                results['changes'].append({
                    'section_number': v2_sec.number,
                    'section_title': v2_sec.title,
                    'v1_content': v1_sec.content,
                    'v2_content': v2_sec.content,
                    'impact': impact,
                    'tie_breaker': match.get('tie_breaker', False),
                    'confidence': match.get('confidence', 1.0),
                    'change_type': change_info['change_type']
                })

    results['total_changes'] = changes_count
    print(f"[OK] Detected {changes_count} substantive changes")

    # Load expected changes and assess alignment
    expected_changes = load_expected_changes(expected_changes_path)
    if expected_changes:
        print("[OK] Assessing alignment with expected changes...")
        detected_sections = set()
        for change in results['changes']:
            section = change['section_number']
            detected_sections.add(section)

    # Add metadata for QA/QC workflow
    results = add_metadata(results)

    # Save results to JSON
    if output_path:
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(results, f, indent=2)
        print(f"[OK] Results saved to: {output_path}")

    return results


def main():
    """Command-line interface."""
    parser = argparse.ArgumentParser(description='Compare two contract document versions (v2.0)')
    parser.add_argument('v1_doc', help='Path to V1 contract document (.docx)')
    parser.add_argument('v2_doc', help='Path to V2 contract document (.docx)')
    parser.add_argument('output', nargs='?', default='comparison_results.json',
                        help='Output JSON path (default: comparison_results.json)')
    parser.add_argument('--expected-changes', type=str,
                        help='Path to JSON file with expected changes from risk report')

    args = parser.parse_args()

    try:
        results = compare_documents(args.v1_doc, args.v2_doc, args.output, args.expected_changes)

        print("\n" + "=" * 60)
        print("COMPARISON SUMMARY")
        print("=" * 60)
        print(f"V1: {results['v1_path']}")
        print(f"V2: {results['v2_path']}")
        print(f"Sections analyzed: {results['sections_analyzed']}")
        print(f"Total changes detected: {results['total_changes']}")
        print("\nImpact breakdown:")

        impact_counts = {}
        for change in results['changes']:
            impact = change['impact']
            impact_counts[impact] = impact_counts.get(impact, 0) + 1

        for impact, count in sorted(impact_counts.items()):
            print(f"  {impact}: {count}")

        print("=" * 60)
        return 0

    except Exception as e:
        print(f"[FAIL] Error during comparison: {e}")
        import traceback
        traceback.print_exc()
        return 1


if __name__ == "__main__":
    sys.exit(main())
