#!/usr/bin/env python3
"""
Redline Exporter
Generates professional Word documents with visual redlines
"""

from docx import Document
from docx.shared import RGBColor
from difflib import SequenceMatcher
from typing import List, Dict
from datetime import datetime


class RedlineExporter:
    """Export redline reviews to professional Word documents"""

    def __init__(self):
        pass

    def export_to_docx(self, clauses: List[Dict], decisions: Dict, modifications: Dict, contract_info: Dict, output_path: str):
        """
        Generate Word document with visual redlines

        Args:
            clauses: List of clause dicts from redline analysis
            decisions: Dict mapping clause_index -> 'approved'|'rejected'|'modified'
            modifications: Dict mapping clause_index -> modified_text
            contract_info: {contract_id, filename, context}
            output_path: Path to save .docx file
        """
        doc = Document()

        # Add title
        title = doc.add_heading('Contract Redline Review', level=0)

        # Add metadata
        doc.add_paragraph(f"Contract: {contract_info.get('filename', 'Unknown')}")
        doc.add_paragraph(f"Contract ID: {contract_info.get('contract_id', 'N/A')}")
        doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

        context = contract_info.get('context', {})
        doc.add_paragraph(f"Position: {context.get('position', 'N/A')}")
        doc.add_paragraph(f"Leverage: {context.get('leverage', 'N/A')}")
        doc.add_paragraph(f"Contract Type: {context.get('contract_type', 'N/A')}")

        doc.add_page_break()

        # Add legend
        legend = doc.add_heading('Legend', level=1)
        doc.add_paragraph("🟢 = Approved revision")
        doc.add_paragraph("✏️ = Modified revision")
        doc.add_paragraph("❌ = Rejected (keeping original)")

        p = doc.add_paragraph()
        run = p.add_run("Deleted text")
        run.font.color.rgb = RGBColor(255, 0, 0)
        run.font.strike = True

        p = doc.add_paragraph()
        run = p.add_run("Inserted text")
        run.font.color.rgb = RGBColor(0, 128, 0)
        run.font.bold = True

        doc.add_page_break()

        # Process each clause
        approved_count = 0
        rejected_count = 0
        modified_count = 0

        for idx, clause in enumerate(clauses):
            decision = decisions.get(idx, 'pending')

            # Only include approved and modified clauses
            if decision not in ['approved', 'modified']:
                continue

            if decision == 'approved':
                approved_count += 1
            elif decision == 'modified':
                modified_count += 1

            # Add clause header
            section_number = clause.get('section_number', 'N/A')
            section_title = clause.get('section_title', 'Unknown')

            heading_text = f"{section_number}: {section_title}"
            if decision == 'approved':
                heading_text = "🟢 " + heading_text
            elif decision == 'modified':
                heading_text = "✏️ " + heading_text

            doc.add_heading(heading_text, level=2)

            # Add risk level and pattern
            info_p = doc.add_paragraph()
            info_p.add_run(f"Risk Level: {clause.get('risk_level', 'N/A')}  |  ")
            info_p.add_run(f"Pattern: {clause.get('pattern_applied', 'N/A')}")

            # Add change metrics
            metrics = clause.get('change_metrics', {})
            metrics_p = doc.add_paragraph()
            metrics_p.add_run(f"Change: {metrics.get('change_ratio', 0):.1%}  |  ")
            metrics_p.add_run(f"Retention: {metrics.get('word_retention', 0):.1%}  |  ")
            metrics_p.add_run(f"Minimal: {'✅' if metrics.get('is_minimal', False) else '❌'}")

            # Add visual redline
            doc.add_heading('Redline', level=3)

            original_text = clause.get('clause_text', '')

            # Determine which revision to use
            if decision == 'modified' and idx in modifications:
                revised_text = modifications[idx]
            else:
                revised_text = clause.get('suggested_revision', '')

            # Generate visual redline in Word format
            self._add_visual_redline(doc, original_text, revised_text)

            doc.add_paragraph()  # Spacing

        # Add summary
        doc.add_page_break()
        doc.add_heading('Summary', level=1)
        doc.add_paragraph(f"Total Clauses: {len(clauses)}")
        doc.add_paragraph(f"Approved Revisions: {approved_count}")
        doc.add_paragraph(f"Modified Revisions: {modified_count}")
        doc.add_paragraph(f"Rejected: {rejected_count}")

        # Save document
        doc.save(output_path)

        return {
            'file_path': output_path,
            'clauses_exported': approved_count + modified_count,
            'approved': approved_count,
            'modified': modified_count
        }

    def _add_visual_redline(self, doc, original: str, revised: str):
        """
        Add visual redline to Word document

        Deletions: Red, strikethrough
        Insertions: Green, bold
        """
        matcher = SequenceMatcher(None, original, revised)

        paragraph = doc.add_paragraph()

        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == 'equal':
                # Unchanged text
                paragraph.add_run(original[i1:i2])
            elif tag == 'delete':
                # Deleted text (red strikethrough)
                deleted_run = paragraph.add_run(original[i1:i2])
                deleted_run.font.color.rgb = RGBColor(255, 0, 0)
                deleted_run.font.strike = True
            elif tag == 'insert':
                # Inserted text (green bold)
                inserted_run = paragraph.add_run(revised[j1:j2])
                inserted_run.font.color.rgb = RGBColor(0, 128, 0)
                inserted_run.font.bold = True
            elif tag == 'replace':
                # Replaced text (deletion + insertion)
                deleted_run = paragraph.add_run(original[i1:i2])
                deleted_run.font.color.rgb = RGBColor(255, 0, 0)
                deleted_run.font.strike = True

                inserted_run = paragraph.add_run(revised[j1:j2])
                inserted_run.font.color.rgb = RGBColor(0, 128, 0)
                inserted_run.font.bold = True


if __name__ == "__main__":
    # Test exporter
    exporter = RedlineExporter()

    # Sample data
    clauses = [
        {
            'section_number': '1.1',
            'section_title': 'Termination',
            'clause_text': 'Company may terminate this Agreement at any time with 15 days notice.',
            'risk_level': 'HIGH',
            'suggested_revision': 'Either party may terminate this Agreement with 30 days written notice.',
            'change_metrics': {
                'change_ratio': 0.25,
                'word_retention': 0.78,
                'is_minimal': True
            },
            'pattern_applied': 'Mutual Termination Rights'
        }
    ]

    decisions = {0: 'approved'}
    modifications: Dict[int, str] = {}

    contract_info = {
        'contract_id': 1,
        'filename': 'test_contract.pdf',
        'context': {
            'position': 'Vendor',
            'leverage': 'Moderate',
            'contract_type': 'Services Agreement'
        }
    }

    result = exporter.export_to_docx(
        clauses, decisions, modifications, contract_info,
        'C:/Users/jrudy/CIP/backend/test_redline_export.docx'
    )

    print(f"\n{'='*60}")
    print("REDLINE EXPORT TEST")
    print(f"{'='*60}")
    print(f"File: {result['file_path']}")
    print(f"Clauses Exported: {result['clauses_exported']}")
    print(f"Approved: {result['approved']}")
    print(f"Modified: {result['modified']}")
