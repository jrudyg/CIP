"""
CIP - Check for Risks Page (API Integrated)
Real-time automated contract risk analysis with backend integration
v1.5: Zone layout retrofit
"""

import streamlit as st
import pandas as pd
import json
import sys
import os
from io import BytesIO
from datetime import datetime

# DOCX generation
try:
    from docx import Document
    from docx.shared import RGBColor, Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Add parent directory to path
sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))

# Import UI components
from ui_components import (
    page_header, section_header,
    toast_success, toast_info, apply_spacing
)
from theme_system import status_badge, apply_theme, inject_cip_logo
from shared_components import contract_selector, api_call_with_spinner
from components.contract_context import (
    init_contract_context,
    get_active_contract,
    get_active_contract_data,
    render_active_contract_header,
    render_recent_contracts_widget
)
from zone_layout import zone_layout, check_system_health, system_status
from cip_components import filter_bar, smart_list


def generate_analysis_docx(analysis: dict, contract_id: int) -> BytesIO:
    """Generate a DOCX report from analysis results."""
    doc = Document()

    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Colors
    COLOR_RED = RGBColor(192, 0, 0)
    COLOR_ORANGE = RGBColor(255, 153, 0)
    COLOR_GREEN = RGBColor(0, 128, 0)
    COLOR_BLUE = RGBColor(31, 73, 125)

    # Title
    title = doc.add_heading('Contract Risk Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Contract ID: {contract_id}").bold = True
    p.add_run(f"\nGenerated: {datetime.now().strftime('%B %d, %Y at %H:%M')}")

    doc.add_paragraph()

    # Executive Summary
    doc.add_heading('Executive Summary', level=1)

    overall_risk = analysis.get('overall_risk', 'UNKNOWN')
    confidence = analysis.get('confidence_score', 0)

    p = doc.add_paragraph()
    p.add_run("Overall Risk Level: ").bold = True
    risk_run = p.add_run(overall_risk)
    risk_run.bold = True
    if overall_risk in ['HIGH', 'CRITICAL']:
        risk_run.font.color.rgb = COLOR_RED
    elif overall_risk == 'MEDIUM':
        risk_run.font.color.rgb = COLOR_ORANGE
    else:
        risk_run.font.color.rgb = COLOR_GREEN

    p = doc.add_paragraph()
    p.add_run("Confidence Score: ").bold = True
    p.add_run(f"{confidence:.0%}")

    # Summary counts
    dealbreakers = analysis.get('dealbreakers', [])
    critical = analysis.get('critical_items', [])
    important = analysis.get('important_items', [])
    standard = analysis.get('standard_items', [])

    p = doc.add_paragraph()
    p.add_run("\nFindings Summary:\n").bold = True
    p.add_run(f"• Dealbreakers: {len(dealbreakers)}\n")
    p.add_run(f"• Critical Items: {len(critical)}\n")
    p.add_run(f"• Important Items: {len(important)}\n")
    p.add_run(f"• Standard Items: {len(standard)}")

    # Dealbreakers Section
    if dealbreakers:
        doc.add_heading('DEALBREAKERS', level=1)
        p = doc.add_paragraph()
        p.add_run("These items present unacceptable risk and should block the deal.").italic = True

        for item in dealbreakers:
            doc.add_heading(f"{item.get('section_title', 'Unknown')} (Section {item.get('section_number', 'N/A')})", level=2)
            p = doc.add_paragraph()
            p.add_run("Category: ").bold = True
            p.add_run(f"{item.get('category', 'N/A')}\n")
            p.add_run("Finding: ").bold = True
            p.add_run(f"{item.get('finding', 'No details')}\n")
            p.add_run("Recommendation: ").bold = True
            p.add_run(f"{item.get('recommendation', 'No recommendation')}")

    # Critical Items Section
    if critical:
        doc.add_heading('CRITICAL ITEMS', level=1)
        p = doc.add_paragraph()
        p.add_run("High-priority issues requiring immediate attention.").italic = True

        for item in critical:
            doc.add_heading(f"{item.get('section_title', 'Unknown')} (Section {item.get('section_number', 'N/A')})", level=2)
            p = doc.add_paragraph()
            p.add_run("Category: ").bold = True
            p.add_run(f"{item.get('category', 'N/A')}\n")
            p.add_run("Finding: ").bold = True
            p.add_run(f"{item.get('finding', 'No details')}\n")
            p.add_run("Recommendation: ").bold = True
            p.add_run(f"{item.get('recommendation', 'No recommendation')}")

    # Important Items Section
    if important:
        doc.add_heading('IMPORTANT ITEMS', level=1)

        for item in important:
            doc.add_heading(f"{item.get('section_title', 'Unknown')} (Section {item.get('section_number', 'N/A')})", level=2)
            p = doc.add_paragraph()
            p.add_run("Category: ").bold = True
            p.add_run(f"{item.get('category', 'N/A')}\n")
            p.add_run("Finding: ").bold = True
            p.add_run(f"{item.get('finding', 'No details')}\n")
            p.add_run("Recommendation: ").bold = True
            p.add_run(f"{item.get('recommendation', 'No recommendation')}")

    # Standard Items Section
    if standard:
        doc.add_heading('STANDARD ITEMS', level=1)

        for item in standard:
            p = doc.add_paragraph()
            p.add_run(f"• {item.get('section_title', 'Unknown')}: ").bold = True
            p.add_run(item.get('finding', 'No details')[:200])

    # Context Section
    context = analysis.get('context', {})
    if context:
        doc.add_heading('Business Context', level=1)
        p = doc.add_paragraph()
        p.add_run("Position: ").bold = True
        p.add_run(f"{context.get('position', 'N/A')}\n")
        p.add_run("Leverage: ").bold = True
        p.add_run(f"{context.get('leverage', 'N/A')}\n")
        if context.get('narrative'):
            p.add_run("Notes: ").bold = True
            p.add_run(context.get('narrative'))

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Generated by Contract Intelligence Platform (CIP)").italic = True

    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# Page config
st.set_page_config(
    page_title="Check for Risks",
    page_icon="🔍",
    layout="wide"
)

apply_spacing()
apply_theme()

# Inject centralized dark theme
from components.theme import inject_dark_theme
inject_dark_theme()

# Page-specific: prevent horizontal overflow for risk findings
st.markdown("""
<style>
.main .block-container { max-width: calc(100vw - 21rem) !important; overflow-x: hidden !important; }
.stCodeBlock, pre, code { max-width: 100% !important; white-space: pre-wrap !important; word-wrap: break-word !important; }
</style>
""", unsafe_allow_html=True)

# Sidebar
with st.sidebar:
    inject_cip_logo()

# Initialize contract context
init_contract_context()

# Check system health
api_healthy, db_healthy, ai_healthy = check_system_health()

# Initialize session state
if 'selected_contract_id' not in st.session_state:
    st.session_state.selected_contract_id = None
if 'analysis_result' not in st.session_state:
    st.session_state.analysis_result = None
if 'analyzing' not in st.session_state:
    st.session_state.analyzing = False
if 'risk_filter' not in st.session_state:
    st.session_state.risk_filter = "All"

# Page header
page_header("🔍 Check for Risks", "Automated risk assessment and clause analysis")
render_active_contract_header()

# Auto-load active contract from context
active_id = get_active_contract()
active_data = get_active_contract_data()

if active_id and not st.session_state.selected_contract_id:
    st.session_state.selected_contract_id = active_id


# Zone content functions
def z1_contract_selector():
    """Z1: Contract selector, scan controls"""
    section_header("Select Contract", "📄")

    if active_id:
        st.info(f"📋 Auto-loaded: #{active_id} - {active_data.get('title', 'Unknown')}")

    selected_contract_id, selected_contract, contracts = contract_selector(
        label="Choose a contract to analyze",
        key="contract_selector_z1",
        show_details=True
    )

    if selected_contract_id:
        st.session_state.selected_contract_id = selected_contract_id

    st.markdown("**Scan Options**")
    col1, col2 = st.columns(2)
    with col1:
        use_patterns = st.checkbox("Use Pattern Library", value=True, key="use_patterns_z1")
    with col2:
        deep_dive = st.checkbox("Deep Dive", value=False, key="deep_dive_z1")

    if st.button("🚀 Run Risk Scan", type="primary", use_container_width=True, key="run_scan_z1"):
        if st.session_state.selected_contract_id:
            st.session_state.analyzing = True
            st.rerun()
        else:
            st.warning("Select a contract first")


def z2_risk_summary():
    """Z2: Risk summary counts by severity"""
    section_header("Risk Summary", "📊")

    if st.session_state.analysis_result:
        analysis = st.session_state.analysis_result

        dealbreakers = len(analysis.get('dealbreakers', []))
        critical = len(analysis.get('critical_items', []))
        important = len(analysis.get('important_items', []))
        standard = len(analysis.get('standard_items', []))

        col1, col2 = st.columns(2)
        with col1:
            st.metric("🚫 Dealbreakers", dealbreakers)
            st.metric("🔴 Critical", critical)
        with col2:
            st.metric("🟡 Important", important)
            st.metric("🟢 Standard", standard)

        overall = analysis.get('overall_risk', 'N/A')
        confidence = analysis.get('confidence_score', 0)
        st.markdown(f"**Overall Risk:** {overall}")
        st.markdown(f"**Confidence:** {confidence:.0%}")
    else:
        st.info("Run a scan to see risk summary")


def z3_risk_filters():
    """Z3: Filter by risk type"""
    section_header("Filter Risks", "🔽")

    risk_filter = st.radio(
        "Show",
        ["All", "Dealbreakers", "Critical", "Important", "Standard"],
        key="risk_filter_z3",
        horizontal=False
    )
    st.session_state.risk_filter = risk_filter

    if st.session_state.analysis_result:
        st.markdown("---")
        if st.button("📤 Export Results", use_container_width=True, key="export_z3"):
            analysis = st.session_state.analysis_result
            json_str = json.dumps(analysis, indent=2)
            st.download_button(
                label="Download JSON",
                data=json_str,
                file_name=f"risk_analysis_{st.session_state.selected_contract_id}.json",
                mime="application/json",
                key="download_json_z3"
            )


def z4_risk_findings():
    """Z4: Risk findings list, clause analysis, severity indicators"""
    section_header("Risk Findings", "⚠️")

    if st.session_state.analysis_result:
        analysis = st.session_state.analysis_result
        risk_filter = st.session_state.risk_filter

        # Collect items based on filter
        items = []
        if risk_filter in ["All", "Dealbreakers"]:
            for item in analysis.get('dealbreakers', []):
                item['_level'] = 'DEALBREAKER'
                items.append(item)
        if risk_filter in ["All", "Critical"]:
            for item in analysis.get('critical_items', []):
                item['_level'] = 'CRITICAL'
                items.append(item)
        if risk_filter in ["All", "Important"]:
            for item in analysis.get('important_items', []):
                item['_level'] = 'IMPORTANT'
                items.append(item)
        if risk_filter in ["All", "Standard"]:
            for item in analysis.get('standard_items', []):
                item['_level'] = 'STANDARD'
                items.append(item)

        if items:
            for item in items:
                level = item.get('_level', 'UNKNOWN')
                icon = {'DEALBREAKER': '🚫', 'CRITICAL': '🔴', 'IMPORTANT': '🟡', 'STANDARD': '🟢'}.get(level, '⚪')

                with st.expander(f"{icon} {item.get('section_title', 'Unknown')} - {level}", expanded=(level in ['DEALBREAKER', 'CRITICAL'])):
                    st.markdown(f"**Section:** {item.get('section_number', 'N/A')}")
                    st.markdown(f"**Category:** {item.get('category', 'N/A')}")

                    if item.get('clause_text'):
                        st.markdown("**Clause Text:**")
                        st.code(item.get('clause_text'), language=None)

                    st.markdown(f"**Finding:** {item.get('finding', 'No details')}")

                    if item.get('redline_suggestion'):
                        st.markdown("**Suggested Redline:**")
                        st.markdown(f"> {item.get('redline_suggestion')}")

                    st.markdown(f"**Recommendation:** {item.get('recommendation', 'N/A')}")

                    if item.get('cascade_impacts'):
                        st.markdown("**Cascade Impacts:**")
                        for impact in item.get('cascade_impacts', []):
                            st.caption(f"→ {impact}")
        else:
            st.info("No findings match current filter")
    else:
        st.info("Run a risk scan to see findings")


def z5_risk_trends():
    """Z5: Risk trend metrics"""
    section_header("Risk Trends", "📈")

    if st.session_state.analysis_result:
        analysis = st.session_state.analysis_result

        # Calculate trend metrics
        dealbreakers = len(analysis.get('dealbreakers', []))
        critical = len(analysis.get('critical_items', []))
        high_risk_total = dealbreakers + critical

        st.metric("High Risk Items", high_risk_total)

        # Category breakdown
        categories = {}
        all_items = (
            analysis.get('dealbreakers', []) +
            analysis.get('critical_items', []) +
            analysis.get('important_items', [])
        )
        for item in all_items:
            cat = item.get('category', 'Other')
            categories[cat] = categories.get(cat, 0) + 1

        if categories:
            st.markdown("**By Category:**")
            for cat, count in sorted(categories.items(), key=lambda x: x[1], reverse=True)[:5]:
                st.caption(f"{cat}: {count}")
    else:
        st.info("Run scan for trends")


def z6_mitigation_tips():
    """Z6: Risk mitigation tips"""
    section_header("Mitigation Tips", "💡")

    if st.session_state.analysis_result:
        analysis = st.session_state.analysis_result

        # Generate tips based on findings
        tips = []
        if analysis.get('dealbreakers'):
            tips.append("Address dealbreakers before proceeding")
        if analysis.get('critical_items'):
            tips.append("Negotiate critical items with counterparty")
        if len(analysis.get('important_items', [])) > 5:
            tips.append("Consider grouping important items for batch negotiation")

        context = analysis.get('context', {})
        if context.get('leverage') == 'Weak':
            tips.append("Focus on non-monetary concessions")
        if context.get('position') == 'Vendor':
            tips.append("Emphasize service limitations")

        if tips:
            for tip in tips:
                st.markdown(f"• {tip}")
        else:
            st.success("No immediate mitigation needed")
    else:
        st.markdown("**General Tips:**")
        st.markdown("• Always review dealbreakers first")
        st.markdown("• Use Pattern Library for consistency")
        st.markdown("• Export findings for team review")


def z7_extra():
    """Z7: Additional controls"""
    pass


# Render zone layout
zone_layout(
    z1=z1_contract_selector,
    z2=z2_risk_summary,
    z3=z3_risk_filters,
    z4=z4_risk_findings,
    z5=z5_risk_trends,
    z6=z6_mitigation_tips,
    z7=z7_extra,
    z7_system_status=True,
    api_healthy=api_healthy,
    db_healthy=db_healthy,
    ai_healthy=ai_healthy
)

# Run analysis if triggered
if st.session_state.analyzing:
    result, error = api_call_with_spinner(
        endpoint="/api/analyze",
        method="POST",
        data={'contract_id': st.session_state.selected_contract_id},
        spinner_message="Analyzing contract with 59-pattern deck... This may take 30-90 seconds...",
        success_message="Risk scan complete!",
        timeout=300,
        result_key=None
    )

    if error:
        st.error(f"Analysis failed: {error}")
    else:
        st.session_state.analysis_result = result.get('analysis') if isinstance(result, dict) else result
        st.session_state.patterns_metadata = result.get('patterns', {}) if isinstance(result, dict) else {}

    st.session_state.analyzing = False
    if result:
        st.rerun()

