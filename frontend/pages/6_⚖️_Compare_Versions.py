"""
CIP - Compare Two Versions Page
v2.0: CIP_UIX_Core_v2 - Progress overlay, ComparisonSnapshot, clause alignment, side-by-side UI
"""

import streamlit as st
import json
import sys
import os
import requests
from datetime import datetime
from io import BytesIO

# Add parent directory to path
sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))

from ui_components import (
    page_header, section_header,
    toast_success, toast_info, toast_error, apply_spacing
)
from shared_components import api_call_with_spinner
from components.contract_context import (
    get_active_contract,
    get_active_contract_data,
    render_active_contract_header,
    init_contract_context
)
from cip_components import split_panel
from zone_layout import zone_layout, check_system_health
from theme_system import apply_theme, inject_cip_logo
from progress_indicators import render_progress_overlay

# Inject centralized dark theme
from components.theme import inject_dark_theme

# ============================================================================
# v2: COMPARE PROGRESS STAGES
# ============================================================================

COMPARE_STAGES = [
    {"name": "Loading V1 & V2", "percent": 10},
    {"name": "Aligning clauses", "percent": 30},
    {"name": "Calculating text deltas", "percent": 55},
    {"name": "Applying risk/severity", "percent": 80},
    {"name": "Preparing view", "percent": 100},
]

# GEM Tier-3 Copy
GEM_COPY = {
    "page_title": "Compare Contract Versions",
    "section_select": "Select Versions",
    "section_comparison": "Clause Comparison",
    "section_impact": "Business Impact",
    "no_changes": "No material changes detected between these versions.",
    "impact_template": "This comparison identified {added} added clauses, {removed} removed clauses, and {modified} modified clauses. Overall risk changed from {old_risk} to {new_risk}.",
    "match_high": "High Match",
    "match_likely": "Likely Match",
    "match_guess": "Best Guess",
    "match_added": "Added in V2",
    "match_removed": "Removed from V1",
    # Phase 4B: Error message keys (CAI taxonomy)
    "compare.network_failure": "We're having trouble reaching the analysis service. Please check your connection and try again.",
    "compare.auth_failure": "Authentication failed. Please verify your API key is configured correctly.",
    "compare.payload_failure": "The contract documents are too large to compare. Try with smaller files.",
    "compare.internal_failure": "Something went wrong during comparison. Please try again in a moment.",
    # Phase 4E: Compare v3 Engine messages
    "compare_v3.title": "Compare v3 (Early Preview)",
    "compare_v3.description": "Next-generation semantic comparison using SAE/ERCE/BIRL pipeline.",
    "compare_v3.placeholder_note": "Phase 4E: Deterministic placeholder outputs. Real intelligence activates in Phase 4F.",
}

API_BASE = "http://localhost:5000/api"


# ============================================================================
# v2: API FUNCTIONS
# ============================================================================

def get_contracts_for_selector():
    """Fetch all contracts for V1/V2 dropdowns"""
    try:
        resp = requests.get(f"{API_BASE}/contracts", timeout=10)
        if resp.ok:
            contracts = resp.json()
            # Format for dropdown: display_id + title
            options = []
            for c in contracts:
                display_id = c.get('display_id') or c.get('id')
                title = c.get('title') or c.get('filename', 'Untitled')
                options.append({
                    'id': c.get('id'),
                    'display_id': display_id,
                    'title': title,
                    'label': f"#{display_id} - {title}"
                })
            return options
        return []
    except Exception:
        return []


def get_contract_clauses(contract_id: int) -> list:
    """Get clause list for a contract (from latest analysis or extraction)"""
    try:
        # Try to get from latest analysis snapshot
        resp = requests.get(f"{API_BASE}/contracts/{contract_id}/analysis/latest", timeout=10)
        if resp.ok:
            data = resp.json()
            return data.get('clauses', [])

        # Fallback: try to get from contract text extraction
        resp = requests.get(f"{API_BASE}/contracts/{contract_id}/clauses", timeout=10)
        if resp.ok:
            return resp.json().get('clauses', [])

        return []
    except Exception:
        return []


def get_latest_snapshot_id(contract_id: int) -> int | None:
    """Get latest analysis snapshot ID for a contract"""
    try:
        resp = requests.get(f"{API_BASE}/contracts/{contract_id}/analysis/latest", timeout=5)
        if resp.ok:
            return resp.json().get('snapshot_id')
        return None
    except Exception:
        return None


def save_comparison_snapshot_api(
    v1_id: int,
    v2_id: int,
    v1_snapshot_id: int | None,
    v2_snapshot_id: int | None,
    similarity_score: float,
    changed_clauses: list,
    risk_delta: list
) -> dict | None:
    """Save comparison snapshot via API"""
    try:
        resp = requests.post(
            f"{API_BASE}/comparisons/snapshot",
            json={
                'v1_contract_id': v1_id,
                'v2_contract_id': v2_id,
                'v1_snapshot_id': v1_snapshot_id,
                'v2_snapshot_id': v2_snapshot_id,
                'similarity_score': similarity_score,
                'changed_clauses': changed_clauses,
                'risk_delta': risk_delta
            },
            timeout=10
        )
        if resp.ok:
            return resp.json()
        return None
    except Exception:
        return None


def export_comparison_docx(comparison_result: dict, v1_title: str, v2_title: str) -> BytesIO:
    """Generate DOCX export (placeholder)"""
    try:
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        doc.add_heading('Contract Comparison Report', 0)
        doc.add_paragraph(f"V1: {v1_title}")
        doc.add_paragraph(f"V2: {v2_title}")
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

        doc.add_heading('Summary', level=1)
        doc.add_paragraph(f"Total changes: {comparison_result.get('total_changes', 0)}")
        doc.add_paragraph(f"Similarity: {comparison_result.get('similarity_score', 0):.1f}%")

        doc.add_heading('Changed Clauses', level=1)
        for clause in comparison_result.get('aligned_clauses', [])[:20]:
            if clause.get('match_type') != 'High Match':
                v1_title = clause.get('v1', {}).get('title', 'N/A') if clause.get('v1') else 'Removed'
                v2_title = clause.get('v2', {}).get('title', 'N/A') if clause.get('v2') else 'Added'
                doc.add_paragraph(f"- {v1_title} -> {v2_title} ({clause.get('match_type')})")

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    except ImportError:
        # Fallback if python-docx not installed
        buffer = BytesIO()
        buffer.write(b"DOCX export requires python-docx package")
        buffer.seek(0)
        return buffer


def export_comparison_json(comparison_result: dict) -> str:
    """Generate JSON export"""
    return json.dumps(comparison_result, indent=2, default=str)


# ============================================================================
# PHASE 4E: COMPARE V3 API FUNCTION
# ============================================================================

def run_compare_v3_api(v1_id: int, v2_id: int) -> dict:
    """
    Call Compare v3 Engine endpoint.

    Phase 4E: Returns deterministic placeholder results.
    Phase 4F: Will return real semantic analysis.

    Args:
        v1_id: V1 contract ID
        v2_id: V2 contract ID

    Returns:
        Dict with success flag and pipeline data
    """
    try:
        resp = requests.post(
            f"{API_BASE}/compare-v3",
            json={
                'v1_contract_id': v1_id,
                'v2_contract_id': v2_id
            },
            timeout=30
        )
        if resp.ok:
            return resp.json()
        else:
            return {
                'success': False,
                'error_message_key': 'compare.internal_failure'
            }
    except requests.exceptions.Timeout:
        return {
            'success': False,
            'error_message_key': 'compare.network_failure'
        }
    except Exception as e:
        return {
            'success': False,
            'error_message_key': 'compare.internal_failure',
            'error_detail': str(e)
        }


# ============================================================================
# PAGE CONFIG
# ============================================================================

st.set_page_config(
    page_title="Compare Versions",
    page_icon="⚖️",
    layout="wide"
)

apply_spacing()
apply_theme()
inject_dark_theme()

# Sidebar
with st.sidebar:
    inject_cip_logo()

# Initialize contract context
init_contract_context()

# Check system health
api_healthy, db_healthy, ai_healthy = check_system_health()

# ============================================================================
# SESSION STATE
# ============================================================================

if 'v1_id' not in st.session_state:
    st.session_state.v1_id = None
if 'v2_id' not in st.session_state:
    st.session_state.v2_id = None
if 'v1_contract' not in st.session_state:
    st.session_state.v1_contract = None
if 'v2_contract' not in st.session_state:
    st.session_state.v2_contract = None
if 'comparing' not in st.session_state:
    st.session_state.comparing = False
if 'comparison_result' not in st.session_state:
    st.session_state.comparison_result = None
if 'compare_progress_stage' not in st.session_state:
    st.session_state.compare_progress_stage = 0
if 'compare_error' not in st.session_state:
    st.session_state.compare_error = None
if 'snapshot_saved' not in st.session_state:
    st.session_state.snapshot_saved = False
# Phase 4E: Compare v3 session state
if 'compare_v3_result' not in st.session_state:
    st.session_state.compare_v3_result = None
if 'compare_v3_running' not in st.session_state:
    st.session_state.compare_v3_running = False

# ============================================================================
# ZONE CONTENT FUNCTIONS
# ============================================================================


def z1_version_selectors():
    """Z1: V1/V2 contract selectors - GEM Section: Select Versions"""
    section_header(GEM_COPY["section_select"], "📄")

    # Auto-load active contract info
    active_id = get_active_contract()
    active_data = get_active_contract_data()

    if active_id:
        st.info(f"📋 Active contract: #{active_id} - {active_data.get('title', 'Unknown')}")

    # Fetch contracts for dropdown
    contracts = get_contracts_for_selector()

    if not contracts:
        st.warning("No contracts available. Upload contracts first.")
        return

    # Build options
    options = {c['label']: c for c in contracts}
    option_labels = ["Select..."] + list(options.keys())

    # V1 selector
    v1_label = st.selectbox(
        "V1 (Original)",
        options=option_labels,
        key="v1_selector_dropdown"
    )

    # V2 selector
    v2_label = st.selectbox(
        "V2 (Revised)",
        options=option_labels,
        key="v2_selector_dropdown"
    )

    # Update session state
    if v1_label != "Select...":
        st.session_state.v1_id = options[v1_label]['id']
        st.session_state.v1_contract = options[v1_label]
    else:
        st.session_state.v1_id = None

    if v2_label != "Select...":
        st.session_state.v2_id = options[v2_label]['id']
        st.session_state.v2_contract = options[v2_label]
    else:
        st.session_state.v2_id = None

    # Validation: two distinct versions required
    v1_id = st.session_state.v1_id
    v2_id = st.session_state.v2_id

    if v1_id and v2_id:
        if v1_id == v2_id:
            st.error("V1 and V2 must be different contracts")
        else:
            st.success("Ready to compare")

            if st.button("🔍 Compare Versions", type="primary", use_container_width=True):
                st.session_state.comparing = True
                st.session_state.comparison_result = None
                st.session_state.compare_error = None
                st.session_state.snapshot_saved = False
                st.rerun()
    else:
        st.caption("Select both V1 and V2 to compare")


def z2_comparison_summary():
    """Z2: Comparison summary stats"""
    section_header("Summary", "📊")

    if st.session_state.comparison_result:
        result = st.session_state.comparison_result

        similarity = result.get('similarity_score', 0)
        total_changes = result.get('total_changes', 0)
        added = result.get('added_count', 0)
        removed = result.get('removed_count', 0)
        modified = result.get('modified_count', 0)

        st.metric("Similarity", f"{similarity:.1f}%")
        st.metric("Total Changes", total_changes)

        col1, col2 = st.columns(2)
        with col1:
            st.metric("➕ Added", added)
        with col2:
            st.metric("➖ Removed", removed)

        st.metric("✏️ Modified", modified)

        # Snapshot saved indicator
        if st.session_state.snapshot_saved:
            st.success("💾 Comparison saved")
    else:
        st.info("Run comparison to see summary")


def z3_match_legend():
    """Z3: Match type legend"""
    section_header("Match Types", "🏷️")

    st.markdown(f"""
    <div style="font-size: 13px; line-height: 1.8;">
        <div>🟢 <strong>{GEM_COPY['match_high']}</strong><br/>
        <span style="color: #94A3B8;">Clauses match closely</span></div>
        <div style="margin-top: 8px;">🟡 <strong>{GEM_COPY['match_likely']}</strong><br/>
        <span style="color: #94A3B8;">Similar content, review recommended</span></div>
        <div style="margin-top: 8px;">🟠 <strong>{GEM_COPY['match_guess']}</strong><br/>
        <span style="color: #94A3B8;">Possible match, verify manually</span></div>
        <div style="margin-top: 8px;">🔵 <strong>{GEM_COPY['match_added']}</strong><br/>
        <span style="color: #94A3B8;">New clause in revised version</span></div>
        <div style="margin-top: 8px;">🔴 <strong>{GEM_COPY['match_removed']}</strong><br/>
        <span style="color: #94A3B8;">Clause removed from original</span></div>
    </div>
    """, unsafe_allow_html=True)


def z4_side_by_side():
    """Z4: Side-by-side clause comparison - GEM Section: Clause Comparison"""
    section_header(GEM_COPY["section_comparison"], "📝")

    if not st.session_state.comparison_result:
        st.info("Select contracts and run comparison to see side-by-side view")
        return

    result = st.session_state.comparison_result
    aligned = result.get('aligned_clauses', [])

    if not aligned:
        st.info(GEM_COPY["no_changes"])
        return

    # Side-by-side columns
    col_v1, col_v2 = st.columns(2)

    with col_v1:
        st.markdown("#### V1: Original")
        v1_contract = st.session_state.v1_contract
        if v1_contract:
            st.caption(v1_contract.get('title', 'Unknown'))

    with col_v2:
        st.markdown("#### V2: Revised")
        v2_contract = st.session_state.v2_contract
        if v2_contract:
            st.caption(v2_contract.get('title', 'Unknown'))

    st.markdown("---")

    # Show aligned clauses
    for idx, pair in enumerate(aligned[:30]):  # Limit to 30 for performance
        match_type = pair.get('match_type', 'Unknown')
        similarity = pair.get('similarity', 0)

        # Match type badge
        badge_colors = {
            "High Match": "#10B981",
            "Likely Match": "#F59E0B",
            "Best Guess": "#F97316",
            "Added": "#3B82F6",
            "No Match": "#EF4444"
        }
        badge_color = badge_colors.get(match_type, "#6B7280")

        with st.expander(f"{match_type} ({similarity:.0%})", expanded=(match_type != "High Match")):
            exp_col1, exp_col2 = st.columns(2)

            with exp_col1:
                v1_clause = pair.get('v1')
                if v1_clause:
                    st.markdown(f"**{v1_clause.get('number', 'N/A')}. {v1_clause.get('title', 'Untitled')}**")
                    st.caption(v1_clause.get('severity', ''))
                else:
                    st.markdown("*— Removed —*")

            with exp_col2:
                v2_clause = pair.get('v2')
                if v2_clause:
                    st.markdown(f"**{v2_clause.get('number', 'N/A')}. {v2_clause.get('title', 'Untitled')}**")
                    st.caption(v2_clause.get('severity', ''))
                else:
                    st.markdown("*— Added —*")

            # Adjust match stub (placeholder for manual override)
            st.markdown(
                f"<a href='#' style='font-size: 12px; color: #60A5FA;'>Adjust match</a>",
                unsafe_allow_html=True
            )

    if len(aligned) > 30:
        st.caption(f"Showing 30 of {len(aligned)} clause pairs")


def z5_business_impact():
    """Z5: Business impact narrative - GEM Section: Business Impact"""
    section_header(GEM_COPY["section_impact"], "💼")

    if not st.session_state.comparison_result:
        st.info("Run comparison for impact analysis")
        return

    result = st.session_state.comparison_result

    added = result.get('added_count', 0)
    removed = result.get('removed_count', 0)
    modified = result.get('modified_count', 0)
    old_risk = result.get('v1_risk', 'Unknown')
    new_risk = result.get('v2_risk', 'Unknown')

    # GEM impact narrative template
    if added == 0 and removed == 0 and modified == 0:
        st.info(GEM_COPY["no_changes"])
    else:
        narrative = GEM_COPY["impact_template"].format(
            added=added,
            removed=removed,
            modified=modified,
            old_risk=old_risk,
            new_risk=new_risk
        )
        st.markdown(narrative)

    # Risk delta visualization
    risk_delta = result.get('risk_delta', [])
    if risk_delta:
        st.markdown("**Risk Changes by Category:**")
        for delta in risk_delta[:6]:
            cat = delta.get('category', 'Unknown')
            v1_risk = delta.get('v1_risk', 'N/A')
            v2_risk = delta.get('v2_risk', 'N/A')

            if v1_risk != v2_risk:
                st.markdown(f"- **{cat}**: {v1_risk} → {v2_risk}")


def z6_export_actions():
    """Z6: Export buttons"""
    section_header("Export", "📤")

    if not st.session_state.comparison_result:
        st.info("Run comparison to enable exports")
        return

    result = st.session_state.comparison_result
    v1_title = st.session_state.v1_contract.get('title', 'V1') if st.session_state.v1_contract else 'V1'
    v2_title = st.session_state.v2_contract.get('title', 'V2') if st.session_state.v2_contract else 'V2'

    # Export DOCX
    if st.button("📄 Export DOCX", use_container_width=True):
        try:
            docx_buffer = export_comparison_docx(result, v1_title, v2_title)
            st.download_button(
                label="Download DOCX",
                data=docx_buffer,
                file_name=f"comparison_{st.session_state.v1_id}_vs_{st.session_state.v2_id}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="docx_download"
            )
            toast_success("DOCX export ready")
        except Exception as e:
            toast_error(f"DOCX export failed: {str(e)}")

    # Export JSON
    if st.button("📋 Export JSON", use_container_width=True):
        try:
            json_str = export_comparison_json(result)
            st.download_button(
                label="Download JSON",
                data=json_str,
                file_name=f"comparison_{st.session_state.v1_id}_vs_{st.session_state.v2_id}.json",
                mime="application/json",
                key="json_download"
            )
            toast_success("JSON export ready")
        except Exception as e:
            toast_error(f"JSON export failed: {str(e)}")

    # ========================================================================
    # PHASE 4E: COMPARE V3 (EARLY PREVIEW) BUTTON
    # ========================================================================
    st.markdown("---")
    section_header(GEM_COPY["compare_v3.title"], "🧪")
    st.caption(GEM_COPY["compare_v3.description"])

    # Compare v3 button - requires valid comparison first
    if st.session_state.v1_id and st.session_state.v2_id:
        if st.button("🧪 Run Compare v3 (Preview)", use_container_width=True, key="compare_v3_btn"):
            st.session_state.compare_v3_running = True
            st.rerun()
    else:
        st.info("Select V1 and V2 contracts first")

    # Display Compare v3 results if available
    if st.session_state.compare_v3_result:
        v3_result = st.session_state.compare_v3_result
        if v3_result.get('success'):
            data = v3_result.get('data', {})
            meta = data.get('_meta', {})

            st.success(f"Compare v3 Engine: {meta.get('engine_version', 'Unknown')}")
            st.caption(GEM_COPY["compare_v3.placeholder_note"])

            # SAE Matches
            with st.expander("SAE Matches (Semantic Alignment)", expanded=False):
                for match in data.get('sae_matches', []):
                    conf = match.get('match_confidence', 'N/A')
                    score = match.get('similarity_score', 0)
                    st.markdown(f"- Clause {match.get('v1_clause_id')} ↔ {match.get('v2_clause_id')}: **{conf}** ({score:.0%})")

            # ERCE Results
            with st.expander("ERCE Results (Risk Classification)", expanded=False):
                for risk in data.get('erce_results', []):
                    cat = risk.get('risk_category', 'N/A')
                    conf = risk.get('confidence', 0)
                    st.markdown(f"- Pair {risk.get('clause_pair_id')}: **{cat}** (confidence: {conf:.0%})")

            # BIRL Narratives
            with st.expander("BIRL Narratives (Business Impact)", expanded=False):
                for narrative in data.get('birl_narratives', []):
                    st.markdown(f"**Pair {narrative.get('clause_pair_id')}**: {narrative.get('narrative')}")
                    dims = ', '.join(narrative.get('impact_dimensions', []))
                    st.caption(f"Impact: {dims}")
        else:
            error_key = v3_result.get('error_message_key', 'compare.internal_failure')
            st.error(GEM_COPY.get(error_key, GEM_COPY["compare.internal_failure"]))


def z7_footer():
    """Z7: Footer"""
    st.markdown("---")
    st.caption(f"Compare v2.0 | {datetime.now().strftime('%Y-%m-%d %H:%M')}")


# ============================================================================
# RENDER ZONE LAYOUT
# ============================================================================

zone_layout(
    z1=z1_version_selectors,
    z2=z2_comparison_summary,
    z3=z3_match_legend,
    z4=z4_side_by_side,
    z5=z5_business_impact,
    z6=z6_export_actions,
    z7=z7_footer,
    z7_system_status=True,
    api_healthy=api_healthy,
    db_healthy=db_healthy,
    ai_healthy=ai_healthy
)


# ============================================================================
# v2: RUN COMPARISON WITH PROGRESS OVERLAY
# ============================================================================

if st.session_state.comparing:
    v1_id = st.session_state.v1_id
    v2_id = st.session_state.v2_id

    if not v1_id or not v2_id or v1_id == v2_id:
        st.error("Two distinct contract versions required")
        st.session_state.comparing = False
    else:
        # Show progress overlay
        current_stage = min(st.session_state.compare_progress_stage, len(COMPARE_STAGES) - 1)
        stage = COMPARE_STAGES[current_stage]
        render_progress_overlay(
            mode="determinate",
            percent=stage["percent"],
            label=stage["name"]
        )

        # Perform comparison via API
        result, error = api_call_with_spinner(
            endpoint="/api/compare",
            method="POST",
            data={
                'v1_contract_id': v1_id,
                'v2_contract_id': v2_id,
                'include_recommendations': True
            },
            spinner_message="Comparing versions...",
            success_message="Comparison complete!",
            timeout=600
        )

        if error:
            # Phase 4B: Handle structured error response
            if isinstance(error, dict):
                error_key = error.get('error_message_key', 'compare.internal_failure')
                user_message = GEM_COPY.get(error_key, GEM_COPY["compare.internal_failure"])
            else:
                user_message = str(error)

            st.session_state.compare_error = user_message
            render_progress_overlay(
                mode="error",
                label="Comparison Failed",
                error_message=user_message
            )
        else:
            # Process result and build aligned clauses
            comparison_data = result if isinstance(result, dict) else {}

            # Get clause lists for alignment
            v1_clauses = get_contract_clauses(v1_id)
            v2_clauses = get_contract_clauses(v2_id)

            # Simple clause alignment
            from difflib import SequenceMatcher

            aligned_clauses = []
            v2_matched = set()
            added_count = 0
            removed_count = 0
            modified_count = 0

            for v1_clause in v1_clauses:
                best_match = None
                best_score = 0.0
                best_idx = -1

                for idx, v2_clause in enumerate(v2_clauses):
                    if idx in v2_matched:
                        continue

                    # Score by number + title similarity
                    number_match = 1.0 if v1_clause.get('number') == v2_clause.get('number') else 0.0
                    title_sim = SequenceMatcher(
                        None,
                        (v1_clause.get('title') or '').lower(),
                        (v2_clause.get('title') or '').lower()
                    ).ratio()

                    score = (number_match * 0.4) + (title_sim * 0.6)

                    if score > best_score:
                        best_score = score
                        best_match = v2_clause
                        best_idx = idx

                # Determine match type
                if best_score >= 0.85:
                    match_type = "High Match"
                elif best_score >= 0.60:
                    match_type = "Likely Match"
                    modified_count += 1
                elif best_score >= 0.40:
                    match_type = "Best Guess"
                    modified_count += 1
                else:
                    match_type = "No Match"
                    removed_count += 1
                    best_match = None

                if best_match and best_idx >= 0:
                    v2_matched.add(best_idx)

                aligned_clauses.append({
                    'v1': v1_clause,
                    'v2': best_match,
                    'match_type': match_type,
                    'similarity': best_score
                })

            # Add unmatched V2 clauses as "Added"
            for idx, v2_clause in enumerate(v2_clauses):
                if idx not in v2_matched:
                    aligned_clauses.append({
                        'v1': None,
                        'v2': v2_clause,
                        'match_type': "Added",
                        'similarity': 0.0
                    })
                    added_count += 1

            # Calculate similarity score
            total_clauses = max(len(v1_clauses), len(v2_clauses), 1)
            high_matches = sum(1 for a in aligned_clauses if a['match_type'] == 'High Match')
            similarity_score = (high_matches / total_clauses) * 100

            # Build result object
            st.session_state.comparison_result = {
                'v1_contract_id': v1_id,
                'v2_contract_id': v2_id,
                'aligned_clauses': aligned_clauses,
                'similarity_score': similarity_score,
                'total_changes': added_count + removed_count + modified_count,
                'added_count': added_count,
                'removed_count': removed_count,
                'modified_count': modified_count,
                'v1_risk': comparison_data.get('v1_risk', 'Unknown'),
                'v2_risk': comparison_data.get('v2_risk', 'Unknown'),
                'risk_delta': comparison_data.get('risk_delta', []),
                'created_at': datetime.now().isoformat()
            }

            # Save ComparisonSnapshot
            v1_snapshot_id = get_latest_snapshot_id(v1_id)
            v2_snapshot_id = get_latest_snapshot_id(v2_id)

            snapshot_result = save_comparison_snapshot_api(
                v1_id=v1_id,
                v2_id=v2_id,
                v1_snapshot_id=v1_snapshot_id,
                v2_snapshot_id=v2_snapshot_id,
                similarity_score=similarity_score,
                changed_clauses=[
                    {
                        'v1_number': a['v1'].get('number') if a['v1'] else None,
                        'v2_number': a['v2'].get('number') if a['v2'] else None,
                        'match_type': a['match_type']
                    }
                    for a in aligned_clauses if a['match_type'] != 'High Match'
                ],
                risk_delta=comparison_data.get('risk_delta', [])
            )

            if snapshot_result:
                st.session_state.snapshot_saved = True

        # Reset comparison state
        st.session_state.comparing = False
        st.session_state.compare_progress_stage = 0

        if st.session_state.comparison_result:
            st.rerun()


# ============================================================================
# PHASE 4E: COMPARE V3 EXECUTION
# ============================================================================

if st.session_state.compare_v3_running:
    v1_id = st.session_state.v1_id
    v2_id = st.session_state.v2_id

    if not v1_id or not v2_id:
        st.error("V1 and V2 contract IDs required for Compare v3")
        st.session_state.compare_v3_running = False
    else:
        with st.spinner("Running Compare v3 pipeline (SAE/ERCE/BIRL)..."):
            result = run_compare_v3_api(v1_id, v2_id)
            st.session_state.compare_v3_result = result
            st.session_state.compare_v3_running = False
            st.rerun()
