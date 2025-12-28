"""
CIP - Compare Two Versions Page
v2.1: Phase 4F - Real intelligence pipeline with SAE/ERCE/BIRL/FAR

Phase 4F Updates:
- SAE: OpenAI text-embedding-3-large with cosine similarity
- ERCE: Pattern Library v1.2 keyword matching
- BIRL: Claude narrative generation (max 150 tokens, temp 0.3)
- FAR: Rule-based flowdown gap analysis
- GEM 4F truncation: SAE max 10, ERCE max 15, BIRL max 5, FAR max 5
- Real intelligence with fallback to placeholders on failure
"""

import streamlit as st
import json
import sys
import os
import requests
from datetime import datetime
from io import BytesIO

# Add parent directory to path
sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))

from ui_components import (
    page_header, section_header,
    toast_success, toast_info, toast_error, apply_spacing
)
from shared_components import api_call_with_spinner, render_trust_banner, clear_trust_banner_dismissal, render_error_card
from components.trust_engine import run_trust_precheck
from components.contract_context import (
    get_active_contract,
    get_active_contract_data,
    render_active_contract_header,
    init_contract_context
)
from cip_components import split_panel
from zone_layout import zone_layout, check_system_health
# GEM Tier-3 Copy
GEM_COPY = {
    "page_title": "Compare Contract Versions",
    "section_select": "Select Versions",
    "section_comparison": "Clause Comparison",
    "section_impact": "Business Impact",
    "no_changes": "No material changes detected between these versions.",
    "impact_template": "This comparison identified {added} added clauses, {removed} removed clauses, and {modified} modified clauses. Overall risk changed from {old_risk} to {new_risk}.",
    "match_high": "High Match",
    "match_likely": "Likely Match",
    "match_guess": "Best Guess",
    "match_added": "Added in V2",
    "match_removed": "Removed from V1",
    # Phase 4B: Error message keys (CAI taxonomy)
    "compare.network_failure": "We're having trouble reaching the analysis service. Please check your connection and try again.",
    "compare.auth_failure": "Authentication failed. Please verify your API key is configured correctly.",
    "compare.payload_failure": "The contract documents are too large to compare. Try with smaller files.",
    "compare.internal_failure": "Something went wrong during comparison. Please try again in a moment.",
    # Phase 4F: Compare v3 Engine messages (upgraded from 4E)
    "compare_v3.title": "Compare v3 Engine",
    "compare_v3.description": "AI-powered semantic comparison using SAE/ERCE/BIRL/FAR pipeline.",
    "compare_v3.placeholder_note": "Phase 4E: Deterministic placeholder outputs. Real intelligence activates in Phase 4F.",
    # Phase 4F copy keys
    "p4f.sae_title": "Semantic Alignment (SAE)",
    "p4f.erce_title": "Risk Classification (ERCE)",
    "p4f.birl_title": "Business Impact (BIRL)",
    "p4f.far_title": "Flowdown Analysis (FAR)",
    "p4f.narrative_dimension": "Impact Dimensions",
    "p4f.truncation_label": "+ {n} more {type}...",
    "p4f.intelligence_active": "Real intelligence active",
    "p4f.intelligence_partial": "Partial intelligence (fallback active)",
    "p4f.intelligence_placeholder": "Placeholder mode",
}

API_BASE = "http://localhost:5000/api"


# ============================================================================
# v2: API FUNCTIONS
# ============================================================================

def get_contracts_for_selector():
    """Fetch all contracts for V1/V2 dropdowns"""
    try:
        resp = requests.get(f"{API_BASE}/contracts", timeout=10)
        if resp.ok:
            data = resp.json()
            # Handle both list and dict response formats
            contracts = data.get('contracts', data) if isinstance(data, dict) else data
            # Format for dropdown: display_id + title
            options = []
            for c in contracts:
                display_id = c.get('display_id') or c.get('id')
                title = c.get('title') or c.get('filename', 'Untitled')
                options.append({
                    'id': c.get('id'),
                    'display_id': display_id,
                    'title': title,
                    'label': f"#{display_id} - {title}"
                })
            return options
        return []
    except Exception:
        return []


def get_contract_clauses(contract_id: int) -> list:
    """Get clause list for a contract"""
    try:
        # Use the correct endpoint: /api/contract/{id} (singular)
        resp = requests.get(f"{API_BASE}/contract/{contract_id}", timeout=10)
        if resp.ok:
            data = resp.json()
            return data.get('clauses', [])
        return []
    except Exception:
        return []


def get_latest_snapshot_id(contract_id: int) -> int | None:
    """Get latest analysis snapshot ID for a contract"""
    try:
        resp = requests.get(f"{API_BASE}/contract/{contract_id}", timeout=5)
        if resp.ok:
            data = resp.json()
            # Get latest assessment ID (assessments ordered by date DESC)
            assessments = data.get('assessments', [])
            if assessments:
                return assessments[0].get('id')
        return None
    except Exception:
        return None


def save_comparison_snapshot_api(
    v1_id: int,
    v2_id: int,
    v1_snapshot_id: int | None,
    v2_snapshot_id: int | None,
    similarity_score: float,
    changed_clauses: list,
    risk_delta: list
) -> dict | None:
    """Save comparison snapshot via API"""
    try:
        resp = requests.post(
            f"{API_BASE}/comparisons/snapshot",
            json={
                'v1_contract_id': v1_id,
                'v2_contract_id': v2_id,
                'v1_snapshot_id': v1_snapshot_id,
                'v2_snapshot_id': v2_snapshot_id,
                'similarity_score': similarity_score,
                'changed_clauses': changed_clauses,
                'risk_delta': risk_delta
            },
            timeout=10
        )
        if resp.ok:
            return resp.json()
        return None
    except Exception:
        return None


def export_comparison_docx(comparison_result: dict, v1_title: str, v2_title: str) -> BytesIO:
    """Generate DOCX export (placeholder)"""
    try:
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        doc.add_heading('Contract Comparison Report', 0)
        doc.add_paragraph(f"V1: {v1_title}")
        doc.add_paragraph(f"V2: {v2_title}")
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

        doc.add_heading('Summary', level=1)
        doc.add_paragraph(f"Total changes: {comparison_result.get('total_changes', 0)}")
        doc.add_paragraph(f"Similarity: {comparison_result.get('similarity_score', 0):.1f}%")

        doc.add_heading('Changed Clauses', level=1)
        for clause in comparison_result.get('aligned_clauses', [])[:20]:
            if clause.get('match_type') != 'High Match':
                v1_title = (clause.get('v1', {}).get('title') or 'N/A') if clause.get('v1') else 'Removed'
                v2_title = (clause.get('v2', {}).get('title') or 'N/A') if clause.get('v2') else 'Added'
                doc.add_paragraph(f"- {v1_title} -> {v2_title} ({clause.get('match_type')})")

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    except ImportError:
        # Fallback if python-docx not installed
        buffer = BytesIO()
        buffer.write(b"DOCX export requires python-docx package")
        buffer.seek(0)
        return buffer


def export_comparison_json(comparison_result: dict) -> str:
    """Generate JSON export"""
    return json.dumps(comparison_result, indent=2, default=str)


# ============================================================================
# PAGE CONFIG - st.set_page_config is in app.py ONLY
# ============================================================================

apply_spacing()

# Initialize contract context
init_contract_context()

# Check system health
api_healthy, db_healthy, ai_healthy = check_system_health()

# ============================================================================
# SESSION STATE
# ============================================================================

if 'v1_id' not in st.session_state:
    st.session_state["v1_id"] = None
if 'v2_id' not in st.session_state:
    st.session_state["v2_id"] = None
if 'v1_contract' not in st.session_state:
    st.session_state["v1_contract"] = None
if 'v2_contract' not in st.session_state:
    st.session_state["v2_contract"] = None
if 'comparing' not in st.session_state:
    st.session_state["comparing"] = False
if 'comparison_result' not in st.session_state:
    st.session_state["comparison_result"] = None
if 'compare_error' not in st.session_state:
    st.session_state["compare_error"] = None
if 'snapshot_saved' not in st.session_state:
    st.session_state["snapshot_saved"] = False
if 'comparison_cached' not in st.session_state:
    st.session_state["comparison_cached"] = False
if 'comparison_hash' not in st.session_state:
    st.session_state["comparison_hash"] = None
# Phase 4E: Compare v3 session state
if 'compare_v3_result' not in st.session_state:
    st.session_state["compare_v3_result"] = None
if 'compare_v3_running' not in st.session_state:
    st.session_state["compare_v3_running"] = False
# Navigation session state for Redline-style view
if 'compare_current_index' not in st.session_state:
    st.session_state["compare_current_index"] = 0

# ============================================================================
# ZONE CONTENT FUNCTIONS
# ============================================================================


def z1_version_selectors():
    """Z1: V1/V2 contract selectors - GEM Section: Select Versions"""
    section_header(GEM_COPY["section_select"], "📄")

    # Auto-load active contract info
    active_id = get_active_contract()
    active_data = get_active_contract_data()

    if active_id:
        st.info(f"📋 Active contract: #{active_id} - {active_data.get('title', 'Unknown')}")

    # Fetch contracts for dropdown
    contracts = get_contracts_for_selector()

    if not contracts:
        st.warning("No contracts available. Upload contracts first.")
        return

    # Build options
    options = {c['label']: c for c in contracts}
    option_labels = ["Select..."] + list(options.keys())

    # Pre-select from session state (e.g., from Portfolio Compare action)
    pre_v1_id = st.session_state.get("v1_id")
    pre_v2_id = st.session_state.get("v2_id")
    v1_default_index = 0
    v2_default_index = 0
    for i, label in enumerate(option_labels):
        if label != "Select..." and label in options:
            if options[label].get('id') == pre_v1_id:
                v1_default_index = i
            if options[label].get('id') == pre_v2_id:
                v2_default_index = i

    # V1 selector
    v1_label = st.selectbox(
        "V1 (Original)",
        options=option_labels,
        index=v1_default_index,
        key="v1_selector_dropdown"
    )

    # V2 selector
    v2_label = st.selectbox(
        "V2 (Revised)",
        options=option_labels,
        index=v2_default_index,
        key="v2_selector_dropdown"
    )

    # Update session state
    if v1_label != "Select...":
        st.session_state["v1_id"] = options[v1_label]['id']
        st.session_state["v1_contract"] = options[v1_label]
    else:
        st.session_state["v1_id"] = None

    if v2_label != "Select...":
        st.session_state["v2_id"] = options[v2_label]['id']
        st.session_state["v2_contract"] = options[v2_label]
    else:
        st.session_state["v2_id"] = None

    # Validation: two distinct versions required
    v1_id = st.session_state["v1_id"]
    v2_id = st.session_state["v2_id"]

    if v1_id and v2_id:
        if v1_id == v2_id:
            st.error("V1 and V2 must be different contracts")
        else:
            st.success("Ready to compare")

            if st.button("🔍 Compare Versions", type="primary", use_container_width=True):
                st.session_state["comparing"] = True
                st.session_state["comparison_result"] = None
                st.session_state["compare_error"] = None
                st.session_state["snapshot_saved"] = False
                st.session_state["comparison_cached"] = False
                st.session_state["comparison_hash"] = None
                st.session_state["compare_current_index"] = 0  # Reset navigation
                st.rerun()
    else:
        st.caption("Select both V1 and V2 to compare")


def z2_comparison_summary():
    """Z2: Comparison summary stats"""
    section_header("Summary", "📊")

    if st.session_state["comparison_result"]:
        result = st.session_state["comparison_result"]

        # Cache indicator
        if st.session_state.get("comparison_cached"):
            st.markdown("""
                <div style="background: linear-gradient(135deg, #1E3A5F 0%, #0D1B2A 100%);
                            border: 1px solid #3B82F6; border-radius: 8px; padding: 8px 12px;
                            margin-bottom: 12px; display: flex; align-items: center; gap: 8px;">
                    <span style="font-size: 18px;">⚡</span>
                    <span style="color: #60A5FA; font-weight: 600;">Cached Result</span>
                    <span style="color: #94A3B8; font-size: 12px; margin-left: auto;">Instant</span>
                </div>
            """, unsafe_allow_html=True)

        similarity = result.get('similarity_score', 0)
        total_changes = result.get('total_changes', 0)
        added = result.get('added_count', 0)
        removed = result.get('removed_count', 0)
        modified = result.get('modified_count', 0)

        st.metric("Similarity", f"{similarity:.1f}%")
        st.metric("Total Changes", total_changes)

        col1, col2 = st.columns(2)
        with col1:
            st.metric("➕ Added", added)
        with col2:
            st.metric("➖ Removed", removed)

        st.metric("✏️ Modified", modified)

        # Snapshot saved indicator
        if st.session_state["snapshot_saved"]:
            st.success("💾 Comparison saved")
    else:
        st.info("Run comparison to see summary")


def z3_match_legend():
    """Z3: Match type legend"""
    section_header("Match Types", "🏷️")

    st.markdown(f"""
    <div style="font-size: 13px; line-height: 1.8;">
        <div>🟢 <strong>{GEM_COPY['match_high']}</strong><br/>
        <span style="color: #94A3B8;">Clauses match closely</span></div>
        <div style="margin-top: 8px;">🟡 <strong>{GEM_COPY['match_likely']}</strong><br/>
        <span style="color: #94A3B8;">Similar content, review recommended</span></div>
        <div style="margin-top: 8px;">🟠 <strong>{GEM_COPY['match_guess']}</strong><br/>
        <span style="color: #94A3B8;">Possible match, verify manually</span></div>
        <div style="margin-top: 8px;">🔵 <strong>{GEM_COPY['match_added']}</strong><br/>
        <span style="color: #94A3B8;">New clause in revised version</span></div>
        <div style="margin-top: 8px;">🔴 <strong>{GEM_COPY['match_removed']}</strong><br/>
        <span style="color: #94A3B8;">Clause removed from original</span></div>
    </div>
    """, unsafe_allow_html=True)


def z4_side_by_side():
    """Z4: Side-by-side clause comparison - Redline-style single-item navigation"""
    section_header(GEM_COPY["section_comparison"], "📝")

    if not st.session_state["comparison_result"]:
        st.info("Select contracts and run comparison to see side-by-side view")
        return

    result = st.session_state["comparison_result"]
    aligned = result.get('aligned_clauses', [])

    if not aligned:
        st.info(GEM_COPY["no_changes"])
        return

    # Get current index
    total = len(aligned)
    idx = st.session_state.get("compare_current_index", 0)
    
    # Clamp to valid range
    if idx >= total:
        idx = total - 1
        st.session_state["compare_current_index"] = idx
    if idx < 0:
        idx = 0
        st.session_state["compare_current_index"] = idx

    pair = aligned[idx]
    match_type = pair.get('match_type', 'Unknown')
    similarity = pair.get('similarity', 0)

    # Match type colors
    match_colors = {
        "High Match": "#10B981",
        "Likely Match": "#F59E0B",
        "Best Guess": "#F97316",
        "Added": "#3B82F6",
        "No Match": "#EF4444"
    }
    match_color = match_colors.get(match_type, "#6B7280")

    # Navigation row
    nav_col1, nav_col2, nav_col3 = st.columns([1, 3, 1])
    
    with nav_col1:
        if st.button("◀ Prev", disabled=(idx == 0), use_container_width=True, key="cmp_prev"):
            st.session_state["compare_current_index"] = idx - 1
            st.rerun()
    
    with nav_col2:
        # Progress indicator
        st.markdown(
            f"<div style='text-align: center; padding: 8px;'>"
            f"<span style='font-size: 14px; color: #94A3B8;'>Clause {idx + 1} of {total}</span>"
            f"</div>",
            unsafe_allow_html=True
        )
    
    with nav_col3:
        if st.button("Next ▶", disabled=(idx >= total - 1), use_container_width=True, key="cmp_next"):
            st.session_state["compare_current_index"] = idx + 1
            st.rerun()

    # Match type badge
    st.markdown(
        f"<div style='background: {match_color}20; border-left: 4px solid {match_color}; "
        f"padding: 8px 12px; border-radius: 4px; margin: 8px 0;'>"
        f"<span style='color: {match_color}; font-weight: 600;'>{match_type}</span>"
        f"<span style='color: #94A3B8; margin-left: 12px;'>Similarity: {similarity:.0%}</span>"
        f"</div>",
        unsafe_allow_html=True
    )

    # Two-column side-by-side layout
    col_v1, col_v2 = st.columns(2)

    with col_v1:
        st.markdown("**V1: Original**")
        v1_clause = pair.get('v1')
        if v1_clause:
            clause_num = v1_clause.get('number') or 'N/A'
            clause_title = v1_clause.get('title') or 'Untitled'
            clause_text = v1_clause.get('text') or v1_clause.get('content') or ''
            severity = v1_clause.get('severity') or ''
            
            st.markdown(f"**{clause_num}. {clause_title}**")
            if severity:
                st.caption(f"Risk: {severity}")
            
            # Display clause text in styled container
            if clause_text:
                st.markdown(
                    f"<div style='background: rgba(30, 41, 59, 0.4); border: 1px solid rgba(255,255,255,0.1); "
                    f"border-radius: 6px; padding: 12px; min-height: 150px; max-height: 300px; "
                    f"overflow-y: auto; font-size: 14px; line-height: 1.6; color: #E2E8F0;'>"
                    f"{clause_text[:1000]}{'...' if len(clause_text) > 1000 else ''}</div>",
                    unsafe_allow_html=True
                )
            else:
                st.info("No clause text available")
        else:
            st.markdown(
                "<div style='background: rgba(239, 68, 68, 0.1); border: 1px dashed #EF4444; "
                "border-radius: 6px; padding: 24px; text-align: center; color: #EF4444;'>"
                "<strong>— Removed in V2 —</strong></div>",
                unsafe_allow_html=True
            )

    with col_v2:
        st.markdown("**V2: Revised**")
        v2_clause = pair.get('v2')
        if v2_clause:
            clause_num = v2_clause.get('number') or 'N/A'
            clause_title = v2_clause.get('title') or 'Untitled'
            clause_text = v2_clause.get('text') or v2_clause.get('content') or ''
            severity = v2_clause.get('severity') or ''
            
            st.markdown(f"**{clause_num}. {clause_title}**")
            if severity:
                st.caption(f"Risk: {severity}")
            
            # Display clause text in styled container
            if clause_text:
                st.markdown(
                    f"<div style='background: rgba(30, 41, 59, 0.4); border: 1px solid rgba(255,255,255,0.1); "
                    f"border-radius: 6px; padding: 12px; min-height: 150px; max-height: 300px; "
                    f"overflow-y: auto; font-size: 14px; line-height: 1.6; color: #E2E8F0;'>"
                    f"{clause_text[:1000]}{'...' if len(clause_text) > 1000 else ''}</div>",
                    unsafe_allow_html=True
                )
            else:
                st.info("No clause text available")
        else:
            st.markdown(
                "<div style='background: rgba(59, 130, 246, 0.1); border: 1px dashed #3B82F6; "
                "border-radius: 6px; padding: 24px; text-align: center; color: #3B82F6;'>"
                "<strong>— Added in V2 —</strong></div>",
                unsafe_allow_html=True
            )

    # Jump to dropdown
    st.markdown("---")
    col_jump, col_filter = st.columns([2, 2])
    
    with col_jump:
        # Build jump options - show match type for each
        jump_options = [f"{i+1}. {aligned[i].get('match_type', '?')}" for i in range(total)]
        selected = st.selectbox(
            "Jump to",
            options=range(total),
            format_func=lambda i: jump_options[i],
            index=idx,
            key="cmp_jump"
        )
        if selected != idx:
            st.session_state["compare_current_index"] = selected
            st.rerun()
    
    with col_filter:
        # Quick filter by match type
        match_types = list(set(p.get('match_type', 'Unknown') for p in aligned))
        match_types.sort()
        filter_type = st.selectbox(
            "Filter by type",
            options=["All"] + match_types,
            key="cmp_filter"
        )
        if filter_type != "All":
            # Find first match of this type
            for i, p in enumerate(aligned):
                if p.get('match_type') == filter_type:
                    if i != idx:
                        st.session_state["compare_current_index"] = i
                        st.rerun()
                    break


def z5_business_impact():
    """Z5: Business impact narrative - GEM Section: Business Impact"""
    section_header(GEM_COPY["section_impact"], "💼")

    if not st.session_state["comparison_result"]:
        st.info("Run comparison for impact analysis")
        return

    result = st.session_state["comparison_result"]

    added = result.get('added_count', 0)
    removed = result.get('removed_count', 0)
    modified = result.get('modified_count', 0)
    old_risk = result.get('v1_risk', 'Unknown')
    new_risk = result.get('v2_risk', 'Unknown')

    # GEM impact narrative template
    if added == 0 and removed == 0 and modified == 0:
        st.info(GEM_COPY["no_changes"])
    else:
        narrative = GEM_COPY["impact_template"].format(
            added=added,
            removed=removed,
            modified=modified,
            old_risk=old_risk,
            new_risk=new_risk
        )
        st.markdown(narrative)

    # Risk delta visualization
    risk_delta = result.get('risk_delta', [])
    if risk_delta:
        st.markdown("**Risk Changes by Category:**")
        for delta in risk_delta[:6]:
            cat = delta.get('category', 'Unknown')
            v1_risk = delta.get('v1_risk', 'N/A')
            v2_risk = delta.get('v2_risk', 'N/A')

            if v1_risk != v2_risk:
                st.markdown(f"- **{cat}**: {v1_risk} → {v2_risk}")


def z6_export_actions():
    """Z6: Export buttons"""
    section_header("Export", "📤")

    if not st.session_state["comparison_result"]:
        st.info("Run comparison to enable exports")
        return

    result = st.session_state["comparison_result"]
    v1_title = st.session_state["v1_contract"].get('title', 'V1') if st.session_state["v1_contract"] else 'V1'
    v2_title = st.session_state["v2_contract"].get('title', 'V2') if st.session_state["v2_contract"] else 'V2'

    # Export DOCX
    if st.button("📄 Export DOCX", use_container_width=True):
        try:
            docx_buffer = export_comparison_docx(result, v1_title, v2_title)
            st.download_button(
                label="Download DOCX",
                data=docx_buffer,
                file_name=f"comparison_{st.session_state["v1_id"]}_vs_{st.session_state["v2_id"]}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="docx_download"
            )
            toast_success("DOCX export ready")
        except Exception as e:
            toast_error(f"DOCX export failed: {str(e)}")

    # Export JSON
    if st.button("📋 Export JSON", use_container_width=True):
        try:
            json_str = export_comparison_json(result)
            st.download_button(
                label="Download JSON",
                data=json_str,
                file_name=f"comparison_{st.session_state["v1_id"]}_vs_{st.session_state["v2_id"]}.json",
                mime="application/json",
                key="json_download"
            )
            toast_success("JSON export ready")
        except Exception as e:
            toast_error(f"JSON export failed: {str(e)}")

    # ========================================================================
    # PHASE 4E: COMPARE V3 (EARLY PREVIEW) BUTTON
    # ========================================================================
    st.markdown("---")
    section_header(GEM_COPY["compare_v3.title"], "🧪")
    st.caption(GEM_COPY["compare_v3.description"])

    # Compare v3 button - requires valid comparison first
    can_compare_v3 = st.session_state["v1_id"] and st.session_state["v2_id"]
    if can_compare_v3:
        # Derive TRUST inputs from active contract (reuse v2 pattern)
        trust_contract_id = st.session_state["v1_id"]  # Use V1 as primary
        trust_last_analyzed = None  # Will be populated from contract data if available
        trust_risk_level = None

        # Try to get contract data for TRUST evaluation
        if st.session_state["v1_contract"]:
            trust_last_analyzed = st.session_state["v1_contract"].get('last_analyzed_at')
            trust_risk_level = st.session_state["v1_contract"].get('risk_level')

        if st.button("🧠 Run Compare v3", use_container_width=True, key="compare_v3_btn", type="secondary", help="Run Compare v3 pipeline with real SAE/ERCE/BIRL/FAR intelligence"):
            # TRUST PRECHECK - runs IMMEDIATELY before backend call
            clear_trust_banner_dismissal()

            trust = run_trust_precheck(
                contract_id=trust_contract_id,
                last_analyzed_at=trust_last_analyzed,
                risk_level=trust_risk_level,
            )
            render_trust_banner(trust.to_dict())

            # Set running flag and rerun
            st.session_state["compare_v3_running"] = True
            st.rerun()
    else:
        st.info("Select V1 and V2 contracts first")

    # Display Compare v3 results if available
    if st.session_state["compare_v3_result"]:
        v3_result = st.session_state["compare_v3_result"]
        if v3_result.get('success'):
            data = v3_result.get('data', {})
            meta = data.get('_meta', {})

            # Phase 4F: Intelligence status display
            engine_version = meta.get('engine_version', 'Unknown')
            intelligence_active = v3_result.get('intelligence_active', False)
            pipeline_status = meta.get('pipeline_status', 'UNKNOWN')

            if intelligence_active:
                st.success(f"Compare v3 Engine v{engine_version} - {GEM_COPY['p4f.intelligence_active']}")
            elif pipeline_status.startswith('PARTIAL'):
                st.warning(f"Compare v3 Engine v{engine_version} - {GEM_COPY['p4f.intelligence_partial']}")
            else:
                st.info(f"Compare v3 Engine v{engine_version} - {GEM_COPY['p4f.intelligence_placeholder']}")

            # GEM 4F Truncation Limits
            SAE_MAX = 10
            ERCE_MAX = 15
            BIRL_MAX = 5
            FAR_MAX = 5

            # SAE Matches (max 10)
            sae_data = data.get('sae_matches', [])
            with st.expander(GEM_COPY["p4f.sae_title"], expanded=False):
                for match in sae_data[:SAE_MAX]:
                    conf = match.get('match_confidence', 'N/A')
                    score = match.get('similarity_score', 0)
                    st.markdown(f"- Clause {match.get('v1_clause_id')} ↔ {match.get('v2_clause_id')}: **{conf}** ({score:.0%})")
                if len(sae_data) > SAE_MAX:
                    st.caption(GEM_COPY["p4f.truncation_label"].format(n=len(sae_data) - SAE_MAX, type="pairs"))

            # ERCE Results (max 15)
            erce_data = data.get('erce_results', [])
            with st.expander(GEM_COPY["p4f.erce_title"], expanded=False):
                for risk in erce_data[:ERCE_MAX]:
                    cat = risk.get('risk_category', 'N/A')
                    conf = risk.get('confidence', 0)
                    pattern = risk.get('pattern_ref', '')
                    pattern_text = f" [{pattern}]" if pattern else ""
                    st.markdown(f"- Pair {risk.get('clause_pair_id')}: **{cat}**{pattern_text} (confidence: {conf:.0%})")
                if len(erce_data) > ERCE_MAX:
                    st.caption(GEM_COPY["p4f.truncation_label"].format(n=len(erce_data) - ERCE_MAX, type="risks"))

            # BIRL Narratives (max 5)
            birl_data = data.get('birl_narratives', [])
            with st.expander(GEM_COPY["p4f.birl_title"], expanded=False):
                for narrative in birl_data[:BIRL_MAX]:
                    st.markdown(f"**Pair {narrative.get('clause_pair_id')}**: {narrative.get('narrative')}")
                    dims = ', '.join(narrative.get('impact_dimensions', []))
                    st.caption(f"{GEM_COPY['p4f.narrative_dimension']}: {dims}")
                if len(birl_data) > BIRL_MAX:
                    st.caption(GEM_COPY["p4f.truncation_label"].format(n=len(birl_data) - BIRL_MAX, type="narratives"))

            # FAR Results (max 5)
            far_data = data.get('flowdown_gaps', [])
            if far_data:
                with st.expander(GEM_COPY["p4f.far_title"], expanded=False):
                    for gap in far_data[:FAR_MAX]:
                        severity = gap.get('severity', 'N/A')
                        gap_type = gap.get('gap_type', 'N/A')
                        st.markdown(f"- **{gap_type}** ({severity}): {gap.get('upstream_value', '')} → {gap.get('downstream_value', '')}")
                        st.caption(f"Recommendation: {gap.get('recommendation', 'N/A')}")
                    if len(far_data) > FAR_MAX:
                        st.caption(GEM_COPY["p4f.truncation_label"].format(n=len(far_data) - FAR_MAX, type="gaps"))
            else:
                with st.expander(GEM_COPY["p4f.far_title"], expanded=False):
                    st.info("No flowdown gaps detected (no contract relationships found)")
        else:
            # Use GEM error card instead of st.error
            render_error_card(v3_result, module_name="Compare v3")


def z7_footer():
    """Z7: Footer"""
    st.markdown("---")
    st.caption(f"Compare v2.1 (Phase 4F) | {datetime.now().strftime('%Y-%m-%d %H:%M')}")


# ============================================================================
# RENDER ZONE LAYOUT
# ============================================================================

zone_layout(
    z1=z1_version_selectors,
    z2=z2_comparison_summary,
    z3=z3_match_legend,
    z4=z4_side_by_side,
    z5=z5_business_impact,
    z6=z6_export_actions,
    z7=z7_footer,
    z7_system_status=True,
    api_healthy=api_healthy,
    db_healthy=db_healthy,
    ai_healthy=ai_healthy
)


# ============================================================================
# v2: RUN COMPARISON WITH PROGRESS OVERLAY
# ============================================================================

if st.session_state["comparing"]:
    v1_id = st.session_state["v1_id"]
    v2_id = st.session_state["v2_id"]

    if not v1_id or not v2_id or v1_id == v2_id:
        st.error("Two distinct contract versions required")
        st.session_state["comparing"] = False
    else:
        # Perform comparison via API
        result, error = api_call_with_spinner(
            endpoint="/api/compare",
            method="POST",
            data={
                'v1_contract_id': v1_id,
                'v2_contract_id': v2_id,
                'include_recommendations': True
            },
            spinner_message="Comparing versions... This may take 30-60 seconds.",
            success_message="Comparison complete!",
            timeout=600
        )

        if error:
            # 404 guard - contract no longer exists
            if isinstance(error, dict) and error.get('status_code') == 404:
                st.error("Selected contract no longer exists. Please reselect.")
                st.session_state["v1_id"] = None
                st.session_state["v2_id"] = None
                st.session_state["comparing"] = False
                st.stop()

            # Phase 4B: Handle structured error response
            if isinstance(error, dict):
                error_key = error.get('error_message_key', 'compare.internal_failure')
                user_message = GEM_COPY.get(error_key, GEM_COPY["compare.internal_failure"])
            else:
                user_message = str(error)

            st.session_state["compare_error"] = user_message
            st.error(f"Comparison Failed: {user_message}")
        else:
            # Process result and build aligned clauses
            comparison_data = result if isinstance(result, dict) else {}

            # Capture cache status from API response
            is_cached = comparison_data.get('cached', False)
            st.session_state["comparison_cached"] = is_cached
            st.session_state["comparison_hash"] = comparison_data.get('comparison_hash')

            # Show toast for cached results
            if is_cached:
                toast_info("Result from cache (instant)")

            # Get clause lists for alignment
            v1_clauses = get_contract_clauses(v1_id)
            v2_clauses = get_contract_clauses(v2_id)

            # Simple clause alignment
            from difflib import SequenceMatcher

            aligned_clauses = []
            v2_matched = set()
            added_count = 0
            removed_count = 0
            modified_count = 0

            for v1_clause in v1_clauses:
                best_match = None
                best_score = 0.0
                best_idx = -1

                for idx, v2_clause in enumerate(v2_clauses):
                    if idx in v2_matched:
                        continue

                    # Score by number + title similarity
                    number_match = 1.0 if v1_clause.get('number') == v2_clause.get('number') else 0.0
                    title_sim = SequenceMatcher(
                        None,
                        (v1_clause.get('title') or '').lower(),
                        (v2_clause.get('title') or '').lower()
                    ).ratio()

                    score = (number_match * 0.4) + (title_sim * 0.6)

                    if score > best_score:
                        best_score = score
                        best_match = v2_clause
                        best_idx = idx

                # Determine match type
                if best_score >= 0.85:
                    match_type = "High Match"
                elif best_score >= 0.60:
                    match_type = "Likely Match"
                    modified_count += 1
                elif best_score >= 0.40:
                    match_type = "Best Guess"
                    modified_count += 1
                else:
                    match_type = "No Match"
                    removed_count += 1
                    best_match = None

                if best_match and best_idx >= 0:
                    v2_matched.add(best_idx)

                aligned_clauses.append({
                    'v1': v1_clause,
                    'v2': best_match,
                    'match_type': match_type,
                    'similarity': best_score
                })

            # Add unmatched V2 clauses as "Added"
            for idx, v2_clause in enumerate(v2_clauses):
                if idx not in v2_matched:
                    aligned_clauses.append({
                        'v1': None,
                        'v2': v2_clause,
                        'match_type': "Added",
                        'similarity': 0.0
                    })
                    added_count += 1

            # Calculate similarity score
            total_clauses = max(len(v1_clauses), len(v2_clauses), 1)
            high_matches = sum(1 for a in aligned_clauses if a['match_type'] == 'High Match')
            similarity_score = (high_matches / total_clauses) * 100

            # Build result object
            st.session_state["comparison_result"] = {
                'v1_contract_id': v1_id,
                'v2_contract_id': v2_id,
                'aligned_clauses': aligned_clauses,
                'similarity_score': similarity_score,
                'total_changes': added_count + removed_count + modified_count,
                'added_count': added_count,
                'removed_count': removed_count,
                'modified_count': modified_count,
                'v1_risk': comparison_data.get('v1_risk', 'Unknown'),
                'v2_risk': comparison_data.get('v2_risk', 'Unknown'),
                'risk_delta': comparison_data.get('risk_delta', []),
                'created_at': datetime.now().isoformat()
            }

            # Save ComparisonSnapshot
            v1_snapshot_id = get_latest_snapshot_id(v1_id)
            v2_snapshot_id = get_latest_snapshot_id(v2_id)

            snapshot_result = save_comparison_snapshot_api(
                v1_id=v1_id,
                v2_id=v2_id,
                v1_snapshot_id=v1_snapshot_id,
                v2_snapshot_id=v2_snapshot_id,
                similarity_score=similarity_score,
                changed_clauses=[
                    {
                        'v1_number': a['v1'].get('number') if a['v1'] else None,
                        'v2_number': a['v2'].get('number') if a['v2'] else None,
                        'match_type': a['match_type']
                    }
                    for a in aligned_clauses if a['match_type'] != 'High Match'
                ],
                risk_delta=comparison_data.get('risk_delta', [])
            )

            if snapshot_result:
                st.session_state["snapshot_saved"] = True

        # Reset comparison state
        st.session_state["comparing"] = False

        if st.session_state["comparison_result"]:
            st.rerun()


# ============================================================================
# PHASE 4E: COMPARE V3 EXECUTION
# ============================================================================

if st.session_state["compare_v3_running"]:
    v1_id = st.session_state["v1_id"]
    v2_id = st.session_state["v2_id"]

    if not v1_id or not v2_id:
        render_error_card({'error_message_key': 'compare.internal_failure'}, module_name="Compare v3")
        st.session_state["compare_v3_running"] = False
    else:
        # Use api_call_with_spinner (same pattern as v2)
        result, error = api_call_with_spinner(
            endpoint="/api/compare-v3",
            method="POST",
            data={'v1_contract_id': v1_id, 'v2_contract_id': v2_id},
            spinner_message="Running Compare v3 pipeline (SAE/ERCE/BIRL)...",
            timeout=30
        )

        if error:
            # 404 guard - contract no longer exists
            if isinstance(error, dict) and error.get('status_code') == 404:
                st.error("Selected contract no longer exists. Please reselect.")
                st.session_state["v1_id"] = None
                st.session_state["v2_id"] = None
                st.session_state["comparing"] = False
                st.stop()

            # Route through GEM error card
            render_error_card(error, module_name="Compare v3")
            st.session_state["compare_v3_result"] = None
        else:
            st.session_state["compare_v3_result"] = result

        st.session_state["compare_v3_running"] = False
        st.rerun()
